#!/usr/bin/env python3
"""
Generate comprehensive Redis 3-Node HA Cluster (with Sentinel) Datadog Monitoring Dashboard Document.
Topology: 3 nodes, each running Redis Server (6379) + Redis Sentinel (26379)
          1 Master + 2 Replicas, Sentinel quorum = 2
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0xA4, 0x1E, 0x22)  # Redis red

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 2'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Redis 8.2 HA Cluster (3-Node + Sentinel)\nDatadog Monitoring & Dashboard Strategy")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xA4, 0x1E, 0x22)

doc.add_paragraph("")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Enterprise Monitoring Architecture & Dashboard Design")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")
doc.add_paragraph("")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Version: 1.0",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "Classification: Internal / Infrastructure Team",
    "Author: Monitoring Architecture Team",
]:
    meta.add_run(line + "\n").font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Executive Summary",
    "2. Monitoring Architecture Overview",
    "3. Datadog Integration Setup",
    "4. Dashboard #1 — Cluster Health & Replication Overview",
    "5. Dashboard #2 — Sentinel Monitoring & Failover Tracking",
    "6. Dashboard #3 — Performance & Throughput (Commands)",
    "7. Dashboard #4 — Memory Management & Eviction",
    "8. Dashboard #5 — Persistence (RDB & AOF)",
    "9. Dashboard #6 — Client & Connection Monitoring",
    "10. Dashboard #7 — Key Space & Data Analytics",
    "11. Dashboard #8 — Latency & Slow Queries",
    "12. Dashboard #9 — Node Resource Utilization (CPU, Disk, Network)",
    "13. Dashboard #10 — Replication Deep-Dive",
    "14. Dashboard #11 — Security & ACL Audit",
    "15. Dashboard #12 — Restart & Availability Tracking",
    "16. Dashboard #13 — TLS / Certificate Monitoring",
    "17. Dashboard #14 — Executive / SLA Summary",
    "18. Alerting Strategy & Escalation Matrix",
    "19. Datadog Monitor Definitions (Complete List)",
    "20. Runbook References",
    "21. Appendix — Full Metric Reference",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document defines the comprehensive Datadog monitoring strategy for our production "
    "3-node Redis HA cluster with Sentinel. Each of the 3 nodes runs a Redis server instance "
    "(port 6379) and a Redis Sentinel instance (port 26379). The topology consists of 1 Master "
    "and 2 Replicas with a Sentinel quorum of 2 for automatic failover."
)
doc.add_paragraph(
    "The strategy covers 14 purpose-built dashboards, 45+ monitors, a tiered alerting framework, "
    "and runbook references — providing complete operational visibility across data operations, "
    "replication, persistence, memory management, security, and Sentinel-managed failover."
)

doc.add_heading("1.1 Cluster Topology", level=2)
add_table(
    ["Property", "Value"],
    [
        ["Cluster Size", "3 Nodes (1 Master + 2 Replicas)"],
        ["Redis Version", "8.2.2"],
        ["Node 1 (Master)", "redis-node-1:6379 + sentinel-1:26379"],
        ["Node 2 (Replica)", "redis-node-2:6379 + sentinel-2:26379"],
        ["Node 3 (Replica)", "redis-node-3:6379 + sentinel-3:26379"],
        ["Sentinel Quorum", "2 of 3"],
        ["Install Path", "/opt/cached/current"],
        ["Data Path", "/opt/cached/current/data"],
        ["Log Path", "/opt/cached/current/logs"],
        ["Persistence", "RDB snapshots + AOF (appendonly)"],
        ["ACL", "Enabled (users.acl)"],
        ["Monitoring Tool", "Datadog Agent + Redis Integration"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  2. MONITORING ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Monitoring Architecture Overview", level=1)

doc.add_heading("2.1 Data Collection Layers", level=2)
add_table(
    ["Layer", "Source", "Collection Method", "Frequency"],
    [
        ["Redis Metrics", "Redis INFO command", "Datadog Redis Integration", "15s"],
        ["Sentinel Metrics", "Sentinel INFO + SENTINEL commands", "Datadog Redis Integration (2nd instance)", "15s"],
        ["System Metrics", "Host OS (CPU, Memory, Disk, Network)", "Datadog Agent", "15s"],
        ["Logs", "Redis logs + Sentinel logs", "Datadog Log Agent", "Real-time"],
        ["Slow Log", "Redis SLOWLOG command", "Datadog Custom Check", "30s"],
        ["Process Checks", "redis-server + redis-sentinel processes", "Datadog Process Check", "30s"],
        ["Custom Checks", "redis-cli INFO, CLIENT LIST, ACL LIST", "Datadog Custom Check", "60s"],
        ["Synthetic Checks", "SET/GET round-trip probe", "Datadog Synthetic Monitor", "60s"],
    ]
)

doc.add_heading("2.2 Tagging Strategy", level=2)
add_table(
    ["Tag Key", "Example Values", "Purpose"],
    [
        ["env", "production, staging, qa", "Environment isolation"],
        ["redis_cluster", "prod-redis-ha-01", "Cluster identification"],
        ["redis_node", "redis-node-1, redis-node-2, redis-node-3", "Per-node drill-down"],
        ["redis_role", "master, replica", "Role-based filtering"],
        ["redis_port", "6379, 26379", "Service port (Redis vs Sentinel)"],
        ["service", "redis, redis-sentinel", "Service catalog"],
        ["team", "platform-sre, cache-team", "Team ownership"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  3. DATADOG INTEGRATION SETUP
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Datadog Integration Setup", level=1)

doc.add_heading("3.1 Redis Server Integration (Each Node)", level=2)
doc.add_paragraph("File: /etc/datadog-agent/conf.d/redisdb.d/conf.yaml")
add_code("""init_config:

instances:
  # Redis Server (port 6379)
  - host: localhost
    port: 6379
    password: <VAULT_SECRET>
    keys:
      - "app:*"
      - "session:*"
      - "cache:*"
    warn_on_missing_keys: false
    command_stats: true
    slowlog-max-len: 128
    tags:
      - env:production
      - redis_cluster:prod-redis-ha-01
      - service:redis

  # Redis Sentinel (port 26379)
  - host: localhost
    port: 26379
    tags:
      - env:production
      - redis_cluster:prod-redis-ha-01
      - service:redis-sentinel""")

doc.add_heading("3.2 Log Collection", level=2)
doc.add_paragraph("File: /etc/datadog-agent/conf.d/redisdb.d/conf.yaml (logs section)")
add_code("""logs:
  # Redis Server logs
  - type: file
    path: /opt/cached/current/logs/redis.log
    service: redis
    source: redis
    tags:
      - env:production
      - redis_cluster:prod-redis-ha-01

  # Sentinel logs
  - type: file
    path: /opt/cached/current/logs/sentinel.log
    service: redis-sentinel
    source: redis
    tags:
      - env:production
      - redis_cluster:prod-redis-ha-01""")

doc.add_heading("3.3 Dedicated Monitoring User (ACL)", level=2)
doc.add_paragraph("Add a read-only monitoring user to /opt/cached/current/conf/users.acl:")
add_code("""# Datadog monitoring user - read-only access
user datadog_monitor on >SecureMonitorPass123 ~* &* +info +ping +config|get +client|list +slowlog|get +slowlog|len +dbsize +lastsave +command|count +latency|latest""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  4. DASHBOARD #1 — CLUSTER HEALTH & REPLICATION OVERVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Dashboard #1 — Cluster Health & Replication Overview", level=1)
doc.add_paragraph(
    "The primary landing dashboard. Shows node status, master/replica roles, replication health, "
    "and overall cluster state. First screen for NOC/SRE teams."
)

doc.add_heading("4.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Cluster Status", "Check Status", "redis.can_connect per node", "All 3 nodes reachable?"],
        ["Node 1 Role", "Query Value", "redis.info.role{node:redis-node-1}", "master or replica"],
        ["Node 2 Role", "Query Value", "redis.info.role{node:redis-node-2}", "master or replica"],
        ["Node 3 Role", "Query Value", "redis.info.role{node:redis-node-3}", "master or replica"],
        ["Master Node Indicator", "Query Value (highlighted)", "Node where role == master", "Who is the current master?"],
        ["Connected Replicas", "Query Value", "redis.replication.connected_slaves", "Expected: 2"],
        ["Replication Offset Lag", "Timeseries (per replica)", "master_repl_offset - slave_repl_offset", "Replication delay in bytes"],
        ["Replication Link Status", "Status Widget", "master_link_status per replica (up/down)", "Replication connectivity"],
        ["Uptime per Node", "Query Value (3x)", "redis.info.uptime_in_seconds", "Current uptime"],
        ["Total Keys", "Query Value", "redis.keys by {redis_node}", "Key count per node"],
        ["Used Memory per Node", "Timeseries (3 lines)", "redis.mem.used by {redis_node}", "Memory consumption"],
        ["Connected Clients", "Query Value (3x)", "redis.net.clients by {redis_node}", "Clients per node"],
        ["Commands Processed/sec", "Timeseries", "rate(redis.net.commands) by {redis_node}", "Throughput per node"],
        ["Keyspace Hit Rate", "Query Value", "hits / (hits + misses) * 100", "Cache effectiveness"],
        ["Redis Version", "Query Value", "redis.info.redis_version", "Running version"],
    ]
)

doc.add_heading("4.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Node Unreachable", "redis.can_connect == 0 for any node > 30s", "P1 - Critical", "Page on-call, check host & process"],
        ["No Master Detected", "No node with role == master for 1 min", "P1 - Critical", "Sentinel failover failed, manual intervention"],
        ["Replica Disconnected", "connected_slaves < 2 for > 1 min", "P1 - Critical", "Check replica & network"],
        ["Replication Lag High", "repl_offset_diff > 10MB for > 2 min", "P2 - Warning", "Check replica I/O & network"],
        ["Master Link Down", "master_link_status == down for > 30s", "P1 - Critical", "Replica can't reach master"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  5. DASHBOARD #2 — SENTINEL MONITORING & FAILOVER
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Dashboard #2 — Sentinel Monitoring & Failover Tracking", level=1)
doc.add_paragraph(
    "Dedicated monitoring for all 3 Sentinel instances. Tracks Sentinel health, quorum status, "
    "failover events, and master election history."
)

doc.add_heading("5.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Sentinel 1 Status", "Check Status", "redis.can_connect{port:26379,node:1}", "Sentinel-1 alive?"],
        ["Sentinel 2 Status", "Check Status", "redis.can_connect{port:26379,node:2}", "Sentinel-2 alive?"],
        ["Sentinel 3 Status", "Check Status", "redis.can_connect{port:26379,node:3}", "Sentinel-3 alive?"],
        ["Active Sentinels Count", "Query Value", "SENTINEL master mymaster -> num-sentinels", "Should be 3"],
        ["Quorum Achieved?", "Status Widget", "num-sentinels >= quorum (2)", "Green if quorum met"],
        ["Current Master (Sentinel View)", "Query Value", "SENTINEL get-master-addr-by-name", "Master IP:port from Sentinel"],
        ["Failover Event Timeline", "Event Timeline", "Log: +switch-master events", "When did failovers happen?"],
        ["Failover Count (30d)", "Query Value", "Count of +switch-master in 30 days", "Failover frequency"],
        ["SDOWN Events", "Event Timeline", "Log: +sdown (subjective down)", "Single sentinel down detection"],
        ["ODOWN Events", "Event Timeline", "Log: +odown (objective down)", "Quorum-agreed down detection"],
        ["Sentinel Voted Leader", "Event Stream", "Log: +elected-leader", "Who led the failover?"],
        ["Failover Duration", "Timeseries", "Time from ODOWN to +switch-master", "How fast is failover?"],
        ["Sentinel Config Epoch", "Query Value per sentinel", "sentinel_master_config_epoch", "Config consistency"],
        ["Sentinel Pending Commands", "Timeseries", "sentinel pending-scripts", "Script queue health"],
    ]
)

doc.add_heading("5.2 Sentinel Log Patterns to Capture", level=2)
add_code("""# Key Sentinel log events for Datadog Log Pipeline:
+sdown         # Subjectively Down — one sentinel thinks master is down
+odown         # Objectively Down — quorum agrees master is down
-odown         # Objectively Up — master recovered
+switch-master # FAILOVER COMPLETE — master changed
+elected-leader # This sentinel leads failover
+failover-state-reconf-slaves  # Reconfiguring replicas after failover
-sdown         # Subjectively Up — sentinel sees node recover
+slave         # New replica registered
+sentinel      # New sentinel registered""")

doc.add_heading("5.3 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Sentinel Down", "redis.can_connect{port:26379} == 0 for any sentinel > 30s", "P1 - Critical", "Restart sentinel, check host"],
        ["Quorum Lost", "Active sentinels < 2", "P1 - Critical", "FAILOVER IMPOSSIBLE — immediate investigation"],
        ["Failover Occurred", "Log: +switch-master detected", "P1 - Critical", "Verify new master, check old master"],
        ["ODOWN Detected", "Log: +odown event", "P1 - Critical", "Master confirmed down by quorum"],
        ["Frequent Failovers", "> 2 failovers in 1 hour", "P1 - Critical", "Flapping — investigate root cause"],
        ["Sentinel Config Mismatch", "config_epoch differs across sentinels", "P2 - Warning", "Sentinel state inconsistency"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  6. DASHBOARD #3 — PERFORMANCE & THROUGHPUT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Dashboard #3 — Performance & Throughput (Commands)", level=1)
doc.add_paragraph(
    "Real-time and historical view of Redis command execution: throughput, command types, "
    "cache hit/miss ratios, and operations per second."
)

doc.add_heading("6.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Commands/sec (cluster total)", "Query Value", "sum(rate(redis.net.commands{*}))", "Total throughput"],
        ["Commands/sec per Node", "Timeseries (3 lines)", "rate(redis.net.commands) by {redis_node}", "Per-node throughput"],
        ["Command Types Distribution", "Pie Chart", "redis.command.calls by {command}", "GET vs SET vs HGET etc."],
        ["Top 20 Commands by Calls/sec", "Top List", "rate(redis.command.calls) by {command} top 20", "Hotspot commands"],
        ["Top 20 Commands by Usec/call", "Top List", "redis.command.usec_per_call by {command} top 20", "Slowest commands"],
        ["Keyspace Hit Rate %", "Query Value (large)", "redis.stats.keyspace_hits / (hits + misses) * 100", "Cache effectiveness (target: >95%)"],
        ["Hit Rate Trend", "Timeseries", "keyspace_hit_rate over time", "Cache health trend"],
        ["Keyspace Hits/sec", "Timeseries", "rate(redis.stats.keyspace_hits)", "Successful lookups/sec"],
        ["Keyspace Misses/sec", "Timeseries", "rate(redis.stats.keyspace_misses)", "Cache misses/sec"],
        ["Ops Latency by Command", "Timeseries (P50/P95/P99)", "redis.command.usec_per_call by {command}", "Command-level latency"],
        ["Pipeline Commands", "Timeseries", "pipelined vs non-pipelined ratio", "Pipeline utilization"],
        ["Throughput Heatmap", "Heatmap", "rate(redis.net.commands) by hour", "Traffic pattern analysis"],
    ]
)

doc.add_heading("6.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Throughput Drop > 50%", "Anomaly: commands/sec drops > 50% for 5 min", "P1 - Critical", "Check client connectivity"],
        ["Hit Rate < 90%", "keyspace_hit_rate < 90% for 10 min", "P2 - Warning", "Review TTL & eviction policy"],
        ["Hit Rate < 80%", "keyspace_hit_rate < 80% for 5 min", "P1 - Critical", "Cache is ineffective, investigate"],
        ["Zero Throughput", "commands/sec == 0 for 5 min", "P1 - Critical", "Redis not serving requests"],
        ["Command Latency Spike", "usec_per_call P99 > 1000us for 5 min", "P2 - Warning", "Check for blocking commands"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  7. DASHBOARD #4 — MEMORY MANAGEMENT & EVICTION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Dashboard #4 — Memory Management & Eviction", level=1)
doc.add_paragraph(
    "Critical for Redis — monitors memory usage, fragmentation, eviction activity, "
    "and memory-related risks that could lead to data loss or OOM."
)

doc.add_heading("7.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Used Memory per Node", "Timeseries (3 lines)", "redis.mem.used by {redis_node}", "Memory consumption trend"],
        ["Max Memory per Node", "Timeseries (3 lines)", "redis.mem.maxmemory by {redis_node}", "Configured limit"],
        ["Memory Usage %", "Gauge per node", "redis.mem.used / maxmemory * 100", "Proximity to limit"],
        ["Memory RSS per Node", "Timeseries", "redis.mem.rss by {redis_node}", "OS-level memory (actual)"],
        ["Memory Fragmentation Ratio", "Timeseries per node", "redis.mem.fragmentation_ratio", "Target: 1.0-1.5; >1.5 = fragmented"],
        ["Peak Memory", "Query Value per node", "redis.mem.peak", "Historical max memory"],
        ["Evicted Keys/sec", "Timeseries per node", "rate(redis.stats.evicted_keys)", "Keys removed by eviction policy"],
        ["Evicted Keys Total", "Query Value per node", "redis.stats.evicted_keys", "Cumulative evictions"],
        ["Eviction Policy", "Query Value", "maxmemory-policy config", "allkeys-lru / volatile-lru etc."],
        ["Memory Breakdown", "Stacked Area", "dataset, overhead, repl_buffer, clients, aof_buffer", "What consumes memory"],
        ["Lua Memory", "Timeseries", "redis.mem.lua", "Lua script memory usage"],
        ["Expired Keys/sec", "Timeseries", "rate(redis.stats.expired_keys)", "TTL-based key removal rate"],
        ["Expired Keys Total", "Query Value", "redis.stats.expired_keys", "Cumulative expirations"],
    ]
)

doc.add_heading("7.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Memory > 85% of Max", "used / maxmemory > 0.85 for 5 min", "P2 - Warning", "Review keys, increase maxmemory or scale"],
        ["Memory > 95% of Max", "used / maxmemory > 0.95 for 2 min", "P1 - Critical", "Evictions imminent or active, immediate action"],
        ["Evictions Active", "evicted_keys rate > 0 for 5 min", "P2 - Warning", "Data loss from eviction, investigate"],
        ["Fragmentation > 1.5", "fragmentation_ratio > 1.5 for 30 min", "P2 - Warning", "Consider restart or MEMORY PURGE"],
        ["Fragmentation < 1.0", "fragmentation_ratio < 1.0 for 10 min", "P1 - Critical", "Swapping to disk — CRITICAL performance impact"],
        ["RSS > Physical Memory 80%", "mem.rss > 80% host memory", "P1 - Critical", "OOM kill risk"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  8. DASHBOARD #5 — PERSISTENCE (RDB & AOF)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Dashboard #5 — Persistence (RDB & AOF)", level=1)
doc.add_paragraph(
    "Monitors Redis persistence mechanisms — RDB snapshots and AOF (Append Only File). "
    "Critical for data durability and recovery capability."
)

doc.add_heading("8.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Last RDB Save Status", "Status Widget per node", "rdb_last_bgsave_status (ok/err)", "Last snapshot success?"],
        ["Time Since Last RDB Save", "Query Value per node", "redis.rdb.last_save_time (seconds ago)", "Data loss window"],
        ["RDB Save Duration", "Timeseries", "redis.rdb.last_bgsave_time_sec", "Snapshot duration trend"],
        ["RDB Changes Since Last Save", "Timeseries", "redis.rdb.changes_since_last_save", "Pending unsaved changes"],
        ["RDB File Size", "Timeseries per node", "Custom: du -b dump.rdb", "Snapshot size trend"],
        ["BGSAVE in Progress", "Status Widget", "rdb_bgsave_in_progress (1/0)", "Currently saving?"],
        ["AOF Enabled", "Status Widget per node", "aof_enabled (1/0)", "AOF active?"],
        ["AOF Current Size", "Timeseries per node", "redis.aof.size", "AOF file growth"],
        ["AOF Base Size", "Timeseries", "redis.aof.base_size", "Size after last rewrite"],
        ["AOF Buffer Length", "Timeseries", "redis.aof.buffer_length", "Pending AOF writes"],
        ["AOF Rewrite in Progress", "Status Widget", "aof_rewrite_in_progress (1/0)", "Currently rewriting?"],
        ["AOF Last Rewrite Duration", "Timeseries", "redis.aof.last_rewrite_time_sec", "Rewrite time trend"],
        ["AOF Last Write Status", "Status Widget", "aof_last_write_status (ok/err)", "Last AOF write success?"],
        ["AOF Fsync Latency", "Timeseries", "aof_delayed_fsync count", "Disk I/O contention"],
    ]
)

doc.add_heading("8.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["RDB Save Failed", "rdb_last_bgsave_status == err", "P1 - Critical", "Check disk space, permissions, fork failure"],
        ["No RDB Save > 1 Hour", "time since last save > 3600s", "P2 - Warning", "Verify BGSAVE schedule"],
        ["RDB Changes Piling Up", "changes_since_last_save > 100,000 for 10 min", "P2 - Warning", "BGSAVE may be failing or slow"],
        ["AOF Write Failed", "aof_last_write_status == err", "P1 - Critical", "Disk full or I/O error"],
        ["AOF File Growing Fast", "AOF size > 2x base_size for 30 min", "P2 - Warning", "Trigger BGREWRITEAOF"],
        ["AOF Fsync Delays", "aof_delayed_fsync > 0 in 5 min", "P2 - Warning", "Disk I/O contention, check IOPS"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  9. DASHBOARD #6 — CLIENT & CONNECTION MONITORING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Dashboard #6 — Client & Connection Monitoring", level=1)
doc.add_paragraph(
    "Monitors client connections, connection lifecycle, blocked clients, "
    "and detects connection leaks or storms."
)

doc.add_heading("9.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Connected Clients per Node", "Timeseries (3 lines)", "redis.net.clients by {redis_node}", "Client count per node"],
        ["Total Connected Clients", "Query Value", "sum(redis.net.clients{*})", "Cluster-wide client count"],
        ["Max Clients Configured", "Query Value", "maxclients config value", "Connection limit"],
        ["Client Usage %", "Gauge per node", "clients / maxclients * 100", "Connection saturation"],
        ["Blocked Clients", "Timeseries per node", "redis.clients.blocked", "Clients in BLPOP/BRPOP/WAIT"],
        ["Connected Clients Trend", "Timeseries (7d)", "redis.net.clients over 7 days", "Connection trend"],
        ["Client Connection Rate", "Timeseries", "rate(redis.stats.total_connections_received)", "New connections/sec"],
        ["Rejected Connections", "Timeseries", "rate(redis.stats.rejected_connections)", "Hitting maxclients limit"],
        ["Client Biggest Input Buffer", "Timeseries", "redis.clients.biggest_input_buf", "Large pending input"],
        ["Client Longest Output List", "Timeseries", "redis.clients.longest_output_list", "Large pending output"],
        ["Client List by Age", "Top List", "Custom: CLIENT LIST sorted by age", "Long-lived connections"],
        ["Client List by Idle", "Top List", "Custom: CLIENT LIST sorted by idle", "Idle connections (leak candidates)"],
    ]
)

doc.add_heading("9.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Clients > 80% of Max", "clients / maxclients > 0.80 for 5 min", "P2 - Warning", "Check for connection leaks"],
        ["Clients > 95% of Max", "clients / maxclients > 0.95 for 1 min", "P1 - Critical", "New connections will be rejected"],
        ["Rejected Connections > 0", "rejected_connections rate > 0 for 1 min", "P1 - Critical", "Maxclients reached, clients failing to connect"],
        ["Blocked Clients > 10", "blocked_clients > 10 for 5 min", "P2 - Warning", "Blocking command overuse"],
        ["Connection Spike", "Anomaly: connection rate > 2x baseline", "P2 - Warning", "Possible reconnect storm"],
        ["Client Output Buffer Large", "longest_output_list > 10000", "P2 - Warning", "Slow consumer or pub/sub backlog"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  10. DASHBOARD #7 — KEYSPACE & DATA ANALYTICS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Dashboard #7 — Keyspace & Data Analytics", level=1)
doc.add_paragraph(
    "Monitors the keyspace: total keys, key distribution by database, TTL coverage, "
    "and key growth trends."
)

doc.add_heading("10.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Total Keys (Master)", "Query Value", "redis.keys{redis_role:master}", "Total key count"],
        ["Keys by Database (db0-db15)", "Bar Chart", "redis.keys by {db}", "Database-level distribution"],
        ["Key Count Trend", "Timeseries", "redis.keys over 30 days", "Growth trend"],
        ["Keys Growth Rate", "Timeseries", "diff(redis.keys)", "Keys added/removed per interval"],
        ["Keys with TTL %", "Query Value", "redis.expires / redis.keys * 100", "TTL coverage (higher = better)"],
        ["Expired Keys/sec", "Timeseries", "rate(redis.stats.expired_keys)", "TTL expiration rate"],
        ["Evicted Keys/sec", "Timeseries", "rate(redis.stats.evicted_keys)", "Eviction rate (memory pressure)"],
        ["Key Size Distribution", "Histogram", "Custom: MEMORY USAGE sampling", "Key size patterns"],
        ["Largest Keys", "Top List", "Custom: redis-cli --bigkeys output", "Hot/large keys"],
        ["Key Namespace Distribution", "Pie Chart", "Custom: prefix analysis (app:*, session:*, cache:*)", "Namespace breakdown"],
        ["Scan Count", "Timeseries", "rate(redis.command.calls{command:scan})", "Full scan operations"],
    ]
)

doc.add_heading("10.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Key Count Anomaly", "Anomaly detection on redis.keys", "P2 - Warning", "Unexpected key growth/drop"],
        ["Low TTL Coverage", "expires / keys < 0.50 for 1 hour", "P3 - Info", "Many keys without TTL — memory risk"],
        ["Key Count Spike", "key growth > 100K/hour", "P2 - Warning", "Application generating excessive keys"],
        ["Large Key Detected", "Key > 10MB found", "P3 - Info", "Performance risk — consider restructuring"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  11. DASHBOARD #8 — LATENCY & SLOW QUERIES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Dashboard #8 — Latency & Slow Queries", level=1)
doc.add_paragraph(
    "Tracks Redis latency metrics, slow log entries, and identifies operations causing "
    "performance degradation."
)

doc.add_heading("11.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Avg Latency per Node", "Timeseries (3 lines)", "redis.info.latency_ms by {redis_node}", "Per-node response time"],
        ["P99 Latency", "Timeseries", "redis.info.latency_p99", "Tail latency"],
        ["Instantaneous Ops Latency", "Query Value", "1 / instantaneous_ops_per_sec * 1000", "Current avg op time"],
        ["Slow Log Count (last 10 min)", "Query Value", "redis.slowlog.count (custom)", "Slow commands detected"],
        ["Slow Log Entries", "Table", "Custom: SLOWLOG GET 20 — command, duration, timestamp", "Actual slow commands"],
        ["Slow Commands by Type", "Pie Chart", "Slow log grouped by command name", "Which commands are slow?"],
        ["Slow Command Duration Trend", "Timeseries", "Max slow log duration over time", "Getting worse or better?"],
        ["KEYS Command Usage", "Timeseries", "rate(redis.command.calls{command:keys})", "KEYS is O(n) — should be 0"],
        ["Blocking Command Duration", "Timeseries", "BLPOP, BRPOP, WAIT durations", "Blocking op latency"],
        ["Fork Latency", "Timeseries", "Custom: latest_fork_usec", "BGSAVE/BGREWRITEAOF fork time"],
        ["Event Loop Latency", "Timeseries", "redis.info.eventloop_latency", "Event loop health"],
    ]
)

doc.add_heading("11.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Latency > 5ms Avg", "avg latency > 5ms for 5 min", "P2 - Warning", "Check for blocking commands, big keys"],
        ["Latency > 20ms Avg", "avg latency > 20ms for 2 min", "P1 - Critical", "Severe degradation, immediate investigation"],
        ["Slow Log Spike", "slow log entries > 10 in 5 min", "P2 - Warning", "Review slow log, optimize commands"],
        ["KEYS Command Used", "command.calls{command:keys} > 0", "P2 - Warning", "Use SCAN instead — blocks Redis"],
        ["Fork Latency > 500ms", "latest_fork_usec > 500000", "P2 - Warning", "Large dataset fork delay"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  12. DASHBOARD #9 — NODE RESOURCE UTILIZATION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Dashboard #9 — Node Resource Utilization (CPU, Disk, Network)", level=1)
doc.add_paragraph(
    "Infrastructure-level monitoring for all 3 nodes. Redis is single-threaded for command processing, "
    "making CPU monitoring especially important."
)

doc.add_heading("12.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["CPU Usage per Node (Host)", "Timeseries (3 lines)", "system.cpu.user by {host}", "Host CPU"],
        ["Redis CPU (User)", "Timeseries per node", "redis.cpu.user", "Redis process CPU"],
        ["Redis CPU (System)", "Timeseries per node", "redis.cpu.sys", "Redis kernel CPU"],
        ["Redis CPU (Children)", "Timeseries per node", "redis.cpu.user_children + sys_children", "BGSAVE/BGREWRITEAOF child CPU"],
        ["Memory RSS per Node", "Timeseries (3 lines)", "redis.mem.rss by {redis_node}", "Actual memory from OS"],
        ["Host Memory Used %", "Gauge per node", "system.mem.used / total * 100", "Host memory"],
        ["Disk Usage %", "Gauge per node", "system.disk.in_use by {host}", "Disk capacity"],
        ["Disk Read/Write IOPS", "Timeseries", "system.io.r_s + w_s by {host}", "Disk I/O (important for AOF/RDB)"],
        ["Disk Write Latency", "Timeseries", "system.io.await by {host}", "Disk latency"],
        ["Network Bytes In/Out", "Timeseries per node", "redis.net.input_bytes + output_bytes", "Redis-level network"],
        ["Host Network In/Out", "Timeseries", "system.net.bytes_rcvd + sent by {host}", "Host network (includes replication)"],
        ["Swap Usage", "Timeseries per node", "system.swap.used by {host}", "CRITICAL: Redis + swap = very bad"],
        ["OS File Descriptors Used", "Gauge per node", "system.fs.file_handles.used", "FD pressure"],
    ]
)

doc.add_heading("12.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Redis CPU > 70%", "redis.cpu.user > 70% for 10 min (single core)", "P2 - Warning", "Near single-thread saturation"],
        ["Redis CPU > 90%", "redis.cpu.user > 90% for 5 min", "P1 - Critical", "Command processing saturated"],
        ["Swap Usage > 0", "system.swap.used > 0 for any Redis host", "P1 - Critical", "Redis swapping — severe latency, disable swap or add memory"],
        ["Host Memory > 90%", "system.mem.used > 90% for 5 min", "P1 - Critical", "OOM kill risk for Redis"],
        ["Disk Usage > 85%", "system.disk.in_use > 85%", "P2 - Warning", "RDB/AOF writes may fail"],
        ["Disk Latency > 10ms", "system.io.await > 10ms for 10 min", "P2 - Warning", "Slow persistence operations"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  13. DASHBOARD #10 — REPLICATION DEEP-DIVE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Dashboard #10 — Replication Deep-Dive", level=1)
doc.add_paragraph(
    "Detailed replication monitoring between master and 2 replicas. Tracks replication lag, "
    "buffer status, partial/full resync events, and replication backlog."
)

doc.add_heading("13.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Master Replication Offset", "Timeseries", "redis.replication.master_repl_offset", "Master write position"],
        ["Replica Offset (per replica)", "Timeseries (2 lines)", "slave0_offset, slave1_offset", "Replica read position"],
        ["Replication Lag Bytes", "Timeseries (2 lines)", "master_offset - slave_offset per replica", "Byte-level lag"],
        ["Replication Lag Seconds", "Timeseries (2 lines)", "redis.replication.slave_repl_offset_lag_seconds", "Time-based lag estimate"],
        ["Master Link Status", "Status Widget per replica", "master_link_status (up/down)", "Connection to master alive?"],
        ["Master Link Down Duration", "Query Value per replica", "master_link_down_since_seconds", "How long disconnected"],
        ["Replication Backlog Size", "Query Value", "repl_backlog_size (configured)", "Partial resync window"],
        ["Replication Backlog Active", "Query Value", "repl_backlog_active (1/0)", "Backlog in use?"],
        ["Partial Resyncs (OK)", "Timeseries", "rate(redis.stats.sync_partial_ok)", "Efficient incremental syncs"],
        ["Partial Resyncs (Err)", "Timeseries", "rate(redis.stats.sync_partial_err)", "Failed partial syncs -> full resync"],
        ["Full Resyncs", "Timeseries", "rate(redis.stats.sync_full)", "EXPENSIVE — full data transfer"],
        ["Replication Buffer Usage", "Timeseries", "Client output buffer for replica", "Buffer pressure"],
        ["Second Replication Offset", "Query Value", "second_repl_offset", "PSYNC2 chain replication"],
    ]
)

doc.add_heading("13.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Replication Lag > 10MB", "offset_diff > 10MB for 2 min", "P2 - Warning", "Check replica I/O and network"],
        ["Replication Lag > 100MB", "offset_diff > 100MB for 1 min", "P1 - Critical", "Replica severely behind, stale reads"],
        ["Master Link Down", "master_link_status == down for 30s", "P1 - Critical", "Replica disconnected from master"],
        ["Full Resync Occurred", "sync_full incremented", "P2 - Warning", "Expensive operation — check why partial failed"],
        ["All Replicas Disconnected", "connected_slaves == 0 for 1 min", "P1 - Critical", "No HA — single point of failure"],
        ["Replication Backlog Too Small", "partial_err increasing", "P2 - Warning", "Increase repl-backlog-size"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  14. DASHBOARD #11 — SECURITY & ACL AUDIT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Dashboard #11 — Security & ACL Audit", level=1)
doc.add_paragraph(
    "Monitors Redis ACL users, authentication events, dangerous command usage, "
    "and security posture compliance."
)

doc.add_heading("14.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Total ACL Users", "Query Value", "Custom: ACL LIST user count", "User count"],
        ["Users by Permission Level", "Pie Chart", "Custom: allcommands, readonly, custom", "Permission distribution"],
        ["Default User Status", "Check Status", "default user disabled?", "Security compliance"],
        ["AUTH Failures", "Timeseries", "Log: 'AUTH failed' OR 'WRONGPASS'", "Failed auth attempts"],
        ["AUTH Failures by Source IP", "Top List", "Auth failures grouped by client IP", "Brute force detection"],
        ["ACL Violations (NOPERM)", "Timeseries", "Log: 'NOPERM' OR 'no permissions'", "Permission violations"],
        ["Dangerous Commands Blocked", "Timeseries", "FLUSHALL, FLUSHDB, DEBUG, KEYS blocked by ACL", "Security enforcement"],
        ["CONFIG SET Usage", "Timeseries", "rate(redis.command.calls{command:config|set})", "Runtime config changes"],
        ["SHUTDOWN Command Attempts", "Event Stream", "Log: SHUTDOWN command detected", "Administrative actions"],
        ["ACL Changes", "Event Stream", "Log: 'ACL SETUSER' OR 'ACL DELUSER'", "User management audit"],
        ["Protected Mode Status", "Check Status per node", "protected-mode yes/no", "Network security"],
        ["Rename Command Coverage", "Table", "Custom: which dangerous commands are renamed", "Command protection"],
    ]
)

doc.add_heading("14.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Auth Failures > 10/min", "AUTH failed count > 10 in 1 min", "P1 - Critical", "Possible brute force, investigate source IPs"],
        ["ACL Violation Detected", "NOPERM events > 0", "P2 - Warning", "Client misconfiguration or unauthorized attempt"],
        ["Default User Active", "default user enabled", "P2 - Warning", "Security risk — should be disabled"],
        ["FLUSHALL/FLUSHDB Executed", "Command detected", "P1 - Critical", "DATA LOSS — verify if intentional"],
        ["ACL User Modified", "ACL SETUSER event", "P2 - Warning", "Audit: who changed what"],
        ["CONFIG SET Executed", "config|set command rate > 0", "P3 - Info", "Runtime config change audit"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  15. DASHBOARD #12 — RESTART & AVAILABILITY TRACKING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Dashboard #12 — Restart & Availability Tracking", level=1)
doc.add_paragraph(
    "Tracks Redis and Sentinel restarts, uptime, availability SLA, and correlates "
    "restarts with failover events."
)

doc.add_heading("15.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["Redis Uptime per Node", "Query Value (3x)", "redis.info.uptime_in_seconds by {node}", "Current uptime in days/hours"],
        ["Sentinel Uptime per Node", "Query Value (3x)", "sentinel uptime_in_seconds by {node}", "Sentinel uptime"],
        ["Uptime Timeline (Redis)", "Timeseries", "redis.info.uptime_in_seconds (drops = restart)", "Visual restart detection"],
        ["Uptime Timeline (Sentinel)", "Timeseries", "sentinel uptime_in_seconds", "Sentinel restart detection"],
        ["Restart Events", "Event Timeline", "Log: 'Redis is starting' OR 'Server started'", "Restart timestamps"],
        ["Restart Count (30d) per Node", "Bar Chart", "uptime resets per node", "Restart frequency"],
        ["Restart vs Failover Correlation", "Event Overlay", "restart events + switch-master events", "Did restart trigger failover?"],
        ["MTBR per Node", "Query Value", "Mean time between restarts", "Stability metric"],
        ["Availability % (30d)", "Query Value", "(total - downtime) / total * 100", "SLA metric (target: 99.99%)"],
        ["Cluster Availability", "Query Value", "Time with functioning master", "Service-level availability"],
        ["Restart Reason", "Log Stream", "Log: 'signal' OR 'shutdown' OR 'systemd' OR 'OOM'", "Root cause classification"],
        ["Process Check (redis-server)", "Check Status per node", "process.up{process:redis-server}", "OS process alive?"],
        ["Process Check (redis-sentinel)", "Check Status per node", "process.up{process:redis-sentinel}", "Sentinel process alive?"],
    ]
)

doc.add_heading("15.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Redis Restarted", "Uptime drop detected", "P1 - Critical", "Verify master status, check failover"],
        ["Sentinel Restarted", "Sentinel uptime drop detected", "P1 - Critical", "Verify quorum maintained"],
        ["Frequent Restarts", "> 2 restarts per node per week", "P1 - Critical", "Investigate root cause"],
        ["Availability < 99.9% (30d)", "Rolling availability < 99.9%", "P2 - Warning", "SLA breach risk"],
        ["Process Down", "process.up == 0 for redis-server or redis-sentinel", "P1 - Critical", "Restart service immediately"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  16. DASHBOARD #13 — TLS / CERTIFICATE MONITORING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Dashboard #13 — TLS / Certificate Monitoring", level=1)
doc.add_paragraph(
    "Monitors TLS certificate health, cipher usage, and TLS configuration compliance "
    "for Redis server and Sentinel connections."
)

doc.add_heading("16.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Datadog Metric / Query", "Purpose"],
    [
        ["TLS Certificate Expiry (Days)", "Query Value per node", "Custom: cert expiry check", "Warn < 30d, Critical < 7d"],
        ["TLS Enabled (Redis)", "Check Status", "tls-port configured vs port 0", "Redis TLS active?"],
        ["TLS Enabled (Sentinel)", "Check Status", "sentinel tls-port configured", "Sentinel TLS active?"],
        ["TLS vs Non-TLS Clients", "Pie Chart", "Connections by TLS status", "Security compliance"],
        ["TLS Protocol Version", "Pie Chart", "TLS 1.2 vs 1.3 connections", "Protocol migration"],
        ["TLS Handshake Failures", "Timeseries", "Log: 'SSL' AND 'error'", "Cert/config issues"],
        ["Replication TLS Status", "Check Status", "tls-replication yes/no", "Master-replica encryption"],
        ["Certificate Chain Valid", "Check Status per node", "Custom: cert chain verification", "CA trust chain intact?"],
    ]
)

doc.add_heading("16.2 Alerts", level=2)
add_table(
    ["Alert Name", "Condition", "Severity", "Action"],
    [
        ["Cert Expiry < 30 Days", "cert_expiry_days < 30", "P2 - Warning", "Schedule cert renewal"],
        ["Cert Expiry < 7 Days", "cert_expiry_days < 7", "P1 - Critical", "Urgent cert renewal"],
        ["Non-TLS Connections", "non_tls_connections > 0", "P2 - Warning", "Security compliance violation"],
        ["TLS Handshake Failures", "handshake failures > 5/min", "P3 - Info", "Client cert issues"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  17. DASHBOARD #14 — EXECUTIVE / SLA SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("17. Dashboard #14 — Executive / SLA Summary", level=1)
doc.add_paragraph(
    "High-level dashboard for management and stakeholders. Shows SLAs, key performance "
    "indicators, and overall system health without technical details."
)

doc.add_heading("17.1 Widgets & Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Cluster Health", "Traffic Light", "Composite score", "At-a-glance status"],
        ["Availability SLA % (30d)", "Query Value (large)", "Uptime % (target: 99.99%)", "SLA compliance"],
        ["Commands Processed Today", "Query Value", "Total commands today", "Daily throughput"],
        ["Cache Hit Rate", "Query Value (large)", "hits / (hits+misses) * 100", "Cache effectiveness KPI"],
        ["Avg Latency", "Query Value", "Average command latency", "Performance KPI"],
        ["P99 Latency", "Query Value", "99th percentile latency", "Tail latency KPI"],
        ["Evicted Keys Today", "Query Value", "Keys evicted today", "Data loss indicator"],
        ["Failovers This Month", "Query Value", "Sentinel switch-master count", "Stability KPI"],
        ["Active Incidents", "Query Value", "P1/P2 alerts firing", "Operational status"],
        ["Memory Utilization", "Gauge", "Used / Max across cluster", "Capacity KPI"],
        ["Key Count Trend (30d)", "Timeseries", "redis.keys over 30 days", "Growth trend"],
        ["Capacity Forecast", "Timeseries + Forecast", "Memory + keys with forecast", "30-day projection"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  18. ALERTING STRATEGY & ESCALATION MATRIX
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("18. Alerting Strategy & Escalation Matrix", level=1)

doc.add_heading("18.1 Severity Levels", level=2)
add_table(
    ["Severity", "Response Time", "Notification", "Escalation", "Examples"],
    [
        ["P1 - Critical", "< 5 min", "PagerDuty + Slack #redis-critical", "SRE Lead in 15 min", "Node down, swap active, master lost, auth attack"],
        ["P2 - Warning", "< 30 min", "Slack #redis-alerts + Email", "Next standup", "Memory high, repl lag, evictions, slow queries"],
        ["P3 - Info", "< 4 hours", "Slack #redis-info", "Weekly review", "Anomalies, TTL coverage, idle connections"],
    ]
)

doc.add_heading("18.2 Alert Routing", level=2)
add_table(
    ["Category", "Primary Team", "Secondary", "Handle"],
    [
        ["Cluster / Nodes / Failover", "Platform SRE", "DevOps", "@pagerduty-redis-cluster"],
        ["Sentinel / Failover", "Platform SRE", "DevOps", "@pagerduty-redis-cluster"],
        ["Memory / Eviction", "Application Team", "Platform SRE", "@slack-redis-alerts"],
        ["Persistence (RDB/AOF)", "Platform SRE", "Storage Team", "@slack-redis-alerts"],
        ["Performance / Latency", "Application Team", "Platform SRE", "@slack-redis-alerts"],
        ["Security / ACL", "Security Team", "Platform SRE", "@pagerduty-security"],
        ["Infrastructure (CPU/Disk/Net)", "Infra Team", "Platform SRE", "@pagerduty-infra"],
    ]
)

doc.add_heading("18.3 Alert Best Practices", level=2)
for practice in [
    "Use composite monitors (e.g., alert only if memory high AND evictions active)",
    "Differentiate thresholds for master vs replica (master is more critical)",
    "Use anomaly detection for throughput and connection count",
    "Include runbook links in every alert notification body",
    "Use Datadog Downtime during planned maintenance/failover tests",
    "Tag all monitors with redis_role (master/replica) for filtered views",
    "Auto-resolve alerts that recover within evaluation window",
    "Review false-positive rate monthly and tune thresholds",
]:
    doc.add_paragraph(practice, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  19. DATADOG MONITOR DEFINITIONS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("19. Datadog Monitor Definitions — Complete List", level=1)

monitors = [
    ["M-001", "Redis Node Unreachable", "redis.can_connect == 0", "P1", "30s"],
    ["M-002", "Sentinel Node Down", "sentinel can_connect == 0", "P1", "30s"],
    ["M-003", "No Master Detected", "No node with role == master", "P1", "1 min"],
    ["M-004", "Replica Disconnected", "connected_slaves < 2", "P1", "1 min"],
    ["M-005", "Master Link Down (replica)", "master_link_status == down", "P1", "30s"],
    ["M-006", "Sentinel Quorum Lost", "Active sentinels < 2", "P1", "30s"],
    ["M-007", "Failover Occurred", "Log: +switch-master", "P1", "0s"],
    ["M-008", "Frequent Failovers", "> 2 failovers in 1 hour", "P1", "0s"],
    ["M-009", "Memory > 95% of Max", "used/maxmemory > 0.95", "P1", "2 min"],
    ["M-010", "Swap Usage > 0", "system.swap.used > 0 on Redis host", "P1", "0s"],
    ["M-011", "Redis CPU > 90%", "redis.cpu.user > 90%", "P1", "5 min"],
    ["M-012", "Rejected Connections", "rejected_connections > 0", "P1", "1 min"],
    ["M-013", "RDB Save Failed", "rdb_last_bgsave_status == err", "P1", "0s"],
    ["M-014", "AOF Write Failed", "aof_last_write_status == err", "P1", "0s"],
    ["M-015", "Auth Failures > 10/min", "AUTH failed > 10 per min", "P1", "1 min"],
    ["M-016", "FLUSHALL/FLUSHDB Executed", "Command detected in logs", "P1", "0s"],
    ["M-017", "Redis Restarted", "Uptime drop detected", "P1", "0s"],
    ["M-018", "All Replicas Disconnected", "connected_slaves == 0", "P1", "1 min"],
    ["M-019", "Replication Lag > 100MB", "offset_diff > 100MB", "P1", "1 min"],
    ["M-020", "Latency > 20ms", "avg latency > 20ms", "P1", "2 min"],
    ["M-021", "Zero Throughput", "commands/sec == 0", "P1", "5 min"],
    ["M-022", "Host Memory > 90%", "system.mem.used > 90%", "P1", "5 min"],
    ["M-023", "TLS Cert Expiry < 7 Days", "cert_expiry_days < 7", "P1", "0s"],
    ["M-024", "Hit Rate < 80%", "keyspace_hit_rate < 80%", "P1", "5 min"],
    ["M-025", "Memory > 85% of Max", "used/maxmemory > 0.85", "P2", "5 min"],
    ["M-026", "Evictions Active", "evicted_keys rate > 0", "P2", "5 min"],
    ["M-027", "Fragmentation > 1.5", "fragmentation_ratio > 1.5", "P2", "30 min"],
    ["M-028", "Fragmentation < 1.0 (Swap)", "fragmentation_ratio < 1.0", "P1", "10 min"],
    ["M-029", "Replication Lag > 10MB", "offset_diff > 10MB", "P2", "2 min"],
    ["M-030", "Full Resync Occurred", "sync_full incremented", "P2", "0s"],
    ["M-031", "Hit Rate < 90%", "keyspace_hit_rate < 90%", "P2", "10 min"],
    ["M-032", "Throughput Drop > 50%", "Anomaly on commands/sec", "P2", "5 min"],
    ["M-033", "Latency > 5ms", "avg latency > 5ms", "P2", "5 min"],
    ["M-034", "Slow Log Spike", "slow log entries > 10 in 5 min", "P2", "5 min"],
    ["M-035", "Clients > 80% of Max", "clients/maxclients > 0.80", "P2", "5 min"],
    ["M-036", "Blocked Clients > 10", "blocked_clients > 10", "P2", "5 min"],
    ["M-037", "No RDB Save > 1 Hour", "Time since save > 3600s", "P2", "0s"],
    ["M-038", "AOF File Growing", "aof_size > 2x base_size", "P2", "30 min"],
    ["M-039", "Disk Usage > 85%", "system.disk.in_use > 85%", "P2", "5 min"],
    ["M-040", "Connection Spike", "Anomaly on connection rate", "P2", "5 min"],
    ["M-041", "Sentinel Config Mismatch", "config_epoch differs", "P2", "5 min"],
    ["M-042", "ACL Violation", "NOPERM events > 0", "P2", "0s"],
    ["M-043", "Default User Active", "default user enabled", "P2", "0s"],
    ["M-044", "TLS Cert Expiry < 30 Days", "cert_expiry_days < 30", "P2", "0s"],
    ["M-045", "Non-TLS Connections", "non_tls_connections > 0", "P2", "0s"],
]

add_table(
    ["Monitor ID", "Name", "Condition", "Severity", "Eval Window"],
    monitors
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  20. RUNBOOK REFERENCES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("20. Runbook References", level=1)

add_table(
    ["Runbook", "Covers Monitors", "Description"],
    [
        ["RB-001: Node Down Recovery", "M-001, M-017", "Check host > Check Redis process > Restart > Verify replication > Confirm Sentinel detects"],
        ["RB-002: Sentinel Failover Investigation", "M-003, M-007, M-008", "Check Sentinel logs > Identify trigger > Verify new master > Check replica reconnection"],
        ["RB-003: Sentinel Quorum Restore", "M-002, M-006", "Identify failed sentinel(s) > Restart > Verify quorum > Check master monitoring"],
        ["RB-004: Memory Pressure Response", "M-009, M-025, M-026", "Check memory breakdown > Identify large keys > Evict/expire > Increase maxmemory"],
        ["RB-005: Swap Elimination", "M-010, M-028", "Identify swap cause > Add memory > Reduce maxmemory > Disable swap on Redis hosts"],
        ["RB-006: Replication Recovery", "M-004, M-005, M-018, M-029, M-030", "Check master health > Restart replica > SLAVEOF new master > Monitor sync progress"],
        ["RB-007: Persistence Failure", "M-013, M-014, M-037, M-038", "Check disk space > Check permissions > Verify fork() > Manual BGSAVE > Check AOF health"],
        ["RB-008: Connection Management", "M-012, M-035, M-036, M-040", "Identify heavy clients > CLIENT KILL idle > Increase maxclients > Fix connection leaks"],
        ["RB-009: Latency Investigation", "M-020, M-033, M-034", "Check SLOWLOG > Identify blocking commands > Check fork latency > Optimize large keys"],
        ["RB-010: Security Incident", "M-015, M-016, M-042, M-043", "Block source IPs > Rotate credentials > Audit ACL changes > Check for data loss"],
        ["RB-011: TLS Certificate Renewal", "M-023, M-044, M-045", "Renew certs > Deploy to all nodes > Rolling restart > Verify handshake"],
        ["RB-012: Full Cluster Recovery", "Multiple", "Start master first > Verify data > Start replicas > SLAVEOF master > Start sentinels > Verify quorum"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  21. APPENDIX — METRIC REFERENCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("21. Appendix — Full Metric Reference", level=1)

doc.add_heading("21.1 Datadog Redis Integration Metrics", level=2)
add_table(
    ["Metric Name", "Type", "Description"],
    [
        ["redis.can_connect", "Service Check", "Can agent connect to Redis?"],
        ["redis.info.uptime_in_seconds", "Gauge", "Seconds since Redis start"],
        ["redis.net.clients", "Gauge", "Connected clients"],
        ["redis.net.commands", "Counter", "Total commands processed"],
        ["redis.net.input_bytes", "Counter", "Input bytes received"],
        ["redis.net.output_bytes", "Counter", "Output bytes sent"],
        ["redis.clients.blocked", "Gauge", "Clients in blocking command"],
        ["redis.stats.keyspace_hits", "Counter", "Successful key lookups"],
        ["redis.stats.keyspace_misses", "Counter", "Failed key lookups"],
        ["redis.stats.evicted_keys", "Counter", "Keys evicted due to maxmemory"],
        ["redis.stats.expired_keys", "Counter", "Keys expired by TTL"],
        ["redis.stats.total_connections_received", "Counter", "Total connections received"],
        ["redis.stats.rejected_connections", "Counter", "Connections rejected (maxclients)"],
        ["redis.stats.sync_full", "Counter", "Full resync count"],
        ["redis.stats.sync_partial_ok", "Counter", "Successful partial resyncs"],
        ["redis.stats.sync_partial_err", "Counter", "Failed partial resyncs"],
        ["redis.mem.used", "Gauge", "Memory used by Redis (bytes)"],
        ["redis.mem.rss", "Gauge", "OS-level memory (bytes)"],
        ["redis.mem.peak", "Gauge", "Peak memory usage (bytes)"],
        ["redis.mem.maxmemory", "Gauge", "Configured max memory"],
        ["redis.mem.fragmentation_ratio", "Gauge", "RSS / used (target: 1.0-1.5)"],
        ["redis.mem.lua", "Gauge", "Lua engine memory"],
        ["redis.keys", "Gauge", "Number of keys per db"],
        ["redis.expires", "Gauge", "Keys with TTL per db"],
        ["redis.replication.connected_slaves", "Gauge", "Number of connected replicas"],
        ["redis.replication.master_repl_offset", "Gauge", "Master replication offset"],
        ["redis.cpu.user", "Counter", "User CPU seconds"],
        ["redis.cpu.sys", "Counter", "System CPU seconds"],
        ["redis.cpu.user_children", "Counter", "Children user CPU (fork)"],
        ["redis.rdb.last_save_time", "Gauge", "Unix time of last RDB save"],
        ["redis.rdb.changes_since_last_save", "Gauge", "Changes since last snapshot"],
        ["redis.rdb.last_bgsave_time_sec", "Gauge", "Duration of last BGSAVE"],
        ["redis.aof.size", "Gauge", "Current AOF file size"],
        ["redis.aof.buffer_length", "Gauge", "AOF buffer length"],
        ["redis.command.calls", "Counter", "Calls per command type"],
        ["redis.command.usec_per_call", "Gauge", "Avg microseconds per call"],
    ]
)

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Redis HA Cluster Datadog Monitoring Strategy - Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output = "/Users/tejasodanapalli/rabbitmq/RabbitMQ/redis-cluster-deployment/docs/Redis_Datadog_Monitoring_Dashboard_Strategy.docx"
doc.save(output)
print(f"Document saved: {output}")
print(f"Pages (estimated): ~45-50 pages")
print(f"Dashboards: 14")
print(f"Monitors: 45")
print(f"Runbooks: 12")
