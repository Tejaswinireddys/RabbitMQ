#!/usr/bin/env python3
"""
Generate Redis 3-Node HA Cluster (with Sentinel) Operations Runbook
Mapped to Datadog Dashboards — For L1/L2 Engineers
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0xA4, 0x1E, 0x22)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 2'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

def add_bold_para(bold_text, normal_text=""):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)

def add_issue_block(issue_id, title, dashboard, severity, symptoms, manual_steps, script, escalation):
    doc.add_heading(f"{issue_id}: {title}", level=2)
    add_table(["Field", "Value"], [
        ["Dashboard", dashboard],
        ["Severity", severity],
        ["L1/L2 Actionable", "Yes"],
        ["Est. Resolution", "5-30 minutes"],
    ])
    doc.add_paragraph("")
    add_bold_para("Symptoms / What You See on Dashboard:")
    for s in symptoms:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph("")
    add_bold_para("Manual Steps (L1/L2):")
    for i, step in enumerate(manual_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Bullet')
    doc.add_paragraph("")
    add_bold_para("Automation Script:")
    add_code(script)
    doc.add_paragraph("")
    add_bold_para("Escalation: ", escalation)
    doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph("")
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Redis 3-Node HA Cluster (with Sentinel)\nOperations Runbook")
run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor(0xA4, 0x1E, 0x22)
doc.add_paragraph("")
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = s.add_run("Dashboard-Mapped Issue Resolution Guide\nFor L1/L2 Engineers")
run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph("")
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["Version: 1.0", f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
             "Classification: Internal / Operations", "Author: Principal Engineering / SRE Team"]:
    m.add_run(line + "\n").font.size = Pt(12)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Environment Variables Reference", level=1)
doc.add_paragraph(
    "All scripts use parameterized environment variables. Source the appropriate "
    "environment file before executing any script."
)
add_code("""# ============================================================
# Redis Environment Configuration
# File: /opt/cached/scripts/env/<environment>.env
# ============================================================

# --- Cluster Nodes ---
export REDIS_NODE1="redis-node-1.example.com"
export REDIS_NODE2="redis-node-2.example.com"
export REDIS_NODE3="redis-node-3.example.com"
export REDIS_NODES="${REDIS_NODE1} ${REDIS_NODE2} ${REDIS_NODE3}"

# --- Ports ---
export REDIS_PORT="6379"
export SENTINEL_PORT="26379"

# --- Authentication ---
export REDIS_AUTH_PASS="<from-vault>"         # NEVER hardcode
export SENTINEL_AUTH_PASS="<from-vault>"      # If sentinel auth enabled

# --- Paths ---
export REDIS_HOME="/opt/cached/current"
export REDIS_CONF_DIR="${REDIS_HOME}/conf"
export REDIS_DATA_DIR="${REDIS_HOME}/data"
export REDIS_LOG_DIR="${REDIS_HOME}/logs"
export REDIS_BIN="${REDIS_HOME}/bin"
export SCRIPTS_DIR="/opt/cached/scripts"

# --- Thresholds ---
export REDIS_MEM_WARN_PCT=80
export REDIS_MEM_CRIT_PCT=90
export REDIS_CONN_WARN_PCT=80
export REDIS_REPL_LAG_WARN_BYTES=10485760    # 10MB
export REDIS_REPL_LAG_CRIT_BYTES=104857600   # 100MB
export REDIS_SLOWLOG_THRESHOLD=10000          # microseconds (10ms)

# --- SSH ---
export SSH_USER="sre-user"
export SSH_KEY="/home/${SSH_USER}/.ssh/id_rsa"
export SSH_OPTS="-o StrictHostKeyChecking=no -o ConnectTimeout=5"

# --- Notification ---
export SLACK_WEBHOOK_URL="https://hooks.slack.com/services/XXX/YYY/ZZZ"
export PAGERDUTY_SERVICE_KEY="<from-vault>"
export OPS_EMAIL="sre-team@company.com"

# --- Environment ---
export ENVIRONMENT="production"
export CLUSTER_NAME="redis-ha-cluster-01"
export SENTINEL_MASTER_NAME="mymaster"

# Usage: source /opt/cached/scripts/env/production.env""")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  COMMON FUNCTIONS LIBRARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Common Functions Library", level=1)
doc.add_paragraph("File: /opt/cached/scripts/lib/redis_common.sh")
add_code("""#!/bin/bash
# ============================================================
# Redis Runbook - Common Functions Library
# File: /opt/cached/scripts/lib/redis_common.sh
# ============================================================
set -euo pipefail

RED='\\033[0;31m'; GREEN='\\033[0;32m'; YELLOW='\\033[1;33m'; NC='\\033[0m'

log_info()  { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] [INFO]  $*"; }
log_warn()  { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${YELLOW}[WARN]${NC}  $*"; }
log_error() { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${RED}[ERROR]${NC} $*"; }
log_ok()    { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${GREEN}[OK]${NC}    $*"; }

validate_env() {
    local required="REDIS_NODE1 REDIS_NODE2 REDIS_NODE3 REDIS_PORT SENTINEL_PORT"
    for var in ${required}; do
        if [[ -z "${!var:-}" ]]; then
            log_error "Variable ${var} not set. Source env file first."
            exit 1
        fi
    done
    log_info "Environment: ${ENVIRONMENT:-unknown} / ${CLUSTER_NAME:-unknown}"
}

# Redis CLI wrapper
redis_cmd() {
    local host="${1}"; local port="${2:-${REDIS_PORT}}"; shift 2
    ${REDIS_BIN}/redis-cli -h "${host}" -p "${port}" -a "${REDIS_AUTH_PASS}" --no-auth-warning "$@" 2>/dev/null
}

# Sentinel CLI wrapper
sentinel_cmd() {
    local host="${1}"; shift
    ${REDIS_BIN}/redis-cli -h "${host}" -p "${SENTINEL_PORT}" ${SENTINEL_AUTH_PASS:+-a "${SENTINEL_AUTH_PASS}"} --no-auth-warning "$@" 2>/dev/null
}

# SSH helper
remote_exec() {
    local node="${1}"; shift
    ssh ${SSH_OPTS} -i "${SSH_KEY}" "${SSH_USER}@${node}" "$@"
}

# Get specific INFO field
redis_info_field() {
    local host="${1}"; local field="${2}"
    redis_cmd "${host}" "${REDIS_PORT}" INFO ALL | grep "^${field}:" | cut -d: -f2 | tr -d '\\r'
}

# Notification
notify_slack() {
    local message="${1}"; local severity="${2:-info}"
    local color="good"
    [[ "${severity}" == "warning" ]] && color="warning"
    [[ "${severity}" == "critical" ]] && color="danger"
    curl -s -X POST "${SLACK_WEBHOOK_URL}" \\
        -H 'Content-type: application/json' \\
        -d "{
            \\"attachments\\": [{
                \\"color\\": \\"${color}\\",
                \\"title\\": \\"Redis Runbook - ${ENVIRONMENT}\\",
                \\"text\\": \\"${message}\\",
                \\"footer\\": \\"Cluster: ${CLUSTER_NAME} | $(date '+%Y-%m-%d %H:%M:%S')\\"
            }]
        }" > /dev/null 2>&1 || true
}

check_node_reachable() {
    local node="${1}"
    ssh ${SSH_OPTS} -i "${SSH_KEY}" "${SSH_USER}@${node}" "echo ok" &>/dev/null
}

report_header() {
    echo "============================================================"
    echo "  Redis Runbook Execution Report"
    echo "  Environment: ${ENVIRONMENT} | Cluster: ${CLUSTER_NAME}"
    echo "  Executed by: $(whoami) at $(date '+%Y-%m-%d %H:%M:%S')"
    echo "============================================================"
}""")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-001: Node Down
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 1: Cluster Health & Node Issues", level=1)
doc.add_paragraph("Dashboard: Cluster Health & Replication Overview")

add_issue_block(
    "RB-REDIS-001", "Redis Node Down / Unreachable",
    "Dashboard #1 — Cluster Health & Replication Overview",
    "P1 - Critical",
    [
        "Node status shows cannot connect for one or more nodes",
        "redis.can_connect check fails",
        "Metrics stop appearing for the affected node",
        "If master is down: Sentinel should auto-failover within 30s",
        "If replica is down: Replication becomes degraded (only 1 replica)",
    ],
    [
        "Verify the alert — identify which node is down from Datadog",
        "SSH to affected node: ssh <user>@<node>",
        "If SSH fails: Host is down — escalate to Infrastructure team",
        "Check if redis-server process is running: ps aux | grep redis-server",
        "Check systemd status: sudo systemctl status redis",
        "Check Redis logs: tail -100 /opt/cached/current/logs/redis.log",
        "If process is down, check for OOM kill: dmesg | grep -i oom | tail -5",
        "Check disk space: df -h /opt/cached/current/data",
        "Attempt restart: sudo systemctl start redis",
        "Verify: redis-cli -h localhost -p 6379 -a <pass> PING (expect PONG)",
        "Check role: redis-cli INFO replication | grep role (should be 'slave' if was replica)",
        "If this was master: Check if Sentinel performed failover — new master elected?",
        "Verify cluster: Check all 3 nodes' replication status",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-001: Redis Node Down - Diagnosis & Recovery
# Usage: ./rb-redis-001-node-down.sh <node_hostname>
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

AFFECTED_NODE="${1:?Usage: $0 <node_hostname>}"
REPORT_FILE="/tmp/redis-node-down-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"
log_info "Investigating node: ${AFFECTED_NODE}" | tee -a "${REPORT_FILE}"

# Step 1: Ping
log_info "Step 1: Host reachability..." | tee -a "${REPORT_FILE}"
if ! ping -c 3 -W 2 "${AFFECTED_NODE}" &>/dev/null; then
    log_error "Host ${AFFECTED_NODE} UNREACHABLE" | tee -a "${REPORT_FILE}"
    notify_slack "REDIS NODE DOWN: ${AFFECTED_NODE} unreachable. Escalate to Infra." "critical"
    exit 1
fi
log_ok "Host reachable" | tee -a "${REPORT_FILE}"

# Step 2: SSH
log_info "Step 2: SSH access..." | tee -a "${REPORT_FILE}"
if ! check_node_reachable "${AFFECTED_NODE}"; then
    log_error "Cannot SSH to ${AFFECTED_NODE}" | tee -a "${REPORT_FILE}"
    exit 1
fi
log_ok "SSH OK" | tee -a "${REPORT_FILE}"

# Step 3: Redis process
log_info "Step 3: Redis process check..." | tee -a "${REPORT_FILE}"
REDIS_PID=$(remote_exec "${AFFECTED_NODE}" "pgrep -f 'redis-server'" 2>/dev/null || echo "")
if [[ -z "${REDIS_PID}" ]]; then
    log_error "redis-server process NOT running" | tee -a "${REPORT_FILE}"

    # Check for OOM
    log_info "Checking for OOM kill..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "sudo dmesg | grep -i 'oom.*redis' | tail -5" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Check disk space
    log_info "Disk space..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "df -h ${REDIS_DATA_DIR}" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Check last log entries
    log_info "Last log entries..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "tail -30 ${REDIS_LOG_DIR}/redis.log" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Attempt restart
    log_info "Attempting restart..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "sudo systemctl start redis" 2>&1 | tee -a "${REPORT_FILE}"
    sleep 10

    # Verify
    PONG=$(redis_cmd "${AFFECTED_NODE}" "${REDIS_PORT}" PING 2>/dev/null || echo "FAIL")
    if [[ "${PONG}" == "PONG" ]]; then
        log_ok "Redis restarted successfully" | tee -a "${REPORT_FILE}"
        ROLE=$(redis_info_field "${AFFECTED_NODE}" "role")
        log_info "Node role: ${ROLE}" | tee -a "${REPORT_FILE}"
        notify_slack "REDIS RECOVERED: ${AFFECTED_NODE} restarted (role: ${ROLE})" "good"
    else
        log_error "Redis FAILED to start" | tee -a "${REPORT_FILE}"
        notify_slack "REDIS CRITICAL: ${AFFECTED_NODE} cannot start. Manual intervention." "critical"
        exit 1
    fi
else
    log_ok "redis-server running (PID: ${REDIS_PID})" | tee -a "${REPORT_FILE}"
    # Check if responding
    PONG=$(redis_cmd "${AFFECTED_NODE}" "${REDIS_PORT}" PING 2>/dev/null || echo "FAIL")
    if [[ "${PONG}" != "PONG" ]]; then
        log_error "Redis process running but NOT responding to PING" | tee -a "${REPORT_FILE}"
        log_warn "May be blocked on slow command or loading data" | tee -a "${REPORT_FILE}"
    else
        log_ok "Redis responding to PING" | tee -a "${REPORT_FILE}"
    fi
fi

# Step 4: Check who is master now
log_info "Step 4: Current master from Sentinel..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    MASTER_INFO=$(sentinel_cmd "${node}" SENTINEL get-master-addr-by-name "${SENTINEL_MASTER_NAME}" 2>/dev/null || echo "failed")
    if [[ "${MASTER_INFO}" != "failed" ]]; then
        log_info "Sentinel on ${node} says master is: ${MASTER_INFO}" | tee -a "${REPORT_FILE}"
        break
    fi
done

# Step 5: Replication status of all nodes
log_info "Step 5: Replication status..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "unknown")
    CONNECTED_SLAVES=$(redis_info_field "${node}" "connected_slaves" 2>/dev/null || echo "?")
    MASTER_LINK=$(redis_info_field "${node}" "master_link_status" 2>/dev/null || echo "N/A")
    log_info "  ${node}: role=${ROLE}, connected_slaves=${CONNECTED_SLAVES}, master_link=${MASTER_LINK}" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If host unreachable: Escalate to Infrastructure team. "
    "If OOM killed: Escalate to L3 — review maxmemory settings. "
    "If master was down and failover happened: Verify application reconnection."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-002: Sentinel Down / Quorum Lost
# ═══════════════════════════════════════════════════════════════════
add_issue_block(
    "RB-REDIS-002", "Sentinel Down / Quorum Lost",
    "Dashboard #2 — Sentinel Monitoring & Failover",
    "P1 - Critical",
    [
        "One or more Sentinel instances not responding",
        "Active Sentinels < 3 (warning) or < 2 (CRITICAL — quorum lost)",
        "If quorum lost: Automatic failover is IMPOSSIBLE",
    ],
    [
        "Identify which Sentinel(s) are down: Check Datadog or query each sentinel",
        "SSH to affected node(s)",
        "Check sentinel process: ps aux | grep redis-sentinel",
        "Check systemd: sudo systemctl status redis-sentinel",
        "Check sentinel logs: tail -50 /opt/cached/current/logs/sentinel.log",
        "Attempt restart: sudo systemctl start redis-sentinel",
        "Verify: redis-cli -h localhost -p 26379 SENTINEL masters",
        "Verify quorum: Check num-other-sentinels in SENTINEL masters output",
        "If all sentinels show same master: Configuration is consistent",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-002: Sentinel Down / Quorum Check
# Usage: ./rb-redis-002-sentinel-quorum.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-sentinel-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

SENTINEL_UP=0
SENTINEL_DOWN_LIST=""

# Step 1: Check each Sentinel
log_info "Step 1: Checking all 3 Sentinels..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    PONG=$(sentinel_cmd "${node}" PING 2>/dev/null || echo "FAIL")
    if [[ "${PONG}" == "PONG" ]]; then
        log_ok "  Sentinel on ${node}: UP" | tee -a "${REPORT_FILE}"
        ((SENTINEL_UP++))

        # Get sentinel info
        MASTER_INFO=$(sentinel_cmd "${node}" SENTINEL master "${SENTINEL_MASTER_NAME}" 2>/dev/null)
        NUM_SLAVES=$(echo "${MASTER_INFO}" | grep -A1 "num-slaves" | tail -1 | tr -d '\\r')
        NUM_SENTINELS=$(echo "${MASTER_INFO}" | grep -A1 "num-other-sentinels" | tail -1 | tr -d '\\r')
        QUORUM=$(echo "${MASTER_INFO}" | grep -A1 "quorum" | tail -1 | tr -d '\\r')
        MASTER_IP=$(echo "${MASTER_INFO}" | grep -A1 "^ip$" | tail -1 | tr -d '\\r')
        log_info "    Master: ${MASTER_IP}, Replicas: ${NUM_SLAVES}, Other Sentinels: ${NUM_SENTINELS}, Quorum: ${QUORUM}" | tee -a "${REPORT_FILE}"
    else
        log_error "  Sentinel on ${node}: DOWN" | tee -a "${REPORT_FILE}"
        SENTINEL_DOWN_LIST="${SENTINEL_DOWN_LIST} ${node}"

        # Try to restart
        if check_node_reachable "${node}"; then
            log_info "  Attempting to restart sentinel on ${node}..." | tee -a "${REPORT_FILE}"
            remote_exec "${node}" "sudo systemctl start redis-sentinel" 2>&1 | tee -a "${REPORT_FILE}"
            sleep 5
            PONG2=$(sentinel_cmd "${node}" PING 2>/dev/null || echo "FAIL")
            if [[ "${PONG2}" == "PONG" ]]; then
                log_ok "  Sentinel restarted on ${node}" | tee -a "${REPORT_FILE}"
                ((SENTINEL_UP++))
            else
                log_error "  Sentinel FAILED to restart on ${node}" | tee -a "${REPORT_FILE}"
            fi
        fi
    fi
done

# Step 2: Quorum assessment
echo "" | tee -a "${REPORT_FILE}"
log_info "Step 2: Quorum Assessment..." | tee -a "${REPORT_FILE}"
if [[ ${SENTINEL_UP} -ge 2 ]]; then
    log_ok "Quorum MAINTAINED: ${SENTINEL_UP}/3 sentinels up (need 2)" | tee -a "${REPORT_FILE}"
    if [[ ${SENTINEL_UP} -lt 3 ]]; then
        notify_slack "Sentinel degraded: ${SENTINEL_UP}/3 up. Down:${SENTINEL_DOWN_LIST}. Quorum OK." "warning"
    fi
else
    log_error "QUORUM LOST: Only ${SENTINEL_UP}/3 sentinels up. FAILOVER IMPOSSIBLE!" | tee -a "${REPORT_FILE}"
    notify_slack "CRITICAL: Sentinel quorum LOST (${SENTINEL_UP}/3). Failover disabled! Immediate action required." "critical"
fi

log_info "Report: ${REPORT_FILE}"
""",
    "If quorum lost (< 2 sentinels): P1 CRITICAL — restore sentinel immediately. "
    "If sentinel keeps crashing: Escalate to L3 to investigate logs."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-003: Failover Occurred
# ═══════════════════════════════════════════════════════════════════
add_issue_block(
    "RB-REDIS-003", "Sentinel Failover Occurred (+switch-master)",
    "Dashboard #2 — Sentinel Monitoring & Failover",
    "P1 - Critical",
    [
        "Sentinel log shows +switch-master event",
        "Failover event count incremented in Datadog",
        "Node roles changed (previous master is now replica or down)",
    ],
    [
        "DO NOT PANIC — Sentinel failover is designed to happen automatically",
        "Identify the NEW master: Query any sentinel: SENTINEL get-master-addr-by-name mymaster",
        "Check WHY failover happened: Review sentinel logs for +odown and +sdown events",
        "Verify new master is accepting writes: redis-cli -h <new-master> SET test:failover ok",
        "Verify replicas reconnected to new master: INFO replication on each node",
        "Check the OLD master: Is it up? Did it rejoin as replica?",
        "If old master is up: Verify it's now replicating from new master (role: slave)",
        "Check applications: Are they connected to the new master? (via Sentinel-aware clients)",
        "If frequent failovers (> 2/hour): Investigate root cause — flapping",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-003: Post-Failover Investigation
# Usage: ./rb-redis-003-failover-investigation.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-failover-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"
log_warn "=== FAILOVER INVESTIGATION ===" | tee -a "${REPORT_FILE}"

# Step 1: Who is master now?
log_info "Step 1: Current master..." | tee -a "${REPORT_FILE}"
MASTER_ADDR=""
for node in ${REDIS_NODES}; do
    RESULT=$(sentinel_cmd "${node}" SENTINEL get-master-addr-by-name "${SENTINEL_MASTER_NAME}" 2>/dev/null)
    if [[ -n "${RESULT}" ]]; then
        MASTER_ADDR=$(echo "${RESULT}" | head -1)
        MASTER_PORT=$(echo "${RESULT}" | tail -1)
        log_info "  Master: ${MASTER_ADDR}:${MASTER_PORT} (from sentinel on ${node})" | tee -a "${REPORT_FILE}"
        break
    fi
done

# Step 2: Role of each node
log_info "Step 2: Role of each node..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "UNREACHABLE")
    UPTIME=$(redis_info_field "${node}" "uptime_in_seconds" 2>/dev/null || echo "?")
    CONN_SLAVES=$(redis_info_field "${node}" "connected_slaves" 2>/dev/null || echo "?")
    MASTER_LINK=$(redis_info_field "${node}" "master_link_status" 2>/dev/null || echo "N/A")
    log_info "  ${node}: role=${ROLE}, uptime=${UPTIME}s, slaves=${CONN_SLAVES}, master_link=${MASTER_LINK}" | tee -a "${REPORT_FILE}"
done

# Step 3: Parse sentinel logs for failover timeline
log_info "Step 3: Failover timeline from Sentinel logs..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    log_info "  --- Sentinel on ${node} ---" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "grep -E '(\\+sdown|\\+odown|\\-odown|\\+switch-master|\\+elected-leader|\\+failover)' ${REDIS_LOG_DIR}/sentinel.log 2>/dev/null | tail -20" 2>/dev/null | tee -a "${REPORT_FILE}"
done

# Step 4: Verify new master accepts writes
log_info "Step 4: Testing write to new master..." | tee -a "${REPORT_FILE}"
if [[ -n "${MASTER_ADDR}" ]]; then
    SET_RESULT=$(redis_cmd "${MASTER_ADDR}" "${REDIS_PORT}" SET "__failover_test__" "$(date +%s)" EX 60 2>/dev/null || echo "FAIL")
    if [[ "${SET_RESULT}" == "OK" ]]; then
        log_ok "New master accepts writes" | tee -a "${REPORT_FILE}"
        redis_cmd "${MASTER_ADDR}" "${REDIS_PORT}" DEL "__failover_test__" &>/dev/null
    else
        log_error "New master NOT accepting writes!" | tee -a "${REPORT_FILE}"
    fi
fi

# Step 5: Replication offset comparison
log_info "Step 5: Replication offsets..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    OFFSET=$(redis_info_field "${node}" "master_repl_offset" 2>/dev/null || echo "?")
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "?")
    log_info "  ${node} (${ROLE}): offset=${OFFSET}" | tee -a "${REPORT_FILE}"
done

notify_slack "Failover investigation complete. New master: ${MASTER_ADDR}. Report: ${REPORT_FILE}" "warning"
log_info "Report: ${REPORT_FILE}"
""",
    "If new master is not accepting writes: Escalate to L3 immediately. "
    "If frequent failovers: Investigate network stability, escalate to Network + L3. "
    "Notify Application team to verify their connections are healthy."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-004: Memory Pressure / Evictions
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 2: Memory & Eviction Issues", level=1)

add_issue_block(
    "RB-REDIS-004", "Memory Pressure / Evictions Active",
    "Dashboard #4 — Memory Management & Eviction",
    "P1 - Critical (>95%) / P2 - Warning (>85%)",
    [
        "Memory usage approaching maxmemory limit (>85%)",
        "Evicted keys counter incrementing (DATA LOSS — keys being removed)",
        "Memory allocated widget shows near max line",
        "Fragmentation ratio abnormal (> 1.5 or < 1.0)",
    ],
    [
        "Check current memory: redis-cli INFO memory",
        "Check used_memory vs maxmemory — how close to limit?",
        "Check eviction policy: CONFIG GET maxmemory-policy",
        "If evictions active: This is DATA LOSS — application data is being removed",
        "Check fragmentation ratio: > 1.5 = fragmented; < 1.0 = SWAPPING (critical!)",
        "Identify biggest keys: redis-cli --bigkeys (WARNING: can be slow on large datasets)",
        "Check for expired key backlog: INFO keyspace — ratio of expires to total keys",
        "Short-term fix: Increase maxmemory if host has available RAM",
        "Long-term fix: Optimize data structures, add TTLs, scale cluster",
        "If fragmentation > 1.5: Consider MEMORY PURGE or scheduled restart",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-004: Memory Pressure - Diagnosis & Mitigation
# Usage: ./rb-redis-004-memory-pressure.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

TARGET="${1:-}"
REPORT_FILE="/tmp/redis-memory-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Check all nodes if no target specified
NODES_TO_CHECK="${TARGET:-${REDIS_NODES}}"

for node in ${NODES_TO_CHECK}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"
    MEM_INFO=$(redis_cmd "${node}" "${REDIS_PORT}" INFO memory 2>/dev/null)

    if [[ -z "${MEM_INFO}" ]]; then
        log_error "Cannot connect to ${node}" | tee -a "${REPORT_FILE}"
        continue
    fi

    USED=$(echo "${MEM_INFO}" | grep "^used_memory:" | cut -d: -f2 | tr -d '\\r')
    USED_HR=$(echo "${MEM_INFO}" | grep "^used_memory_human:" | cut -d: -f2 | tr -d '\\r')
    RSS=$(echo "${MEM_INFO}" | grep "^used_memory_rss:" | cut -d: -f2 | tr -d '\\r')
    RSS_HR=$(echo "${MEM_INFO}" | grep "^used_memory_rss_human:" | cut -d: -f2 | tr -d '\\r')
    PEAK=$(echo "${MEM_INFO}" | grep "^used_memory_peak_human:" | cut -d: -f2 | tr -d '\\r')
    MAXMEM=$(echo "${MEM_INFO}" | grep "^maxmemory:" | cut -d: -f2 | tr -d '\\r')
    MAXMEM_HR=$(echo "${MEM_INFO}" | grep "^maxmemory_human:" | cut -d: -f2 | tr -d '\\r')
    FRAG=$(echo "${MEM_INFO}" | grep "^mem_fragmentation_ratio:" | cut -d: -f2 | tr -d '\\r')
    POLICY=$(redis_cmd "${node}" "${REDIS_PORT}" CONFIG GET maxmemory-policy 2>/dev/null | tail -1)

    # Calculate percentage
    if [[ "${MAXMEM}" -gt 0 ]] 2>/dev/null; then
        PCT=$(echo "scale=1; ${USED} * 100 / ${MAXMEM}" | bc 2>/dev/null || echo "?")
    else
        PCT="no limit set"
    fi

    echo "  Used Memory:     ${USED_HR} (${PCT}% of max)" | tee -a "${REPORT_FILE}"
    echo "  RSS Memory:      ${RSS_HR}" | tee -a "${REPORT_FILE}"
    echo "  Peak Memory:     ${PEAK}" | tee -a "${REPORT_FILE}"
    echo "  Max Memory:      ${MAXMEM_HR}" | tee -a "${REPORT_FILE}"
    echo "  Fragmentation:   ${FRAG}" | tee -a "${REPORT_FILE}"
    echo "  Eviction Policy: ${POLICY}" | tee -a "${REPORT_FILE}"

    # Eviction stats
    EVICTED=$(redis_info_field "${node}" "evicted_keys")
    EXPIRED=$(redis_info_field "${node}" "expired_keys")
    echo "  Evicted Keys:    ${EVICTED}" | tee -a "${REPORT_FILE}"
    echo "  Expired Keys:    ${EXPIRED}" | tee -a "${REPORT_FILE}"

    # Warnings
    if echo "${FRAG}" | awk '{exit ($1 < 1.0) ? 0 : 1}' 2>/dev/null; then
        log_error "  CRITICAL: Fragmentation < 1.0 — Redis is SWAPPING to disk!" | tee -a "${REPORT_FILE}"
    elif echo "${FRAG}" | awk '{exit ($1 > 1.5) ? 0 : 1}' 2>/dev/null; then
        log_warn "  WARNING: Fragmentation > 1.5 — Memory is fragmented" | tee -a "${REPORT_FILE}"
    fi

    # Keyspace stats
    log_info "  Keyspace:" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" INFO keyspace 2>/dev/null | grep "^db" | tee -a "${REPORT_FILE}"

    # MEMORY DOCTOR
    log_info "  Memory Doctor:" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" MEMORY DOCTOR 2>/dev/null | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If evictions active: Notify Application team — data is being lost. "
    "If fragmentation < 1.0 (swapping): P1 EMERGENCY — add memory or reduce maxmemory. "
    "If > 95% of maxmemory: Increase maxmemory (requires approval) or scale."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-005: Swap Usage
# ═══════════════════════════════════════════════════════════════════
add_issue_block(
    "RB-REDIS-005", "Swap Usage Detected on Redis Host",
    "Dashboard #9 — Node Resource Utilization",
    "P1 - Critical (EMERGENCY)",
    [
        "system.swap.used > 0 on any Redis host",
        "Latency spikes visible on latency dashboard",
        "Memory fragmentation ratio < 1.0",
        "Redis becomes extremely slow (100x-1000x normal latency)",
    ],
    [
        "THIS IS AN EMERGENCY for Redis — swap destroys Redis performance",
        "Verify swap usage: free -h on affected host",
        "Check what's using memory: top, ps aux --sort=-%mem",
        "Check Redis RSS vs used_memory: redis-cli INFO memory",
        "If fragmentation < 1.0: Redis data is being paged to swap",
        "Immediate fix: Reduce maxmemory so Redis uses less RAM: CONFIG SET maxmemory <lower-value>",
        "Or: Add more physical RAM to the host",
        "Long-term: Disable swap entirely on Redis hosts (production best practice)",
        "After fixing: Verify swap usage drops to 0",
        "Best practice: vm.swappiness=1 or swapoff -a on Redis hosts",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-005: Swap Detection & Emergency Response
# Usage: ./rb-redis-005-swap-check.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-swap-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"
log_error "=== SWAP USAGE CHECK (CRITICAL for Redis) ===" | tee -a "${REPORT_FILE}"

SWAP_FOUND=false
for node in ${REDIS_NODES}; do
    log_info "--- ${node} ---" | tee -a "${REPORT_FILE}"

    # Check swap
    SWAP_INFO=$(remote_exec "${node}" "free -h | grep -i swap" 2>/dev/null)
    echo "  Swap: ${SWAP_INFO}" | tee -a "${REPORT_FILE}"

    SWAP_USED=$(remote_exec "${node}" "free -b | grep -i swap | awk '{print \\$3}'" 2>/dev/null || echo "0")
    if [[ "${SWAP_USED}" -gt 0 ]] 2>/dev/null; then
        log_error "  SWAP IN USE: ${SWAP_USED} bytes" | tee -a "${REPORT_FILE}"
        SWAP_FOUND=true

        # Check Redis memory
        FRAG=$(redis_info_field "${node}" "mem_fragmentation_ratio" 2>/dev/null || echo "?")
        USED=$(redis_info_field "${node}" "used_memory_human" 2>/dev/null || echo "?")
        RSS=$(redis_info_field "${node}" "used_memory_rss_human" 2>/dev/null || echo "?")
        echo "  Redis used_memory: ${USED}" | tee -a "${REPORT_FILE}"
        echo "  Redis RSS:         ${RSS}" | tee -a "${REPORT_FILE}"
        echo "  Fragmentation:     ${FRAG}" | tee -a "${REPORT_FILE}"

        # Check swappiness
        SWAPPINESS=$(remote_exec "${node}" "cat /proc/sys/vm/swappiness" 2>/dev/null || echo "?")
        echo "  vm.swappiness:     ${SWAPPINESS}" | tee -a "${REPORT_FILE}"

        # Redis process swap usage
        REDIS_PID=$(remote_exec "${node}" "pgrep -f redis-server" 2>/dev/null || echo "")
        if [[ -n "${REDIS_PID}" ]]; then
            REDIS_SWAP=$(remote_exec "${node}" "sudo cat /proc/${REDIS_PID}/status | grep VmSwap" 2>/dev/null || echo "unknown")
            echo "  Redis VmSwap:      ${REDIS_SWAP}" | tee -a "${REPORT_FILE}"
        fi
    else
        log_ok "  No swap usage" | tee -a "${REPORT_FILE}"
    fi
done

if ${SWAP_FOUND}; then
    echo "" | tee -a "${REPORT_FILE}"
    log_error "IMMEDIATE ACTIONS:" | tee -a "${REPORT_FILE}"
    echo "  1. Reduce maxmemory: redis-cli CONFIG SET maxmemory <lower_value>" | tee -a "${REPORT_FILE}"
    echo "  2. Set vm.swappiness=1: sudo sysctl vm.swappiness=1" | tee -a "${REPORT_FILE}"
    echo "  3. Long-term: Add RAM or disable swap: sudo swapoff -a" | tee -a "${REPORT_FILE}"
    notify_slack "EMERGENCY: Swap usage detected on Redis hosts. Performance severely degraded." "critical"
fi

log_info "Report: ${REPORT_FILE}"
""",
    "ALWAYS escalate swap on Redis to L3/Principal. "
    "Infrastructure team for RAM addition. "
    "This is a P1 EMERGENCY — Redis latency is 100x-1000x worse with swap."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-006: Replication Lag
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 3: Replication Issues", level=1)

add_issue_block(
    "RB-REDIS-006", "Replication Lag High / Link Down",
    "Dashboard #10 — Replication Deep-Dive",
    "P2 - Warning (>10MB lag) / P1 - Critical (>100MB or link down)",
    [
        "Replication offset lag > 10MB between master and replica(s)",
        "master_link_status showing 'down' for any replica",
        "Full resync events occurring (very expensive)",
        "Replication delay chart showing sustained lag",
    ],
    [
        "Check master_link_status on each replica: redis-cli INFO replication",
        "If link down: Check network between master and replica",
        "Check master load: Is master overloaded (high CPU, many commands)?",
        "Check replica I/O: Is replica disk slow (AOF/RDB writes blocking replication)?",
        "Check output buffer: Is replica output buffer full?",
        "If full resync happened: Check repl-backlog-size — may be too small",
        "If persistent lag: Consider increasing repl-backlog-size",
        "If replica keeps disconnecting: Check timeout settings (repl-timeout)",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-006: Replication Lag Analysis
# Usage: ./rb-redis-006-replication-lag.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-replication-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Find master
MASTER_NODE=""
MASTER_OFFSET=0
for node in ${REDIS_NODES}; do
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "unknown")
    if [[ "${ROLE}" == "master" ]]; then
        MASTER_NODE="${node}"
        MASTER_OFFSET=$(redis_info_field "${node}" "master_repl_offset" 2>/dev/null || echo "0")
        CONNECTED_SLAVES=$(redis_info_field "${node}" "connected_slaves" 2>/dev/null || echo "0")
        BACKLOG_SIZE=$(redis_info_field "${node}" "repl_backlog_size" 2>/dev/null || echo "0")
        BACKLOG_ACTIVE=$(redis_info_field "${node}" "repl_backlog_active" 2>/dev/null || echo "0")
        log_info "Master: ${node}" | tee -a "${REPORT_FILE}"
        echo "  Master offset:     ${MASTER_OFFSET}" | tee -a "${REPORT_FILE}"
        echo "  Connected slaves:  ${CONNECTED_SLAVES}" | tee -a "${REPORT_FILE}"
        echo "  Backlog size:      $((${BACKLOG_SIZE:-0} / 1048576))MB" | tee -a "${REPORT_FILE}"
        echo "  Backlog active:    ${BACKLOG_ACTIVE}" | tee -a "${REPORT_FILE}"

        # Show slave details from master perspective
        for i in 0 1; do
            SLAVE_INFO=$(redis_cmd "${node}" "${REDIS_PORT}" INFO replication 2>/dev/null | grep "^slave${i}:" | cut -d: -f2-)
            if [[ -n "${SLAVE_INFO}" ]]; then
                echo "  Slave${i}: ${SLAVE_INFO}" | tee -a "${REPORT_FILE}"
            fi
        done
        break
    fi
done

# Check each replica
echo "" | tee -a "${REPORT_FILE}"
log_info "Replica Details:" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "unknown")
    if [[ "${ROLE}" == "slave" ]]; then
        LINK_STATUS=$(redis_info_field "${node}" "master_link_status" 2>/dev/null || echo "?")
        LINK_DOWN_SINCE=$(redis_info_field "${node}" "master_link_down_since_seconds" 2>/dev/null || echo "0")
        SLAVE_OFFSET=$(redis_info_field "${node}" "slave_repl_offset" 2>/dev/null || echo "0")
        SLAVE_READ_ONLY=$(redis_info_field "${node}" "slave_read_only" 2>/dev/null || echo "?")
        SYNC_FULL=$(redis_info_field "${node}" "sync_full" 2>/dev/null || echo "0")
        SYNC_PARTIAL_OK=$(redis_info_field "${node}" "sync_partial_ok" 2>/dev/null || echo "0")
        SYNC_PARTIAL_ERR=$(redis_info_field "${node}" "sync_partial_err" 2>/dev/null || echo "0")

        LAG_BYTES=$((${MASTER_OFFSET:-0} - ${SLAVE_OFFSET:-0}))
        LAG_MB=$((LAG_BYTES / 1048576))

        log_info "  ${node} (replica):" | tee -a "${REPORT_FILE}"
        echo "    master_link_status:  ${LINK_STATUS}" | tee -a "${REPORT_FILE}"
        echo "    slave_repl_offset:   ${SLAVE_OFFSET}" | tee -a "${REPORT_FILE}"
        echo "    Lag:                 ${LAG_BYTES} bytes (${LAG_MB}MB)" | tee -a "${REPORT_FILE}"
        echo "    sync_full:           ${SYNC_FULL}" | tee -a "${REPORT_FILE}"
        echo "    sync_partial_ok:     ${SYNC_PARTIAL_OK}" | tee -a "${REPORT_FILE}"
        echo "    sync_partial_err:    ${SYNC_PARTIAL_ERR}" | tee -a "${REPORT_FILE}"

        if [[ "${LINK_STATUS}" != "up" ]]; then
            log_error "    LINK DOWN for ${LINK_DOWN_SINCE}s!" | tee -a "${REPORT_FILE}"
        fi
        if [[ ${LAG_BYTES} -gt ${REDIS_REPL_LAG_CRIT_BYTES:-104857600} ]]; then
            log_error "    LAG CRITICAL: ${LAG_MB}MB" | tee -a "${REPORT_FILE}"
        elif [[ ${LAG_BYTES} -gt ${REDIS_REPL_LAG_WARN_BYTES:-10485760} ]]; then
            log_warn "    LAG WARNING: ${LAG_MB}MB" | tee -a "${REPORT_FILE}"
        fi
    fi
done

log_info "Report: ${REPORT_FILE}"
""",
    "If master_link_status down: Check network, escalate to Network team. "
    "If full resyncs frequent: Increase repl-backlog-size (L3 decision). "
    "If all replicas disconnected: P1 — single point of failure."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-007: Persistence Failure
# ═══════════════════════════════════════════════════════════════════
add_issue_block(
    "RB-REDIS-007", "Persistence Failure (RDB/AOF)",
    "Dashboard #5 — Persistence (RDB & AOF)",
    "P1 - Critical",
    [
        "rdb_last_bgsave_status shows 'err'",
        "aof_last_write_status shows 'err'",
        "Time since last RDB save > 1 hour",
        "AOF file growing without rewrite",
        "AOF fsync delays detected",
    ],
    [
        "Check last save status: redis-cli INFO persistence",
        "Check disk space: df -h /opt/cached/current/data",
        "If disk full: Free space (see RB for disk alarm)",
        "Check file permissions: ls -la /opt/cached/current/data/",
        "If RDB failed: Check for fork failure (not enough memory for COW)",
        "Check overcommit setting: cat /proc/sys/vm/overcommit_memory (should be 1)",
        "Trigger manual save: redis-cli BGSAVE (non-blocking)",
        "For AOF issues: Check redis-cli CONFIG GET appendonly and aof-use-rdb-preamble",
        "If AOF corruption: redis-check-aof --fix <aof-file>",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-007: Persistence Health Check & Recovery
# Usage: ./rb-redis-007-persistence.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

TARGET="${1:-}"
REPORT_FILE="/tmp/redis-persistence-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

NODES_TO_CHECK="${TARGET:-${REDIS_NODES}}"
for node in ${NODES_TO_CHECK}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"
    PERSIST=$(redis_cmd "${node}" "${REDIS_PORT}" INFO persistence 2>/dev/null)

    if [[ -z "${PERSIST}" ]]; then
        log_error "Cannot connect to ${node}" | tee -a "${REPORT_FILE}"
        continue
    fi

    # RDB status
    RDB_STATUS=$(echo "${PERSIST}" | grep "^rdb_last_bgsave_status:" | cut -d: -f2 | tr -d '\\r')
    RDB_LAST_TIME=$(echo "${PERSIST}" | grep "^rdb_last_save_time:" | cut -d: -f2 | tr -d '\\r')
    RDB_CHANGES=$(echo "${PERSIST}" | grep "^rdb_changes_since_last_save:" | cut -d: -f2 | tr -d '\\r')
    RDB_DURATION=$(echo "${PERSIST}" | grep "^rdb_last_bgsave_time_sec:" | cut -d: -f2 | tr -d '\\r')
    RDB_IN_PROG=$(echo "${PERSIST}" | grep "^rdb_bgsave_in_progress:" | cut -d: -f2 | tr -d '\\r')

    SECS_SINCE_SAVE=$(($(date +%s) - ${RDB_LAST_TIME:-0}))
    log_info "  RDB Status:" | tee -a "${REPORT_FILE}"
    echo "    Last save status:       ${RDB_STATUS}" | tee -a "${REPORT_FILE}"
    echo "    Seconds since last save: ${SECS_SINCE_SAVE}" | tee -a "${REPORT_FILE}"
    echo "    Changes since save:      ${RDB_CHANGES}" | tee -a "${REPORT_FILE}"
    echo "    Last save duration:      ${RDB_DURATION}s" | tee -a "${REPORT_FILE}"
    echo "    BGSAVE in progress:      ${RDB_IN_PROG}" | tee -a "${REPORT_FILE}"

    if [[ "${RDB_STATUS}" == "err" ]]; then
        log_error "    RDB SAVE FAILED!" | tee -a "${REPORT_FILE}"
    fi
    if [[ ${SECS_SINCE_SAVE} -gt 3600 ]]; then
        log_warn "    No save for > 1 hour!" | tee -a "${REPORT_FILE}"
    fi

    # AOF status
    AOF_ENABLED=$(echo "${PERSIST}" | grep "^aof_enabled:" | cut -d: -f2 | tr -d '\\r')
    AOF_WRITE_STATUS=$(echo "${PERSIST}" | grep "^aof_last_write_status:" | cut -d: -f2 | tr -d '\\r')
    AOF_SIZE=$(echo "${PERSIST}" | grep "^aof_current_size:" | cut -d: -f2 | tr -d '\\r')
    AOF_BASE_SIZE=$(echo "${PERSIST}" | grep "^aof_base_size:" | cut -d: -f2 | tr -d '\\r')
    AOF_DELAYED=$(echo "${PERSIST}" | grep "^aof_delayed_fsync:" | cut -d: -f2 | tr -d '\\r')

    log_info "  AOF Status:" | tee -a "${REPORT_FILE}"
    echo "    AOF enabled:        ${AOF_ENABLED}" | tee -a "${REPORT_FILE}"
    echo "    Last write status:  ${AOF_WRITE_STATUS}" | tee -a "${REPORT_FILE}"
    echo "    Current size:       $((${AOF_SIZE:-0} / 1048576))MB" | tee -a "${REPORT_FILE}"
    echo "    Base size:          $((${AOF_BASE_SIZE:-0} / 1048576))MB" | tee -a "${REPORT_FILE}"
    echo "    Delayed fsync:      ${AOF_DELAYED}" | tee -a "${REPORT_FILE}"

    if [[ "${AOF_WRITE_STATUS}" == "err" ]]; then
        log_error "    AOF WRITE FAILED!" | tee -a "${REPORT_FILE}"
    fi

    # Disk space
    log_info "  Disk space:" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "df -h ${REDIS_DATA_DIR}" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Overcommit setting
    OVERCOMMIT=$(remote_exec "${node}" "cat /proc/sys/vm/overcommit_memory" 2>/dev/null || echo "?")
    echo "    vm.overcommit_memory: ${OVERCOMMIT} (should be 1 for Redis)" | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If disk full: Escalate to Infrastructure for volume expansion. "
    "If fork failure (not enough memory): Escalate to L3 for memory configuration. "
    "If AOF corruption: Escalate to L3 — data integrity at risk."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-008: High Latency / Slow Commands
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 4: Performance Issues", level=1)

add_issue_block(
    "RB-REDIS-008", "High Latency / Slow Commands",
    "Dashboard #8 — Latency & Slow Queries",
    "P2 - Warning (>5ms) / P1 - Critical (>20ms)",
    [
        "Average latency > 5ms or P99 > 20ms",
        "Slow log entries increasing",
        "KEYS command usage detected (O(n) — blocks Redis)",
        "Fork latency spikes (during BGSAVE/BGREWRITEAOF)",
    ],
    [
        "Check slow log: redis-cli SLOWLOG GET 20",
        "Identify which commands are slow — look for O(n) commands",
        "Check for KEYS command usage: redis-cli INFO commandstats | grep keys",
        "Check for large key operations (HGETALL on huge hash, SMEMBERS on huge set)",
        "Check fork latency: redis-cli INFO persistence | grep latest_fork_usec",
        "If BGSAVE causing spikes: Consider scheduling BGSAVE during off-peak",
        "If KEYS used: Replace with SCAN (non-blocking alternative)",
        "Check if swap is active (see RB-REDIS-005) — #1 cause of sudden latency",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-008: Latency & Slow Command Analysis
# Usage: ./rb-redis-008-latency.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

TARGET="${1:-}"
REPORT_FILE="/tmp/redis-latency-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

NODES_TO_CHECK="${TARGET:-${REDIS_NODES}}"
for node in ${NODES_TO_CHECK}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"

    # Slow log
    log_info "  Slow Log (last 20 entries):" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" SLOWLOG GET 20 2>/dev/null | tee -a "${REPORT_FILE}"

    # Slow log length
    SLOWLOG_LEN=$(redis_cmd "${node}" "${REDIS_PORT}" SLOWLOG LEN 2>/dev/null || echo "?")
    echo "  Total slow log entries: ${SLOWLOG_LEN}" | tee -a "${REPORT_FILE}"

    # Dangerous command usage
    log_info "  Dangerous command stats (O(n) operations):" | tee -a "${REPORT_FILE}"
    for cmd in keys flushall flushdb; do
        STATS=$(redis_cmd "${node}" "${REDIS_PORT}" INFO commandstats 2>/dev/null | grep "cmdstat_${cmd}:" || echo "  ${cmd}: not used")
        echo "    ${STATS}" | tee -a "${REPORT_FILE}"
    done

    # Fork latency
    FORK_USEC=$(redis_info_field "${node}" "latest_fork_usec" 2>/dev/null || echo "?")
    echo "  Latest fork latency: ${FORK_USEC} microseconds" | tee -a "${REPORT_FILE}"

    # Ops per second
    OPS=$(redis_info_field "${node}" "instantaneous_ops_per_sec" 2>/dev/null || echo "?")
    echo "  Current ops/sec: ${OPS}" | tee -a "${REPORT_FILE}"

    # Latency history (if available)
    log_info "  Latency latest:" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" LATENCY LATEST 2>/dev/null | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If KEYS command found: Escalate to Application team to replace with SCAN. "
    "If large key operations: Application team to restructure data. "
    "If fork latency > 500ms: L3 to review BGSAVE schedule."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-009: Connection Issues
# ═══════════════════════════════════════════════════════════════════
add_issue_block(
    "RB-REDIS-009", "Connection Exhaustion / Rejected Connections",
    "Dashboard #6 — Client & Connection Monitoring",
    "P1 - Critical (rejected > 0) / P2 - Warning (>80% of max)",
    [
        "connected_clients approaching maxclients",
        "rejected_connections counter > 0 (clients failing to connect)",
        "Blocked clients count high",
    ],
    [
        "Check current clients: redis-cli INFO clients",
        "Check maxclients: redis-cli CONFIG GET maxclients",
        "If near limit: Identify who is consuming connections",
        "List connections: redis-cli CLIENT LIST | sort by IP",
        "If rejected_connections > 0: Applications ARE failing",
        "Short-term: Increase maxclients: redis-cli CONFIG SET maxclients <higher>",
        "Long-term: Fix connection leaks in application code (use connection pooling)",
        "Kill idle connections: redis-cli CLIENT KILL ID <id>",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-009: Connection Analysis
# Usage: ./rb-redis-009-connections.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

TARGET="${1:-}"
REPORT_FILE="/tmp/redis-connections-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

NODES_TO_CHECK="${TARGET:-${REDIS_NODES}}"
for node in ${NODES_TO_CHECK}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"

    CLIENTS=$(redis_info_field "${node}" "connected_clients")
    BLOCKED=$(redis_info_field "${node}" "blocked_clients")
    MAXCLIENTS=$(redis_cmd "${node}" "${REDIS_PORT}" CONFIG GET maxclients 2>/dev/null | tail -1)
    REJECTED=$(redis_info_field "${node}" "rejected_connections")
    TOTAL_CONNS=$(redis_info_field "${node}" "total_connections_received")

    PCT=0
    if [[ "${MAXCLIENTS}" -gt 0 ]] 2>/dev/null; then
        PCT=$(echo "scale=1; ${CLIENTS:-0} * 100 / ${MAXCLIENTS}" | bc 2>/dev/null || echo "?")
    fi

    echo "  Connected clients: ${CLIENTS} / ${MAXCLIENTS} (${PCT}%)" | tee -a "${REPORT_FILE}"
    echo "  Blocked clients:   ${BLOCKED}" | tee -a "${REPORT_FILE}"
    echo "  Rejected conns:    ${REJECTED}" | tee -a "${REPORT_FILE}"
    echo "  Total received:    ${TOTAL_CONNS}" | tee -a "${REPORT_FILE}"

    if [[ "${REJECTED}" -gt 0 ]] 2>/dev/null; then
        log_error "  CONNECTIONS BEING REJECTED!" | tee -a "${REPORT_FILE}"
    fi

    # Connections by IP (top 15)
    log_info "  Connections by source IP (top 15):" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" CLIENT LIST 2>/dev/null | \
        grep -oP 'addr=\\K[^:]+' | sort | uniq -c | sort -rn | head -15 | \
        while read count ip; do
            echo "    ${ip}: ${count} connections" | tee -a "${REPORT_FILE}"
        done

    # Idle connections (>300s)
    IDLE_COUNT=$(redis_cmd "${node}" "${REDIS_PORT}" CLIENT LIST 2>/dev/null | \
        grep -oP 'idle=\\K[0-9]+' | awk '$1 > 300' | wc -l)
    echo "  Idle connections (>300s): ${IDLE_COUNT}" | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If rejected > 0: P1 — increase maxclients immediately, notify Application team. "
    "If connection leak identified: Application team to fix client connection pooling."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-010: Security / Auth Failures
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 5: Security Issues", level=1)

add_issue_block(
    "RB-REDIS-010", "Authentication Failures / Security Events",
    "Dashboard #11 — Security & ACL Audit",
    "P1 - Critical (spike or FLUSHALL) / P2 - Warning (occasional)",
    [
        "AUTH failure count spiking in logs",
        "NOPERM errors in logs (ACL violations)",
        "FLUSHALL or FLUSHDB command detected",
        "Unknown source IPs attempting connections",
    ],
    [
        "Check Redis logs for AUTH failures: grep 'AUTH\\|WRONGPASS\\|NOPERM' redis.log",
        "Identify source IPs of failed attempts",
        "If credential rotation: Coordinate with application teams",
        "If FLUSHALL/FLUSHDB detected: VERIFY IF INTENTIONAL — this is DATA LOSS",
        "Check ACL users: redis-cli ACL LIST",
        "Check if default user is disabled: redis-cli ACL GETUSER default",
        "If unauthorized access: Block IP at firewall, rotate credentials",
        "Audit ACL changes: Check logs for ACL SETUSER/DELUSER",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-010: Security Audit
# Usage: ./rb-redis-010-security-audit.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-security-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

for node in ${REDIS_NODES}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"

    # ACL Users
    log_info "  ACL Users:" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" ACL LIST 2>/dev/null | tee -a "${REPORT_FILE}"

    # Default user status
    log_info "  Default user:" | tee -a "${REPORT_FILE}"
    redis_cmd "${node}" "${REDIS_PORT}" ACL GETUSER default 2>/dev/null | head -5 | tee -a "${REPORT_FILE}"

    # Dangerous command stats
    log_info "  Dangerous command usage:" | tee -a "${REPORT_FILE}"
    for cmd in flushall flushdb shutdown debug config; do
        STAT=$(redis_cmd "${node}" "${REDIS_PORT}" INFO commandstats 2>/dev/null | grep "cmdstat_${cmd}:" || echo "")
        if [[ -n "${STAT}" ]]; then
            log_warn "    ${STAT}" | tee -a "${REPORT_FILE}"
        fi
    done

    # Auth failures from logs
    log_info "  Recent auth failures from logs:" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "grep -iE 'auth|wrongpass|noperm|denied' ${REDIS_LOG_DIR}/redis.log 2>/dev/null | tail -20" 2>/dev/null | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If FLUSHALL detected without approval: P1 EMERGENCY — potential data loss/breach. "
    "If auth failures from unknown IPs: Escalate to Security team. "
    "Rotate credentials if compromise suspected."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  RB-REDIS-011: Full Cluster Recovery
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 6: Full Cluster Recovery", level=1)

add_issue_block(
    "RB-REDIS-011", "Full Cluster Recovery (All Nodes Down)",
    "Dashboard #12 — Restart & Availability",
    "P1 - Critical",
    [
        "All 3 Redis nodes are down or all recently restarted",
        "Sentinel has no master to monitor",
        "Applications cannot connect to Redis",
    ],
    [
        "CRITICAL ORDER OF OPERATIONS:",
        "1. Start the node that was MASTER last (it has the most recent data)",
        "2. If unsure, start the node with the most recent RDB/AOF file",
        "3. Wait for it to fully load data: redis-cli PING should return PONG",
        "4. Start Sentinel on the master node FIRST",
        "5. Start the other 2 Redis nodes — they will replicate from master",
        "6. Start Sentinel on the remaining 2 nodes",
        "7. Verify: redis-cli -h <master> INFO replication — should show 2 slaves",
        "8. Verify: redis-cli -p 26379 SENTINEL masters — should show correct master",
        "9. Test write/read: SET test:recovery ok, GET test:recovery",
        "10. Notify application teams to verify connectivity",
    ],
    """#!/bin/bash
# ============================================================
# RB-REDIS-011: Full Cluster Recovery (Ordered Startup)
# Usage: ./rb-redis-011-cluster-recovery.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-cluster-recovery-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"
log_error "=== FULL CLUSTER RECOVERY ===" | tee -a "${REPORT_FILE}"
notify_slack "CLUSTER RECOVERY started for ${CLUSTER_NAME}" "critical"

# Step 1: Determine which node to start first
log_info "Step 1: Finding most recent data..." | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    if check_node_reachable "${node}"; then
        # Check RDB file timestamp
        RDB_TIME=$(remote_exec "${node}" "stat -c %Y ${REDIS_DATA_DIR}/dump.rdb 2>/dev/null" || echo "0")
        AOF_TIME=$(remote_exec "${node}" "stat -c %Y ${REDIS_DATA_DIR}/appendonly.aof* 2>/dev/null | sort -rn | head -1" || echo "0")
        LATEST=$((RDB_TIME > AOF_TIME ? RDB_TIME : AOF_TIME))
        LATEST_HR=$(date -d "@${LATEST}" '+%Y-%m-%d %H:%M:%S' 2>/dev/null || echo "unknown")
        log_info "  ${node}: Latest data file: ${LATEST_HR} (epoch: ${LATEST})" | tee -a "${REPORT_FILE}"
    else
        log_error "  ${node}: UNREACHABLE" | tee -a "${REPORT_FILE}"
    fi
done

echo "" | tee -a "${REPORT_FILE}"
log_warn "=== RECOVERY PROCEDURE ===" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "Execute these commands IN ORDER:" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "# Step 1: Start the node with most recent data as MASTER" | tee -a "${REPORT_FILE}"
echo "# (Replace NODE1 with the correct node if different)" | tee -a "${REPORT_FILE}"
echo "ssh ${SSH_USER}@${REDIS_NODE1}" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis" | tee -a "${REPORT_FILE}"
echo "  # Wait for data to load:" | tee -a "${REPORT_FILE}"
echo "  ${REDIS_BIN}/redis-cli -p ${REDIS_PORT} -a '\${REDIS_AUTH_PASS}' PING" | tee -a "${REPORT_FILE}"
echo "  # Verify it's master:" | tee -a "${REPORT_FILE}"
echo "  ${REDIS_BIN}/redis-cli -p ${REDIS_PORT} -a '\${REDIS_AUTH_PASS}' INFO replication | grep role" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "# Step 2: Start Sentinel on master node" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis-sentinel" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "# Step 3: Start remaining Redis nodes (they join as replicas)" | tee -a "${REPORT_FILE}"
echo "ssh ${SSH_USER}@${REDIS_NODE2}" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis" | tee -a "${REPORT_FILE}"
echo "  sleep 10" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis-sentinel" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "ssh ${SSH_USER}@${REDIS_NODE3}" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis" | tee -a "${REPORT_FILE}"
echo "  sleep 10" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start redis-sentinel" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "# Step 4: Verify" | tee -a "${REPORT_FILE}"
echo "${REDIS_BIN}/redis-cli -h ${REDIS_NODE1} -p ${REDIS_PORT} -a '\${REDIS_AUTH_PASS}' INFO replication" | tee -a "${REPORT_FILE}"
echo "${REDIS_BIN}/redis-cli -h ${REDIS_NODE1} -p ${SENTINEL_PORT} SENTINEL masters" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "# Step 5: Test" | tee -a "${REPORT_FILE}"
echo "${REDIS_BIN}/redis-cli -h ${REDIS_NODE1} -p ${REDIS_PORT} -a '\${REDIS_AUTH_PASS}' SET __recovery_test__ ok EX 60" | tee -a "${REPORT_FILE}"
echo "${REDIS_BIN}/redis-cli -h ${REDIS_NODE1} -p ${REDIS_PORT} -a '\${REDIS_AUTH_PASS}' GET __recovery_test__" | tee -a "${REPORT_FILE}"

log_info "Report: ${REPORT_FILE}"
""",
    "If data appears lost after recovery: Escalate to L3/Principal immediately. "
    "If nodes won't join cluster: Escalate to L3. "
    "Create post-incident review after recovery."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  COMPREHENSIVE HEALTH CHECK
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Comprehensive Health Check Script", level=1)
doc.add_paragraph("Run before/after maintenance or for routine health verification.")

add_code("""#!/bin/bash
# ============================================================
# Redis Comprehensive Health Check
# Usage: ./redis-health-check.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/redis_common.sh"
validate_env

REPORT_FILE="/tmp/redis-health-$(date +%Y%m%d-%H%M%S).log"
EXIT_CODE=0
report_header | tee "${REPORT_FILE}"

# ---- Redis Server Health ----
log_info "=== REDIS SERVER HEALTH ===" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    PONG=$(redis_cmd "${node}" "${REDIS_PORT}" PING 2>/dev/null || echo "FAIL")
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "UNKNOWN")
    UPTIME=$(redis_info_field "${node}" "uptime_in_seconds" 2>/dev/null || echo "0")
    UPTIME_HR="$((${UPTIME:-0}/3600))h $((${UPTIME:-0}%3600/60))m"

    if [[ "${PONG}" == "PONG" ]]; then
        log_ok "${node}: UP (role=${ROLE}, uptime=${UPTIME_HR})" | tee -a "${REPORT_FILE}"
    else
        log_error "${node}: DOWN" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
    fi
done

# ---- Sentinel Health ----
log_info "=== SENTINEL HEALTH ===" | tee -a "${REPORT_FILE}"
SENTINEL_COUNT=0
for node in ${REDIS_NODES}; do
    PONG=$(sentinel_cmd "${node}" PING 2>/dev/null || echo "FAIL")
    if [[ "${PONG}" == "PONG" ]]; then
        log_ok "${node}:${SENTINEL_PORT}: Sentinel UP" | tee -a "${REPORT_FILE}"
        ((SENTINEL_COUNT++))
    else
        log_error "${node}:${SENTINEL_PORT}: Sentinel DOWN" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
    fi
done
if [[ ${SENTINEL_COUNT} -lt 2 ]]; then
    log_error "QUORUM AT RISK: Only ${SENTINEL_COUNT}/3 sentinels" | tee -a "${REPORT_FILE}"
fi

# ---- Memory ----
log_info "=== MEMORY ===" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    USED=$(redis_info_field "${node}" "used_memory_human" 2>/dev/null || echo "?")
    MAXMEM=$(redis_info_field "${node}" "maxmemory_human" 2>/dev/null || echo "?")
    FRAG=$(redis_info_field "${node}" "mem_fragmentation_ratio" 2>/dev/null || echo "?")
    EVICTED=$(redis_info_field "${node}" "evicted_keys" 2>/dev/null || echo "?")
    log_info "${node}: used=${USED}, max=${MAXMEM}, frag=${FRAG}, evicted=${EVICTED}" | tee -a "${REPORT_FILE}"
done

# ---- Replication ----
log_info "=== REPLICATION ===" | tee -a "${REPORT_FILE}"
MASTER_OFFSET=0
for node in ${REDIS_NODES}; do
    ROLE=$(redis_info_field "${node}" "role" 2>/dev/null || echo "?")
    if [[ "${ROLE}" == "master" ]]; then
        MASTER_OFFSET=$(redis_info_field "${node}" "master_repl_offset" 2>/dev/null || echo "0")
        SLAVES=$(redis_info_field "${node}" "connected_slaves" 2>/dev/null || echo "0")
        log_info "Master ${node}: offset=${MASTER_OFFSET}, slaves=${SLAVES}" | tee -a "${REPORT_FILE}"
        if [[ "${SLAVES}" -lt 2 ]]; then
            log_warn "Expected 2 replicas, got ${SLAVES}" | tee -a "${REPORT_FILE}"
        fi
    elif [[ "${ROLE}" == "slave" ]]; then
        SLAVE_OFFSET=$(redis_info_field "${node}" "slave_repl_offset" 2>/dev/null || echo "0")
        LINK=$(redis_info_field "${node}" "master_link_status" 2>/dev/null || echo "?")
        LAG=$(( ${MASTER_OFFSET:-0} - ${SLAVE_OFFSET:-0} ))
        log_info "Replica ${node}: offset=${SLAVE_OFFSET}, lag=${LAG}bytes, link=${LINK}" | tee -a "${REPORT_FILE}"
    fi
done

# ---- Persistence ----
log_info "=== PERSISTENCE ===" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    RDB_STATUS=$(redis_info_field "${node}" "rdb_last_bgsave_status" 2>/dev/null || echo "?")
    AOF_STATUS=$(redis_info_field "${node}" "aof_last_write_status" 2>/dev/null || echo "?")
    log_info "${node}: RDB=${RDB_STATUS}, AOF=${AOF_STATUS}" | tee -a "${REPORT_FILE}"
    if [[ "${RDB_STATUS}" == "err" || "${AOF_STATUS}" == "err" ]]; then
        log_error "PERSISTENCE FAILURE on ${node}!" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
    fi
done

# ---- Performance ----
log_info "=== PERFORMANCE ===" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    OPS=$(redis_info_field "${node}" "instantaneous_ops_per_sec" 2>/dev/null || echo "?")
    HITS=$(redis_info_field "${node}" "keyspace_hits" 2>/dev/null || echo "0")
    MISSES=$(redis_info_field "${node}" "keyspace_misses" 2>/dev/null || echo "0")
    if [[ $((HITS + MISSES)) -gt 0 ]]; then
        HIT_RATE=$(echo "scale=1; ${HITS} * 100 / (${HITS} + ${MISSES})" | bc 2>/dev/null || echo "?")
    else
        HIT_RATE="N/A"
    fi
    CLIENTS=$(redis_info_field "${node}" "connected_clients" 2>/dev/null || echo "?")
    log_info "${node}: ops/s=${OPS}, hit_rate=${HIT_RATE}%, clients=${CLIENTS}" | tee -a "${REPORT_FILE}"
done

# ---- Swap Check ----
log_info "=== SWAP CHECK ===" | tee -a "${REPORT_FILE}"
for node in ${REDIS_NODES}; do
    SWAP=$(remote_exec "${node}" "free -b | grep -i swap | awk '{print \\$3}'" 2>/dev/null || echo "?")
    if [[ "${SWAP}" -gt 0 ]] 2>/dev/null; then
        log_error "${node}: SWAP IN USE (${SWAP} bytes) — CRITICAL!" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
    else
        log_ok "${node}: No swap" | tee -a "${REPORT_FILE}"
    fi
done

echo "" | tee -a "${REPORT_FILE}"
if [[ ${EXIT_CODE} -eq 0 ]]; then
    log_ok "=== ALL CHECKS PASSED ===" | tee -a "${REPORT_FILE}"
else
    log_error "=== ISSUES FOUND — SEE ABOVE ===" | tee -a "${REPORT_FILE}"
fi

log_info "Report: ${REPORT_FILE}"
exit ${EXIT_CODE}
""")

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Redis HA Cluster Operations Runbook — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output = "/Users/tejasodanapalli/rabbitmq/RabbitMQ/redis-cluster-deployment/docs/Redis_Operations_Runbook.docx"
doc.save(output)
print(f"Document saved: {output}")
