#!/usr/bin/env python3
"""
Generate Datadog Monitoring Gap Analysis Document
for RabbitMQ 3-Node Cluster and Redis 3-Node HA Cluster.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

def add_colored_table(headers, rows):
    """Table with color coding: green for Present, red/orange for Missing."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Datadog Monitoring Gap Analysis\nRabbitMQ & Redis Clusters")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph("")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Current State Assessment — What Is Monitored vs What Is Missing")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")
doc.add_paragraph("")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Version: 1.0",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "Classification: Internal / Infrastructure Team",
    "Prepared by: Monitoring Architecture Team",
]:
    meta.add_run(line + "\n").font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Executive Summary",
    "2. RabbitMQ — Currently Configured Widgets",
    "3. RabbitMQ — Missing Widgets (Critical)",
    "4. RabbitMQ — Missing Widgets (Important)",
    "5. RabbitMQ — Missing Widgets (Operational)",
    "6. RabbitMQ — Coverage Scorecard",
    "7. Redis — Currently Configured Widgets",
    "8. Redis — Missing Widgets (Critical)",
    "9. Redis — Missing Widgets (Important)",
    "10. Redis — Missing Widgets (Operational)",
    "11. Redis — Coverage Scorecard",
    "12. Combined Summary & Risk Assessment",
    "13. Recommended Implementation Priority",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document provides a gap analysis of the current Datadog monitoring dashboards "
    "for the RabbitMQ 3-node cluster and the Redis 3-node HA cluster (with Sentinel). "
    "It identifies what is currently being monitored versus what is missing, categorized "
    "by severity of the monitoring gap."
)

doc.add_heading("Quick Summary", level=2)
add_table(
    ["System", "Widgets Present", "Critical Missing", "Important Missing", "Operational Missing", "Coverage %"],
    [
        ["RabbitMQ", "9", "12", "8", "6", "~25%"],
        ["Redis", "12", "10", "9", "7", "~32%"],
    ]
)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Key Risk: ")
run.bold = True
p.add_run(
    "Both dashboards focus primarily on resource metrics (CPU, Memory, Threads) and basic throughput. "
    "Neither dashboard has adequate coverage for availability detection (node up/down), "
    "alarm/alert states, security monitoring, or persistence health. A node could go down, "
    "a memory alarm could block all publishers, or a security breach could occur — and the "
    "current dashboards would not show it."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  2. RABBITMQ — CURRENTLY CONFIGURED
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("2. RabbitMQ — Currently Configured Widgets", level=1)
doc.add_paragraph(
    "The following 9 widgets are currently present on the Datadog dashboard for RabbitMQ:"
)

add_table(
    ["#", "Widget Name", "Metric Type", "Scope", "Assessment"],
    [
        ["1", "Total CPU Utilization%", "Resource — CPU", "Per node (3 nodes)", "Good — shows per-node CPU with avg/max/value"],
        ["2", "RSS Memory Utilization%", "Resource — Memory", "Per node (3 nodes)", "Good — shows RSS with 90% threshold line"],
        ["3", "Total Threads", "Resource — Threads", "Per node (3 nodes)", "OK — less critical for RabbitMQ (Erlang uses processes not threads)"],
        ["4", "Queue Depth", "Queue — Total Messages", "Per queue (top queues)", "Good — shows depth by queue name over time"],
        ["5", "Ready Messages in Queue", "Queue — Ready", "Per queue (top queues)", "Good — ready messages pending delivery"],
        ["6", "Message Deliver Rate", "Throughput — Deliver", "Per queue (msg/sec)", "Good — delivery throughput"],
        ["7", "Message Publish Rate", "Throughput — Publish", "Per queue (msg/sec)", "Good — publish throughput"],
        ["8", "Messages Published to Exchanges", "Throughput — Exchange In", "Per exchange (cumulative)", "OK — shows total published to exchanges"],
        ["9", "Messages Routed from Exchanges", "Throughput — Exchange Out", "Per exchange (cumulative)", "OK — shows total routed from exchanges"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  3. RABBITMQ — MISSING CRITICAL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("3. RabbitMQ — Missing Widgets (Critical)", level=1)
doc.add_paragraph(
    "These are high-severity gaps. Without these, you cannot detect outages, "
    "resource alarms, or data loss conditions."
)

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Why It's Critical", "Risk If Missing"],
    [
        ["1", "Node Up/Down Status",
         "rabbitmq.node.running / rabbitmq.aliveness",
         "No visibility if any of the 3 nodes goes down",
         "A node can be down for hours without anyone knowing"],
        ["2", "Memory Alarm Status",
         "rabbitmq.node.mem_alarm",
         "When triggered, ALL publishers are BLOCKED cluster-wide",
         "Complete publishing outage with no dashboard indication"],
        ["3", "Disk Alarm Status",
         "rabbitmq.node.disk_alarm",
         "When triggered, ALL publishers are BLOCKED cluster-wide",
         "Complete publishing outage with no dashboard indication"],
        ["4", "Unacknowledged Messages",
         "rabbitmq.queue.messages.unacknowledged",
         "Messages delivered but not acked = stuck/crashed consumers",
         "Silent consumer failure — messages pile up, never processed"],
        ["5", "Connected Clients / Connections",
         "rabbitmq.connections",
         "Connection leaks or storms crash the cluster",
         "Cluster crash from FD/memory exhaustion with no warning"],
        ["6", "Channel Count",
         "rabbitmq.channels",
         "Channel leaks are a top RabbitMQ failure mode",
         "Cluster crash from channel leak with no warning"],
        ["7", "File Descriptor Usage %",
         "rabbitmq.node.fd_used / fd_total",
         "FD exhaustion kills the node — cannot accept connections",
         "Node becomes unreachable with no prior warning on dashboard"],
        ["8", "Disk Free Space",
         "rabbitmq.node.disk_free",
         "Disk full triggers disk alarm, blocking all publishers",
         "Disk alarm triggers with no visibility into remaining space"],
        ["9", "Network Partition Status",
         "rabbitmq.node.partitions",
         "Most dangerous failure — split-brain causes DATA LOSS",
         "Data loss and message duplication without any alert"],
        ["10", "Node Uptime",
         "rabbitmq.node.uptime",
         "Cannot detect when a node restarted (planned or crash)",
         "No restart detection, no RCA capability"],
        ["11", "Consumer Count per Queue",
         "rabbitmq.queue.consumers",
         "If consumers die, queue depth grows but root cause invisible",
         "Queue depth alert triggers but you don't know WHY (no consumers)"],
        ["12", "Redelivered Messages Rate",
         "rabbitmq.overview.messages.redeliver.rate",
         "Messages being rejected/nacked by consumers repeatedly",
         "Consumer processing failures go undetected"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  4. RABBITMQ — MISSING IMPORTANT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("4. RabbitMQ — Missing Widgets (Important)", level=1)
doc.add_paragraph(
    "These provide deeper operational visibility and are needed for effective troubleshooting "
    "and capacity planning."
)

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Why It's Important"],
    [
        ["1", "Publish vs Deliver Rate Overlay",
         "publish.rate vs deliver.rate (overlaid)",
         "When publish > deliver, backlog is growing — currently you have them in separate charts making gap detection difficult"],
        ["2", "Unroutable / Returned Messages",
         "rabbitmq.overview.messages.return_unroutable",
         "Messages published to exchange with no matching queue = SILENT data loss"],
        ["3", "Dead Letter Queue (DLQ) Depth",
         "rabbitmq.queue.messages{queue:*dlq*}",
         "Failed messages accumulating — no visibility into processing failures"],
        ["4", "Blocked Connections",
         "rabbitmq.connections.blocked",
         "Connections blocked by flow control = publishers impacted"],
        ["5", "Consumer Utilization %",
         "rabbitmq.queue.consumer_utilisation",
         "Are consumers actually processing or sitting idle?"],
        ["6", "Erlang Process Usage %",
         "rabbitmq.node.proc_used / proc_total",
         "Erlang process exhaustion crashes the node (different from OS threads)"],
        ["7", "Quorum Queue Leader Distribution",
         "Queue leaders per node",
         "Leader imbalance = one node handling disproportionate load"],
        ["8", "Acknowledge Rate",
         "rabbitmq.overview.messages.ack.rate",
         "Consumer acknowledgment speed — gap between deliver and ack = processing backlog"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  5. RABBITMQ — MISSING OPERATIONAL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("5. RabbitMQ — Missing Widgets (Operational)", level=1)
doc.add_paragraph("Nice-to-have for mature monitoring, cleanup, and long-term operations.")

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Purpose"],
    [
        ["1", "Total Queue Count", "rabbitmq.overview.queues", "Detect queue leak (apps creating queues without cleanup)"],
        ["2", "Queues with 0 Consumers", "consumers == 0 AND messages > 0", "Orphaned queues accumulating messages"],
        ["3", "Connection Churn Rate", "rate(connections opened/closed)", "Reconnect storm detection"],
        ["4", "Erlang GC Rate", "rabbitmq_erlang_gc_runs_total", "Memory pressure and GC overhead"],
        ["5", "IO Read/Write Latency", "rabbitmq.node.io_read/write_avg_time", "Disk bottleneck detection for persistence"],
        ["6", "Cluster Event / Error Log Stream", "source:rabbitmq status:error", "Real-time error visibility for correlation"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  6. RABBITMQ — COVERAGE SCORECARD
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("6. RabbitMQ — Coverage Scorecard", level=1)

add_table(
    ["Monitoring Category", "Configured", "Missing", "Coverage %", "Risk Level"],
    [
        ["Node Health & Availability", "3 (CPU, Memory, Threads)", "7 (Up/Down, FD, Disk, Uptime, Erlang Procs, Partitions, Alarms)", "30%", "HIGH"],
        ["Queue Monitoring", "2 (Depth, Ready)", "5 (Unacked, Consumers, DLQ, Queue Count, Orphaned Queues)", "29%", "HIGH"],
        ["Message Throughput", "4 (Deliver, Publish, Exchange Pub, Exchange Route)", "4 (Overlay, Redeliver, Unroutable, Ack Rate)", "50%", "MEDIUM"],
        ["Connection & Channel", "0", "4 (Connections, Channels, Blocked, Churn)", "0%", "CRITICAL"],
        ["Persistence / Replication", "0", "3 (Quorum Leaders, Raft Elections, Replica Status)", "0%", "HIGH"],
        ["Security & Audit", "0", "Not on this dashboard", "0%", "HIGH"],
        ["OVERALL", "9 widgets", "~26 missing", "~25%", "HIGH"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  7. REDIS — CURRENTLY CONFIGURED
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Redis — Currently Configured Widgets", level=1)
doc.add_paragraph(
    "The following 12 widgets are currently present on the Datadog dashboard for Redis:"
)

add_table(
    ["#", "Widget Name", "Metric Type", "Scope", "Assessment"],
    [
        ["1", "Redis CPU",
         "Resource — CPU",
         "Per node (3 nodes, avg/min/max)",
         "Good — shows Redis process CPU per node"],
        ["2", "Total Threads",
         "Resource — Threads",
         "Per node (3 nodes)",
         "OK — less relevant for Redis (single-threaded command processing)"],
        ["3", "RSS Memory Utilization%",
         "Resource — Memory",
         "Per node (3 nodes, 90% threshold)",
         "Good — shows RSS memory with threshold line"],
        ["4", "Hit Rate",
         "Performance — Cache",
         "Cluster aggregate (87%)",
         "Good — cache effectiveness at a glance"],
        ["5", "Primary Link Down",
         "Replication — Status",
         "Cluster aggregate (0 secs)",
         "Good — shows how long primary link has been down"],
        ["6", "Redis Slaves Connected",
         "Replication — Count",
         "Per node (2 conns each, min/max/sum/value)",
         "Good — shows connected replica count per node"],
        ["7", "Blocked Clients",
         "Client — Blocked",
         "Cluster aggregate (0 conns)",
         "Good — clients in blocking commands (BLPOP etc.)"],
        ["8", "Keys to Expire",
         "Keyspace — TTL",
         "Cluster aggregate (42.94k keys)",
         "Good — keys with TTL set"],
        ["9", "Max Connected Clients",
         "Client — Connections",
         "Per node (64k conns, min/max/sum/value)",
         "Good — shows connection count with max capacity"],
        ["10", "Latency to Connect with Master",
         "Replication — Latency",
         "Per replica (avg/min/max in ms)",
         "Good — replication connection latency"],
        ["11", "Replication Delay",
         "Replication — Lag",
         "Per node (offsets, avg/min/max)",
         "Good — replication offset delay between master and replicas"],
        ["12", "Memory Allocated by Redis",
         "Resource — Memory",
         "Per node (with max memory reached line)",
         "Good — shows used memory vs maxmemory limit"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  8. REDIS — MISSING CRITICAL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Redis — Missing Widgets (Critical)", level=1)
doc.add_paragraph(
    "These are high-severity gaps. Without these, you cannot detect node failures, "
    "data loss conditions, or critical performance degradation."
)

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Why It's Critical", "Risk If Missing"],
    [
        ["1", "Node Up/Down Status (Redis Server)",
         "redis.can_connect",
         "No explicit status showing if each Redis server is reachable",
         "A node can be down — you only see it indirectly via missing data points"],
        ["2", "Sentinel Status (All 3 Sentinels)",
         "redis.can_connect{port:26379}",
         "No Sentinel monitoring at all — Sentinel manages failover",
         "Sentinel quorum lost = automatic failover IMPOSSIBLE, no warning"],
        ["3", "Current Master Identity",
         "Sentinel MASTER info / redis.info.role",
         "No visibility into which node is the current master",
         "After failover, team doesn't know who the new master is"],
        ["4", "Failover Event Tracking",
         "Log: +switch-master event",
         "No visibility when Sentinel performs a failover",
         "Master changed silently — applications may have stale connections"],
        ["5", "Swap Usage",
         "system.swap.used",
         "Redis + swap = catastrophic performance (100x-1000x latency)",
         "Redis swapping to disk with no alert — severe latency invisible"],
        ["6", "Evicted Keys Rate",
         "rate(redis.stats.evicted_keys)",
         "Keys being removed due to memory pressure = DATA LOSS",
         "Application data silently evicted with no visibility"],
        ["7", "Memory Usage % of Max",
         "redis.mem.used / maxmemory * 100",
         "Memory Allocated widget exists but percentage proximity to limit is clearer",
         "Hard to see how close to eviction/OOM at a glance"],
        ["8", "Rejected Connections",
         "redis.stats.rejected_connections",
         "Connections rejected because maxclients reached",
         "Applications failing to connect — no dashboard visibility"],
        ["9", "Commands/sec (Throughput)",
         "rate(redis.net.commands)",
         "No throughput metric — cannot see if Redis is processing requests",
         "Redis could stop serving and dashboard wouldn't show it"],
        ["10", "RDB/AOF Persistence Status",
         "rdb_last_bgsave_status / aof_last_write_status",
         "No persistence health monitoring",
         "Backups failing silently — data recovery impossible after crash"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  9. REDIS — MISSING IMPORTANT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Redis — Missing Widgets (Important)", level=1)
doc.add_paragraph(
    "These provide deeper operational insight for troubleshooting and capacity planning."
)

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Why It's Important"],
    [
        ["1", "Memory Fragmentation Ratio",
         "redis.mem.fragmentation_ratio",
         "> 1.5 = memory fragmented (waste); < 1.0 = SWAPPING (critical). Currently invisible"],
        ["2", "Slow Log Count & Entries",
         "SLOWLOG GET / custom metric",
         "Slow commands causing latency spikes — no visibility into what commands are slow"],
        ["3", "Node Uptime",
         "redis.info.uptime_in_seconds",
         "Cannot detect when a Redis node restarted — no restart tracking"],
        ["4", "Full Resync Count",
         "redis.stats.sync_full",
         "Full resyncs are expensive (full data transfer) — not tracked"],
        ["5", "Keyspace Misses/sec",
         "rate(redis.stats.keyspace_misses)",
         "Hit Rate is shown but misses trend is not — can't see miss rate spikes"],
        ["6", "Total Key Count",
         "redis.keys",
         "Keys to Expire is shown but total key count is missing — can't calculate TTL coverage %"],
        ["7", "Disk I/O Latency",
         "system.io.await",
         "Disk latency affects RDB saves and AOF writes — not monitored"],
        ["8", "Client Output Buffer Size",
         "redis.clients.longest_output_list",
         "Large output buffers = slow consumers or pub/sub backlog"],
        ["9", "Expired Keys/sec",
         "rate(redis.stats.expired_keys)",
         "Keys to Expire count is shown but the RATE of expiration is not — can't detect expiry storms"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  10. REDIS — MISSING OPERATIONAL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Redis — Missing Widgets (Operational)", level=1)
doc.add_paragraph("Nice-to-have for mature monitoring and long-term operations.")

add_table(
    ["#", "Missing Widget", "Datadog Metric", "Purpose"],
    [
        ["1", "Command Type Distribution", "redis.command.calls by {command}", "Which commands are used most (GET vs SET vs HGET etc.)"],
        ["2", "Connection Churn Rate", "rate(redis.stats.total_connections_received)", "Detect reconnect storms"],
        ["3", "ACL / Auth Failure Events", "Log: AUTH failed / WRONGPASS", "Security monitoring for unauthorized access attempts"],
        ["4", "AOF File Size Trend", "redis.aof.size", "AOF growth tracking for disk capacity planning"],
        ["5", "RDB Save Duration", "redis.rdb.last_bgsave_time_sec", "Are snapshots taking too long?"],
        ["6", "Lua Script Memory", "redis.mem.lua", "Lua scripts consuming excessive memory"],
        ["7", "TLS Certificate Expiry", "Custom check", "Certificate management for encrypted connections"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  11. REDIS — COVERAGE SCORECARD
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Redis — Coverage Scorecard", level=1)

add_table(
    ["Monitoring Category", "Configured", "Missing", "Coverage %", "Risk Level"],
    [
        ["Node Health & Availability", "3 (CPU, Threads, RSS Memory)", "3 (Up/Down, Uptime, Swap)", "50%", "HIGH"],
        ["Replication", "4 (Link Down, Slaves Connected, Latency, Delay)", "2 (Full Resync, Master Identity)", "67%", "MEDIUM"],
        ["Memory Management", "2 (Memory Allocated, Keys to Expire)", "4 (Usage %, Fragmentation, Evictions, Expired Rate)", "33%", "HIGH"],
        ["Sentinel & Failover", "0", "4 (Sentinel Status, Quorum, Failover Events, Master Tracking)", "0%", "CRITICAL"],
        ["Performance & Throughput", "1 (Hit Rate)", "4 (Commands/sec, Slow Log, Command Distribution, Latency)", "20%", "HIGH"],
        ["Client & Connections", "2 (Max Connected, Blocked)", "3 (Rejected, Churn, Output Buffers)", "40%", "MEDIUM"],
        ["Persistence (RDB/AOF)", "0", "5 (Save Status, AOF Status, Save Duration, File Size, Fsync)", "0%", "CRITICAL"],
        ["Security", "0", "3 (Auth Failures, ACL Events, TLS Cert)", "0%", "HIGH"],
        ["OVERALL", "12 widgets", "~26 missing", "~32%", "HIGH"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  12. COMBINED SUMMARY & RISK ASSESSMENT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Combined Summary & Risk Assessment", level=1)

doc.add_heading("12.1 Overall Monitoring Posture", level=2)
add_table(
    ["System", "Present", "Missing (Critical)", "Missing (Important)", "Missing (Operational)", "Total Missing", "Coverage"],
    [
        ["RabbitMQ", "9", "12", "8", "6", "26", "~25%"],
        ["Redis", "12", "10", "9", "7", "26", "~32%"],
        ["TOTAL", "21", "22", "17", "13", "52", "~29%"],
    ]
)

doc.add_heading("12.2 Critical Blind Spots (Both Systems)", level=2)
doc.add_paragraph(
    "The following critical capabilities are missing from BOTH the RabbitMQ and Redis dashboards:"
)
blind_spots = [
    ("Node Availability Detection", "Neither dashboard has an explicit node UP/DOWN status widget. A node can be down and the team relies on indirect indicators (missing data points) to notice."),
    ("Sentinel / Failover Monitoring (Redis)", "Zero Sentinel monitoring. The system that manages automatic failover is completely unmonitored. If Sentinel quorum is lost, automatic failover becomes impossible with no alert."),
    ("Persistence / Data Durability (Redis)", "No RDB or AOF health monitoring. If persistence fails, a crash results in complete data loss."),
    ("Memory Alarms (RabbitMQ)", "Memory and disk alarms block ALL publishers cluster-wide. These critical states have zero visibility on the dashboard."),
    ("Connection / Channel Monitoring (RabbitMQ)", "Zero connection or channel monitoring. Connection leaks and channel storms are top RabbitMQ failure causes."),
    ("Security & Access Monitoring", "Neither system has authentication failure tracking, permission change auditing, or ACL violation monitoring."),
    ("Restart / Uptime Tracking", "Neither system tracks node restarts. When a node restarts (planned or crash), there is no event captured on the dashboard."),
]
for title_text, desc in blind_spots:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("12.3 What the Current Dashboards Do Well", level=2)
strengths = [
    "Resource metrics (CPU, Memory, Threads) are covered for both systems across all nodes",
    "RabbitMQ message throughput (publish/deliver rates, exchange metrics) provides good traffic visibility",
    "Redis replication monitoring (slaves connected, replication delay, link down, latency) is the strongest area",
    "Redis hit rate provides cache effectiveness at a glance",
    "Redis blocked clients and max connected clients provide basic connection awareness",
    "Both dashboards show per-node breakdown allowing node-level comparison",
]
for s in strengths:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  13. RECOMMENDED IMPLEMENTATION PRIORITY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Recommended Implementation Priority", level=1)

doc.add_heading("13.1 Phase 1 — Immediate (Week 1-2): Critical Gaps", level=2)
doc.add_paragraph("Address the blind spots that can cause undetected outages:")

add_table(
    ["Priority", "System", "Widget to Add", "Effort"],
    [
        ["1", "RabbitMQ", "Node Up/Down Status (rabbitmq.node.running)", "Low — native integration metric"],
        ["2", "RabbitMQ", "Memory Alarm + Disk Alarm Status", "Low — native integration metric"],
        ["3", "RabbitMQ", "Connected Clients + Channel Count", "Low — native integration metric"],
        ["4", "RabbitMQ", "File Descriptor Usage %", "Low — native integration metric"],
        ["5", "RabbitMQ", "Unacknowledged Messages", "Low — native integration metric"],
        ["6", "RabbitMQ", "Consumer Count per Queue", "Low — native integration metric"],
        ["7", "Redis", "Sentinel Status (all 3 sentinels)", "Medium — need 2nd Redis integration instance on port 26379"],
        ["8", "Redis", "Failover Event Tracking (log-based)", "Medium — needs log pipeline + pattern matching"],
        ["9", "Redis", "Commands/sec (Throughput)", "Low — native integration metric"],
        ["10", "Redis", "RDB Save Status + AOF Write Status", "Low — native integration metric"],
        ["11", "Redis", "Evicted Keys Rate", "Low — native integration metric"],
        ["12", "Redis", "Swap Usage per Host", "Low — Datadog Agent system metric"],
    ]
)

doc.add_heading("13.2 Phase 2 — Short-Term (Week 3-4): Important Gaps", level=2)
add_table(
    ["Priority", "System", "Widget to Add", "Effort"],
    [
        ["13", "RabbitMQ", "Network Partition Status", "Low — native metric"],
        ["14", "RabbitMQ", "Node Uptime + Restart Detection", "Low — native metric"],
        ["15", "RabbitMQ", "Redelivered Messages Rate", "Low — native metric"],
        ["16", "RabbitMQ", "Publish vs Deliver Overlay Chart", "Low — dashboard edit only"],
        ["17", "RabbitMQ", "Dead Letter Queue Depth", "Low — filter existing queue metric"],
        ["18", "Redis", "Memory Fragmentation Ratio", "Low — native metric"],
        ["19", "Redis", "Node Uptime + Restart Detection", "Low — native metric"],
        ["20", "Redis", "Slow Log Entries", "Medium — custom check needed"],
        ["21", "Redis", "Memory Usage % Gauge", "Low — formula widget"],
        ["22", "Redis", "Rejected Connections", "Low — native metric"],
        ["23", "Redis", "Total Key Count + TTL Coverage %", "Low — native metric"],
    ]
)

doc.add_heading("13.3 Phase 3 — Medium-Term (Month 2): Operational Gaps", level=2)
add_table(
    ["Priority", "System", "Widget to Add", "Effort"],
    [
        ["24", "RabbitMQ", "Queue Count Total + Orphaned Queues", "Low"],
        ["25", "RabbitMQ", "Erlang Process Usage %", "Low"],
        ["26", "RabbitMQ", "Connection Churn Rate", "Low"],
        ["27", "RabbitMQ", "Quorum Queue Leader Distribution", "Medium"],
        ["28", "Redis", "Command Type Distribution", "Low"],
        ["29", "Redis", "Connection Churn Rate", "Low"],
        ["30", "Redis", "AOF File Size Trend", "Low"],
        ["31", "Redis", "Auth Failure Monitoring", "Medium — log-based"],
        ["32", "Both", "Error Log Stream Widget", "Medium"],
        ["33", "Both", "Alerting / Monitor Setup", "Medium — separate effort"],
    ]
)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Note: ")
run.bold = True
p.add_run(
    "Most Phase 1 and Phase 2 items are Low effort because they use metrics already collected "
    "by the Datadog Redis and RabbitMQ integrations — they just need to be added as widgets to "
    "the existing dashboard. No new data collection is needed for the majority of these."
)

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Datadog Monitoring Gap Analysis — RabbitMQ & Redis — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output = "/Users/tejasodanapalli/rabbitmq/RabbitMQ/docs/monitoring/Datadog_Monitoring_Gap_Analysis_RabbitMQ_Redis.docx"
doc.save(output)
print(f"Document saved: {output}")
