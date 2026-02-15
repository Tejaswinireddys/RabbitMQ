#!/usr/bin/env python3
"""
Generate RabbitMQ Operations Runbook Document
Mapped to each Datadog Dashboard — For L1/L2 Engineers
Includes manual steps + automation scripts (parameterized)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

def add_bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)
    return p

def add_issue_block(issue_id, title, dashboard, severity, symptoms, manual_steps, script, escalation):
    """Standard issue block for every runbook entry."""
    doc.add_heading(f"{issue_id}: {title}", level=2)

    add_table(
        ["Field", "Value"],
        [
            ["Dashboard", dashboard],
            ["Severity", severity],
            ["L1/L2 Actionable", "Yes"],
            ["Estimated Resolution Time", "5-30 minutes"],
        ]
    )
    doc.add_paragraph("")

    add_bold_para("Symptoms / What You See on Dashboard:", "")
    for s in symptoms:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph("")
    add_bold_para("Manual Steps (L1/L2):", "")
    for i, step in enumerate(manual_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Bullet')

    doc.add_paragraph("")
    add_bold_para("Automation Script:", "")
    add_code(script)

    doc.add_paragraph("")
    add_bold_para("Escalation:", "")
    doc.add_paragraph(escalation)
    doc.add_paragraph("")

# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph("")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RabbitMQ 3-Node Cluster\nOperations Runbook")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph("")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Dashboard-Mapped Issue Resolution Guide\nFor L1/L2 Engineers")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Version: 1.0",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "Classification: Internal / Operations",
    "Author: Principal Engineering / SRE Team",
]:
    meta.add_run(line + "\n").font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  ENVIRONMENT VARIABLES REFERENCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Environment Variables Reference", level=1)
doc.add_paragraph(
    "All scripts in this runbook use parameterized environment variables. Before executing "
    "any script, source the appropriate environment file or set these variables:"
)

add_code("""# ============================================================
# RabbitMQ Environment Configuration
# File: /opt/rabbitmq/scripts/env/<environment>.env
# ============================================================

# --- Cluster Nodes ---
export RMQ_NODE1="rabbitmq-node-1.example.com"
export RMQ_NODE2="rabbitmq-node-2.example.com"
export RMQ_NODE3="rabbitmq-node-3.example.com"
export RMQ_NODES="${RMQ_NODE1} ${RMQ_NODE2} ${RMQ_NODE3}"

# --- Ports ---
export RMQ_AMQP_PORT="5672"
export RMQ_MANAGEMENT_PORT="15672"
export RMQ_PROMETHEUS_PORT="15692"

# --- Authentication ---
export RMQ_ADMIN_USER="admin"
export RMQ_ADMIN_PASS="<from-vault>"       # NEVER hardcode in scripts
export RMQ_MONITOR_USER="monitoring"
export RMQ_MONITOR_PASS="<from-vault>"

# --- Paths ---
export RMQ_HOME="/opt/rabbitmq"
export RMQ_CONF_DIR="/etc/rabbitmq"
export RMQ_LOG_DIR="/var/log/rabbitmq"
export RMQ_DATA_DIR="/var/lib/rabbitmq"
export RMQ_SCRIPTS_DIR="/opt/rabbitmq/scripts"

# --- Thresholds ---
export RMQ_QUEUE_DEPTH_WARN=10000
export RMQ_QUEUE_DEPTH_CRIT=50000
export RMQ_MEM_WARN_PCT=80
export RMQ_MEM_CRIT_PCT=90
export RMQ_FD_WARN_PCT=80
export RMQ_DISK_FREE_WARN_GB=5

# --- SSH ---
export SSH_USER="sre-user"
export SSH_KEY="/home/${SSH_USER}/.ssh/id_rsa"
export SSH_OPTS="-o StrictHostKeyChecking=no -o ConnectTimeout=5"

# --- Notification ---
export SLACK_WEBHOOK_URL="https://hooks.slack.com/services/XXX/YYY/ZZZ"
export PAGERDUTY_SERVICE_KEY="<from-vault>"
export OPS_EMAIL="sre-team@company.com"

# --- Environment Identifier ---
export ENVIRONMENT="production"    # production / staging / qa / dev
export CLUSTER_NAME="rmq-cluster-01"

# Usage: source /opt/rabbitmq/scripts/env/production.env""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  COMMON FUNCTIONS LIBRARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Common Functions Library", level=1)
doc.add_paragraph(
    "All automation scripts source this shared library. Place at "
    "/opt/rabbitmq/scripts/lib/common.sh"
)

add_code("""#!/bin/bash
# ============================================================
# RabbitMQ Runbook - Common Functions Library
# File: /opt/rabbitmq/scripts/lib/common.sh
# ============================================================

set -euo pipefail

# --- Colors ---
RED='\\033[0;31m'; GREEN='\\033[0;32m'; YELLOW='\\033[1;33m'; NC='\\033[0m'

# --- Logging ---
log_info()  { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] [INFO]  $*"; }
log_warn()  { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${YELLOW}[WARN]${NC}  $*"; }
log_error() { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${RED}[ERROR]${NC} $*"; }
log_ok()    { echo -e "[$(date '+%Y-%m-%d %H:%M:%S')] ${GREEN}[OK]${NC}    $*"; }

# --- Validate environment loaded ---
validate_env() {
    local required_vars="RMQ_NODE1 RMQ_NODE2 RMQ_NODE3 RMQ_MANAGEMENT_PORT RMQ_ADMIN_USER"
    for var in ${required_vars}; do
        if [[ -z "${!var:-}" ]]; then
            log_error "Environment variable ${var} is not set. Source your env file first."
            log_error "Usage: source /opt/rabbitmq/scripts/env/<environment>.env"
            exit 1
        fi
    done
    log_info "Environment validated: ${ENVIRONMENT:-unknown} / ${CLUSTER_NAME:-unknown}"
}

# --- API call helper ---
rmq_api() {
    local node="${1}"; local endpoint="${2}"; shift 2
    curl -s -u "${RMQ_ADMIN_USER}:${RMQ_ADMIN_PASS}" \\
        --connect-timeout 5 --max-time 10 \\
        "http://${node}:${RMQ_MANAGEMENT_PORT}/api/${endpoint}" "$@"
}

# --- SSH helper ---
remote_exec() {
    local node="${1}"; shift
    ssh ${SSH_OPTS} -i "${SSH_KEY}" "${SSH_USER}@${node}" "$@"
}

# --- Notification ---
notify_slack() {
    local message="${1}"; local severity="${2:-info}"
    local color="good"
    [[ "${severity}" == "warning" ]] && color="warning"
    [[ "${severity}" == "critical" ]] && color="danger"

    curl -s -X POST "${SLACK_WEBHOOK_URL}" \\
        -H 'Content-type: application/json' \\
        -d "{
            \\"attachments\\": [{
                \\"color\\": \\"${color}\\",
                \\"title\\": \\"RabbitMQ Runbook - ${ENVIRONMENT}\\",
                \\"text\\": \\"${message}\\",
                \\"footer\\": \\"Cluster: ${CLUSTER_NAME} | $(date '+%Y-%m-%d %H:%M:%S')\\",
                \\"mrkdwn_in\\": [\\"text\\"]
            }]
        }" > /dev/null 2>&1 || true
}

# --- Check if node is reachable ---
check_node_reachable() {
    local node="${1}"
    if ssh ${SSH_OPTS} -i "${SSH_KEY}" "${SSH_USER}@${node}" "echo ok" &>/dev/null; then
        return 0
    else
        return 1
    fi
}

# --- Check RabbitMQ service status ---
check_rmq_status() {
    local node="${1}"
    remote_exec "${node}" "sudo rabbitmqctl status" 2>/dev/null
}

# --- Get cluster status ---
get_cluster_status() {
    local node="${1:-${RMQ_NODE1}}"
    remote_exec "${node}" "sudo rabbitmqctl cluster_status --formatter json" 2>/dev/null
}

# --- Generate report header ---
report_header() {
    echo "============================================================"
    echo "  RabbitMQ Runbook Execution Report"
    echo "  Environment: ${ENVIRONMENT} | Cluster: ${CLUSTER_NAME}"
    echo "  Executed by: $(whoami) at $(date '+%Y-%m-%d %H:%M:%S')"
    echo "============================================================"
}""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 1 — CLUSTER HEALTH RUNBOOK
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 1: Cluster Health & Node Issues", level=1)
doc.add_paragraph("Dashboard: Cluster Health Overview")
doc.add_paragraph(
    "This runbook covers issues detected on the primary cluster health dashboard: "
    "node failures, resource alarms, and Erlang VM problems."
)

# Issue 1.1 - Node Down
add_issue_block(
    "RB-RMQ-001", "RabbitMQ Node Down",
    "Dashboard #1 — Cluster Health Overview",
    "P1 - Critical",
    [
        "Node status widget shows RED / 0 for one or more nodes",
        "Uptime widget shows 0 or N/A for affected node",
        "Metrics from that node stop appearing on all charts",
        "Connected clients drop on that node",
    ],
    [
        "Verify the alert — check which node is reported down in Datadog",
        "SSH to the affected node: ssh <user>@<node>",
        "Check if the host is reachable: ping <node>. If not, escalate to Infrastructure team",
        "Check if RabbitMQ process is running: sudo systemctl status rabbitmq-server",
        "Check Erlang beam process: ps aux | grep beam.smp",
        "Review recent logs: sudo tail -100 /var/log/rabbitmq/rabbit@<node>.log",
        "If process is down, attempt restart: sudo systemctl start rabbitmq-server",
        "Wait 30 seconds, then verify: sudo rabbitmqctl status",
        "Check if node rejoined cluster: sudo rabbitmqctl cluster_status",
        "If node won't start, check for mnesia lock: ls -la /var/lib/rabbitmq/mnesia/",
        "If last node to stop, may need force boot: sudo rabbitmqctl force_boot && sudo systemctl start rabbitmq-server",
        "Verify queue synchronization: sudo rabbitmqctl list_queues name type leader online",
        "Verify in Datadog that node status returns to GREEN",
        "Document findings in incident ticket",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-001: RabbitMQ Node Down - Diagnosis & Recovery
# Usage: ./rb-rmq-001-node-down.sh <node_hostname>
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

AFFECTED_NODE="${1:?Usage: $0 <node_hostname>}"
REPORT_FILE="/tmp/rmq-node-down-$(date +%Y%m%d-%H%M%S).log"

report_header | tee "${REPORT_FILE}"
log_info "Investigating node: ${AFFECTED_NODE}" | tee -a "${REPORT_FILE}"

# Step 1: Check host reachability
log_info "Step 1: Checking host reachability..." | tee -a "${REPORT_FILE}"
if ! ping -c 3 -W 2 "${AFFECTED_NODE}" &>/dev/null; then
    log_error "Host ${AFFECTED_NODE} is NOT reachable via ping" | tee -a "${REPORT_FILE}"
    notify_slack "NODE DOWN: ${AFFECTED_NODE} is unreachable via ping. Escalate to Infra." "critical"
    echo "ACTION: Escalate to Infrastructure team — host is down" | tee -a "${REPORT_FILE}"
    exit 1
fi
log_ok "Host is reachable via ping" | tee -a "${REPORT_FILE}"

# Step 2: Check SSH
log_info "Step 2: Checking SSH access..." | tee -a "${REPORT_FILE}"
if ! check_node_reachable "${AFFECTED_NODE}"; then
    log_error "Cannot SSH to ${AFFECTED_NODE}" | tee -a "${REPORT_FILE}"
    notify_slack "NODE DOWN: Cannot SSH to ${AFFECTED_NODE}. Check SSH service." "critical"
    exit 1
fi
log_ok "SSH access confirmed" | tee -a "${REPORT_FILE}"

# Step 3: Check RabbitMQ process
log_info "Step 3: Checking RabbitMQ process..." | tee -a "${REPORT_FILE}"
RMQ_PID=$(remote_exec "${AFFECTED_NODE}" "pgrep -f beam.smp" 2>/dev/null || echo "")
if [[ -z "${RMQ_PID}" ]]; then
    log_error "RabbitMQ process (beam.smp) is NOT running" | tee -a "${REPORT_FILE}"

    # Step 4: Check why it stopped
    log_info "Step 4: Checking recent logs for shutdown reason..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "sudo tail -50 ${RMQ_LOG_DIR}/rabbit@*.log" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Step 5: Check disk space
    log_info "Step 5: Checking disk space..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "df -h ${RMQ_DATA_DIR}" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Step 6: Attempt restart
    log_info "Step 6: Attempting to start RabbitMQ..." | tee -a "${REPORT_FILE}"
    remote_exec "${AFFECTED_NODE}" "sudo systemctl start rabbitmq-server" 2>&1 | tee -a "${REPORT_FILE}"

    sleep 30

    # Step 7: Verify
    log_info "Step 7: Verifying restart..." | tee -a "${REPORT_FILE}"
    RMQ_STATUS=$(remote_exec "${AFFECTED_NODE}" "sudo rabbitmqctl status" 2>&1)
    if echo "${RMQ_STATUS}" | grep -q "uptime"; then
        log_ok "RabbitMQ started successfully" | tee -a "${REPORT_FILE}"
        notify_slack "NODE RECOVERED: ${AFFECTED_NODE} RabbitMQ restarted successfully." "good"
    else
        log_error "RabbitMQ failed to start. Attempting force_boot..." | tee -a "${REPORT_FILE}"
        remote_exec "${AFFECTED_NODE}" "sudo rabbitmqctl force_boot" 2>&1 | tee -a "${REPORT_FILE}"
        remote_exec "${AFFECTED_NODE}" "sudo systemctl start rabbitmq-server" 2>&1 | tee -a "${REPORT_FILE}"
        sleep 30
        RMQ_STATUS=$(remote_exec "${AFFECTED_NODE}" "sudo rabbitmqctl status" 2>&1)
        if echo "${RMQ_STATUS}" | grep -q "uptime"; then
            log_ok "RabbitMQ started after force_boot" | tee -a "${REPORT_FILE}"
            notify_slack "NODE RECOVERED (force_boot): ${AFFECTED_NODE}. Verify data integrity." "warning"
        else
            log_error "RabbitMQ FAILED to start even after force_boot" | tee -a "${REPORT_FILE}"
            notify_slack "NODE CRITICAL: ${AFFECTED_NODE} cannot start. Manual intervention required." "critical"
            exit 1
        fi
    fi
else
    log_ok "RabbitMQ process is running (PID: ${RMQ_PID})" | tee -a "${REPORT_FILE}"
fi

# Step 8: Verify cluster membership
log_info "Step 8: Checking cluster membership..." | tee -a "${REPORT_FILE}"
remote_exec "${AFFECTED_NODE}" "sudo rabbitmqctl cluster_status" 2>&1 | tee -a "${REPORT_FILE}"

# Step 9: Check queue sync
log_info "Step 9: Checking queue synchronization..." | tee -a "${REPORT_FILE}"
remote_exec "${AFFECTED_NODE}" "sudo rabbitmqctl list_queues name type messages consumers leader online" 2>&1 | tee -a "${REPORT_FILE}"

log_info "Report saved to: ${REPORT_FILE}"
log_info "Runbook RB-RMQ-001 complete."
""",
    "If node cannot be recovered after force_boot: Escalate to L3/Principal Engineer. "
    "If host is unreachable: Escalate to Infrastructure team. "
    "If multiple nodes are down: Follow Cluster Recovery runbook RB-RMQ-012."
)

doc.add_page_break()

# Issue 1.2 - Memory Alarm
add_issue_block(
    "RB-RMQ-002", "Memory Alarm Triggered — Publishers Blocked",
    "Dashboard #1 — Cluster Health Overview",
    "P1 - Critical",
    [
        "Memory alarm widget shows RED / Active for one or more nodes",
        "Memory usage gauge shows > 90% of watermark",
        "Publish rate drops to zero or near-zero",
        "Connected publishers may show blocked state",
    ],
    [
        "Identify which node(s) have the memory alarm: Check Datadog or run rabbitmqctl status",
        "SSH to affected node and check memory breakdown: rabbitmqctl status | grep -A20 memory",
        "Identify top memory consumers: Check if queues, connections, or binaries are using most memory",
        "Check for large queues: rabbitmqctl list_queues name messages memory --sort-by memory --reverse --limit 20",
        "Check for connection/channel leaks: rabbitmqctl list_connections name channels send_oct recv_oct",
        "If queue buildup: Check consumer health (are consumers running? are they stuck?)",
        "If safe to purge: rabbitmqctl purge_queue <queue_name> (ONLY with approval)",
        "Temporary relief: Increase memory watermark: rabbitmqctl set_vm_memory_high_watermark 0.7",
        "Monitor: Watch memory usage drop below watermark in Datadog",
        "Root cause: Investigate why memory grew (consumer failure, traffic spike, memory leak)",
        "Revert temporary watermark change after root cause is fixed",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-002: Memory Alarm - Diagnosis & Mitigation
# Usage: ./rb-rmq-002-memory-alarm.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

TARGET_NODE="${1:-${RMQ_NODE1}}"
REPORT_FILE="/tmp/rmq-memory-alarm-$(date +%Y%m%d-%H%M%S).log"

report_header | tee "${REPORT_FILE}"
log_info "Investigating memory alarm on: ${TARGET_NODE}" | tee -a "${REPORT_FILE}"

# Step 1: Check memory status on all nodes
log_info "Step 1: Memory status across cluster..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    log_info "--- Node: ${node} ---" | tee -a "${REPORT_FILE}"
    MEM_INFO=$(rmq_api "${node}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
for n in nodes:
    name=n.get('name','?')
    mem_used=n.get('mem_used',0)
    mem_limit=n.get('mem_limit',1)
    mem_alarm=n.get('mem_alarm',False)
    pct=round(mem_used/mem_limit*100,1)
    print(f'{name}: {mem_used//1048576}MB / {mem_limit//1048576}MB ({pct}%) alarm={mem_alarm}')
" 2>/dev/null || echo "Failed to query ${node}")
    echo "${MEM_INFO}" | tee -a "${REPORT_FILE}"
done

# Step 2: Top 20 queues by memory
log_info "Step 2: Top 20 queues by memory usage..." | tee -a "${REPORT_FILE}"
rmq_api "${TARGET_NODE}" "queues?sort=memory&sort_reverse=true&page_size=20&columns=name,vhost,messages,memory,consumers,state" 2>/dev/null | \
    python3 -c "
import json,sys
queues=json.load(sys.stdin)
items=queues if isinstance(queues,list) else queues.get('items',[])
print(f'{\"Queue\":<50} {\"Messages\":>10} {\"Memory(MB)\":>12} {\"Consumers\":>10}')
print('-'*85)
for q in items[:20]:
    name=q.get('name','?')[:48]
    msgs=q.get('messages',0)
    mem=round(q.get('memory',0)/1048576,1)
    cons=q.get('consumers',0)
    print(f'{name:<50} {msgs:>10} {mem:>12} {cons:>10}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 3: Connection count and channel count
log_info "Step 3: Connection and channel counts..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    OVERVIEW=$(rmq_api "${node}" "overview" 2>/dev/null)
    CONNS=$(echo "${OVERVIEW}" | python3 -c "import json,sys; d=json.load(sys.stdin); print(d.get('object_totals',{}).get('connections',0))" 2>/dev/null || echo "?")
    CHANS=$(echo "${OVERVIEW}" | python3 -c "import json,sys; d=json.load(sys.stdin); print(d.get('object_totals',{}).get('channels',0))" 2>/dev/null || echo "?")
    log_info "${node}: connections=${CONNS}, channels=${CHANS}" | tee -a "${REPORT_FILE}"
done

# Step 4: Check for queues with 0 consumers
log_info "Step 4: Queues with messages but 0 consumers..." | tee -a "${REPORT_FILE}"
rmq_api "${TARGET_NODE}" "queues?page_size=500&columns=name,messages,consumers" 2>/dev/null | \
    python3 -c "
import json,sys
queues=json.load(sys.stdin)
items=queues if isinstance(queues,list) else queues.get('items',[])
orphaned=[q for q in items if q.get('consumers',0)==0 and q.get('messages',0)>0]
if orphaned:
    print(f'FOUND {len(orphaned)} queues with messages but NO consumers:')
    for q in sorted(orphaned, key=lambda x: x.get('messages',0), reverse=True)[:10]:
        print(f'  {q[\"name\"]}: {q[\"messages\"]} messages')
else:
    print('All queues with messages have consumers.')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 5: If alarm still active, offer temporary watermark increase
log_info "Step 5: Checking if alarm is still active..." | tee -a "${REPORT_FILE}"
ALARM_ACTIVE=$(rmq_api "${TARGET_NODE}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
alarms=[n['name'] for n in nodes if n.get('mem_alarm',False)]
print(','.join(alarms) if alarms else 'none')
" 2>/dev/null || echo "unknown")

if [[ "${ALARM_ACTIVE}" != "none" ]]; then
    log_warn "Memory alarm STILL ACTIVE on: ${ALARM_ACTIVE}" | tee -a "${REPORT_FILE}"
    log_warn "To temporarily increase watermark (requires approval):" | tee -a "${REPORT_FILE}"
    echo "  ssh <node> sudo rabbitmqctl set_vm_memory_high_watermark 0.7" | tee -a "${REPORT_FILE}"
    notify_slack "MEMORY ALARM active on ${ALARM_ACTIVE}. Publishers BLOCKED. Investigation report: ${REPORT_FILE}" "critical"
else
    log_ok "No memory alarms currently active" | tee -a "${REPORT_FILE}"
fi

log_info "Report saved to: ${REPORT_FILE}"
""",
    "If memory cannot be freed and alarm persists > 15 min: Escalate to L3. "
    "If caused by single large queue that cannot be purged: Escalate to Application team. "
    "If caused by connection/channel leak: Escalate to Application team to fix client code."
)

doc.add_page_break()

# Issue 1.3 - Disk Alarm
add_issue_block(
    "RB-RMQ-003", "Disk Alarm Triggered — Publishers Blocked",
    "Dashboard #1 — Cluster Health Overview",
    "P1 - Critical",
    [
        "Disk alarm widget shows RED / Active",
        "Disk free space gauge shows critically low",
        "Publish rate drops to zero",
    ],
    [
        "Identify which node(s) have the disk alarm",
        "SSH to affected node: check disk space with df -h",
        "Check RabbitMQ data directory size: du -sh /var/lib/rabbitmq/mnesia/",
        "Check log directory size: du -sh /var/log/rabbitmq/",
        "Rotate/compress old logs: logrotate -f /etc/logrotate.d/rabbitmq",
        "Remove old crash dumps if present: rm /var/log/rabbitmq/erl_crash.dump",
        "Check for large queue data files in mnesia directory",
        "If disk still full: Expand disk volume (AWS EBS, etc.) or add disk",
        "Verify disk alarm clears in Datadog",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-003: Disk Alarm - Diagnosis & Recovery
# Usage: ./rb-rmq-003-disk-alarm.sh [node_hostname]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

TARGET_NODE="${1:-${RMQ_NODE1}}"
REPORT_FILE="/tmp/rmq-disk-alarm-$(date +%Y%m%d-%H%M%S).log"

report_header | tee "${REPORT_FILE}"

# Step 1: Disk space on all nodes
log_info "Step 1: Disk space on all nodes..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    log_info "--- ${node} ---" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "df -h ${RMQ_DATA_DIR} ${RMQ_LOG_DIR} /tmp" 2>/dev/null | tee -a "${REPORT_FILE}"
done

# Step 2: Largest directories
log_info "Step 2: Largest directories in RabbitMQ data..." | tee -a "${REPORT_FILE}"
remote_exec "${TARGET_NODE}" "sudo du -sh ${RMQ_DATA_DIR}/mnesia/* 2>/dev/null | sort -rh | head -20" | tee -a "${REPORT_FILE}"

# Step 3: Log file sizes
log_info "Step 3: Log files..." | tee -a "${REPORT_FILE}"
remote_exec "${TARGET_NODE}" "sudo ls -lh ${RMQ_LOG_DIR}/" | tee -a "${REPORT_FILE}"

# Step 4: Cleanup old logs
log_info "Step 4: Rotating and cleaning old logs..." | tee -a "${REPORT_FILE}"
remote_exec "${TARGET_NODE}" "sudo logrotate -f /etc/logrotate.d/rabbitmq 2>/dev/null || true" | tee -a "${REPORT_FILE}"

# Remove crash dumps
CRASH_DUMP="${RMQ_LOG_DIR}/erl_crash.dump"
if remote_exec "${TARGET_NODE}" "test -f ${CRASH_DUMP}" 2>/dev/null; then
    DUMP_SIZE=$(remote_exec "${TARGET_NODE}" "du -h ${CRASH_DUMP}" 2>/dev/null)
    log_warn "Found crash dump: ${DUMP_SIZE}. Removing..." | tee -a "${REPORT_FILE}"
    remote_exec "${TARGET_NODE}" "sudo rm -f ${CRASH_DUMP}" 2>/dev/null
    log_ok "Crash dump removed" | tee -a "${REPORT_FILE}"
fi

# Remove old compressed logs
remote_exec "${TARGET_NODE}" "sudo find ${RMQ_LOG_DIR} -name '*.gz' -mtime +7 -delete 2>/dev/null" | tee -a "${REPORT_FILE}"
log_info "Cleaned logs older than 7 days" | tee -a "${REPORT_FILE}"

# Step 5: Check /tmp
log_info "Step 5: Checking /tmp..." | tee -a "${REPORT_FILE}"
remote_exec "${TARGET_NODE}" "sudo du -sh /tmp/* 2>/dev/null | sort -rh | head -10" | tee -a "${REPORT_FILE}"

# Step 6: Recheck disk space
log_info "Step 6: Rechecking disk space..." | tee -a "${REPORT_FILE}"
remote_exec "${TARGET_NODE}" "df -h ${RMQ_DATA_DIR}" | tee -a "${REPORT_FILE}"

# Step 7: Check alarm status
ALARM=$(rmq_api "${TARGET_NODE}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
for n in nodes:
    if n.get('disk_free_alarm',False):
        print(f'ACTIVE on {n[\"name\"]}: disk_free={n.get(\"disk_free\",0)//1048576}MB, limit={n.get(\"disk_free_limit\",0)//1048576}MB')
" 2>/dev/null || echo "unknown")

if [[ -n "${ALARM}" && "${ALARM}" != "" ]]; then
    log_warn "Disk alarm still active: ${ALARM}" | tee -a "${REPORT_FILE}"
    notify_slack "DISK ALARM still active after cleanup. May need volume expansion. ${ALARM}" "critical"
else
    log_ok "No disk alarms active" | tee -a "${REPORT_FILE}"
    notify_slack "Disk alarm resolved after cleanup on ${TARGET_NODE}." "good"
fi

log_info "Report: ${REPORT_FILE}"
""",
    "If disk cannot be freed sufficiently: Escalate to Infrastructure team for volume expansion. "
    "If data directory is consuming excessive space: Escalate to L3 to investigate queue data."
)

doc.add_page_break()

# Issue 1.4 - Network Partition
add_issue_block(
    "RB-RMQ-004", "Network Partition Detected (Split-Brain)",
    "Dashboard #1 — Cluster Health / Dashboard #10 — Partition",
    "P0 - EMERGENCY",
    [
        "Partition status shows > 0 partitions",
        "Nodes show different cluster membership views",
        "Error logs show 'Mnesia partitioned' or 'inconsistent_cluster'",
    ],
    [
        "DO NOT PANIC. Do NOT restart nodes randomly — this can make it worse",
        "Identify which nodes are in which partition: rabbitmqctl cluster_status on EACH node",
        "Check network connectivity between all node pairs: ping, telnet port 4369 and 25672",
        "Determine the partition handling strategy configured: Check rabbitmq.conf for cluster_partition_handling",
        "If autoheal: RabbitMQ should auto-recover. Monitor and verify",
        "If pause_minority: Minority-side nodes will pause. Identify and restart them after network is fixed",
        "If ignore: Manual resolution required — you must choose which partition to keep",
        "Fix the underlying network issue FIRST before attempting recovery",
        "After network is fixed: Restart the nodes that were in the MINORITY partition",
        "Verify cluster is unified: rabbitmqctl cluster_status shows all 3 nodes",
        "Verify queue synchronization: rabbitmqctl list_queues name type leader online",
        "CHECK FOR DATA INCONSISTENCY: Compare message counts before and after",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-004: Network Partition - Diagnosis & Resolution
# Usage: ./rb-rmq-004-network-partition.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-partition-$(date +%Y%m%d-%H%M%S).log"

report_header | tee "${REPORT_FILE}"
log_error "=== NETWORK PARTITION INVESTIGATION ===" | tee -a "${REPORT_FILE}"
notify_slack "EMERGENCY: Network partition investigation started on ${CLUSTER_NAME}" "critical"

# Step 1: Get cluster status from EACH node
log_info "Step 1: Cluster status from each node's perspective..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    log_info "=== View from ${node} ===" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "sudo rabbitmqctl cluster_status" 2>&1 | tee -a "${REPORT_FILE}"
    echo "" | tee -a "${REPORT_FILE}"
done

# Step 2: Network connectivity matrix
log_info "Step 2: Network connectivity between nodes..." | tee -a "${REPORT_FILE}"
NODES_ARRAY=(${RMQ_NODES})
for src in "${NODES_ARRAY[@]}"; do
    for dst in "${NODES_ARRAY[@]}"; do
        if [[ "${src}" != "${dst}" ]]; then
            # AMQP port
            AMQP_OK=$(remote_exec "${src}" "timeout 3 bash -c 'echo > /dev/tcp/${dst}/${RMQ_AMQP_PORT}' 2>/dev/null && echo OK || echo FAIL" 2>/dev/null || echo "SSH_FAIL")
            # Erlang distribution port
            DIST_OK=$(remote_exec "${src}" "timeout 3 bash -c 'echo > /dev/tcp/${dst}/25672' 2>/dev/null && echo OK || echo FAIL" 2>/dev/null || echo "SSH_FAIL")
            # EPMD port
            EPMD_OK=$(remote_exec "${src}" "timeout 3 bash -c 'echo > /dev/tcp/${dst}/4369' 2>/dev/null && echo OK || echo FAIL" 2>/dev/null || echo "SSH_FAIL")
            # Latency
            LATENCY=$(remote_exec "${src}" "ping -c 3 -W 2 ${dst} 2>/dev/null | tail -1 | awk -F/ '{print \\$5}'" 2>/dev/null || echo "?")
            log_info "${src} -> ${dst}: AMQP=${AMQP_OK}, Dist=${DIST_OK}, EPMD=${EPMD_OK}, Latency=${LATENCY}ms" | tee -a "${REPORT_FILE}"
        fi
    done
done

# Step 3: Check partition handling mode
log_info "Step 3: Partition handling strategy..." | tee -a "${REPORT_FILE}"
PARTITION_MODE=$(remote_exec "${RMQ_NODE1}" "sudo rabbitmqctl eval 'application:get_env(rabbit, cluster_partition_handling).' 2>/dev/null" || echo "unknown")
log_info "Partition handling mode: ${PARTITION_MODE}" | tee -a "${REPORT_FILE}"

# Step 4: Check for alarms
log_info "Step 4: Active alarms..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    ALARMS=$(remote_exec "${node}" "sudo rabbitmqctl status 2>/dev/null | grep -A5 'alarms'" || echo "unknown")
    log_info "${node} alarms: ${ALARMS}" | tee -a "${REPORT_FILE}"
done

log_error "=== ACTION REQUIRED ===" | tee -a "${REPORT_FILE}"
echo "1. Fix network issue between partitioned nodes FIRST" | tee -a "${REPORT_FILE}"
echo "2. After network is fixed, restart MINORITY partition nodes:" | tee -a "${REPORT_FILE}"
echo "   sudo rabbitmqctl stop_app && sudo rabbitmqctl start_app" | tee -a "${REPORT_FILE}"
echo "3. Verify unified cluster: sudo rabbitmqctl cluster_status" | tee -a "${REPORT_FILE}"
echo "4. Verify queue sync: sudo rabbitmqctl list_queues name type leader online" | tee -a "${REPORT_FILE}"
echo "5. CHECK FOR DATA LOSS — compare message counts" | tee -a "${REPORT_FILE}"

notify_slack "Partition report generated: ${REPORT_FILE}. Manual resolution required." "critical"
log_info "Report: ${REPORT_FILE}"
""",
    "ALWAYS escalate to L3/Principal Engineer for network partitions. "
    "Network team must be involved to fix underlying connectivity. "
    "DO NOT restart nodes until network is verified fixed."
)

doc.add_page_break()

# Issue 1.5 - FD Exhaustion
add_issue_block(
    "RB-RMQ-005", "File Descriptor Usage High (>80%)",
    "Dashboard #1 — Cluster Health Overview",
    "P2 - Warning (P1 if >95%)",
    [
        "FD usage gauge shows > 80% for one or more nodes",
        "New connections may start failing",
        "Error logs may show 'too many open files'",
    ],
    [
        "Check current FD usage: rabbitmqctl status | grep file_descriptors",
        "Check what is consuming FDs — connections are the primary consumer",
        "List connections: rabbitmqctl list_connections name peer_host channels state --limit 50",
        "Look for connection leaks: Many connections from same client/IP",
        "If connection leak identified: Notify application team to fix client code",
        "Temporary relief: Close idle connections: rabbitmqctl close_connection <pid> 'FD pressure'",
        "Long-term fix: Increase FD limit in systemd unit file: LimitNOFILE=500000",
        "Restart required for new FD limit to take effect",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-005: File Descriptor Usage High
# Usage: ./rb-rmq-005-fd-high.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-fd-usage-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

for node in ${RMQ_NODES}; do
    log_info "=== ${node} ===" | tee -a "${REPORT_FILE}"

    # FD stats from API
    FD_INFO=$(rmq_api "${node}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
for n in nodes:
    fd_used=n.get('fd_used',0)
    fd_total=n.get('fd_total',0)
    sock_used=n.get('sockets_used',0)
    sock_total=n.get('sockets_total',0)
    fd_pct=round(fd_used/fd_total*100,1) if fd_total>0 else 0
    print(f'{n[\"name\"]}: FD={fd_used}/{fd_total} ({fd_pct}%), Sockets={sock_used}/{sock_total}')
" 2>/dev/null || echo "API query failed")
    echo "${FD_INFO}" | tee -a "${REPORT_FILE}"

    # Top connections by channel count
    log_info "Top connections by channels..." | tee -a "${REPORT_FILE}"
    rmq_api "${node}" "connections?sort=channels&sort_reverse=true&page_size=10&columns=name,peer_host,channels,state" 2>/dev/null | \
        python3 -c "
import json,sys
try:
    data=json.load(sys.stdin)
    conns=data if isinstance(data,list) else data.get('items',[])
    for c in conns[:10]:
        print(f'  {c.get(\"peer_host\",\"?\")}: channels={c.get(\"channels\",0)}, state={c.get(\"state\",\"?\")}')
except: print('  Could not parse connections')
" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Connections per source IP
    log_info "Connections grouped by source IP..." | tee -a "${REPORT_FILE}"
    rmq_api "${node}" "connections?page_size=1000&columns=peer_host" 2>/dev/null | \
        python3 -c "
import json,sys
from collections import Counter
try:
    data=json.load(sys.stdin)
    conns=data if isinstance(data,list) else data.get('items',[])
    ip_counts=Counter(c.get('peer_host','?') for c in conns)
    for ip,count in ip_counts.most_common(10):
        print(f'  {ip}: {count} connections')
except: print('  Could not parse')
" 2>/dev/null | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If FD > 95%: Escalate to L3 for emergency connection cleanup. "
    "If leak identified: Escalate to Application team with source IP/client details."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 2 — QUEUE MONITORING RUNBOOK
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 2: Queue Monitoring Issues", level=1)
doc.add_paragraph("Dashboard: Queue Monitoring & Analytics")

add_issue_block(
    "RB-RMQ-006", "Queue Depth Critical (>50K messages)",
    "Dashboard #2 — Queue Monitoring & Analytics",
    "P1 - Critical",
    [
        "Queue depth widget shows one or more queues with > 50,000 messages",
        "Queue growth rate is positive and sustained",
        "Ready messages count is very high",
    ],
    [
        "Identify the affected queue(s) from Datadog top list widget",
        "Check consumer count for those queues: rabbitmqctl list_queues name messages consumers",
        "If consumers == 0: Check if consumer service is running, restart it",
        "If consumers > 0 but messages growing: Consumers are too slow",
        "Check consumer acknowledgment rate vs publish rate",
        "Check if consumers are redelivering (nack loop): rabbitmqctl list_queues name messages_unacknowledged",
        "Consider scaling consumers (add more instances)",
        "Check if there's a traffic spike causing the buildup",
        "As LAST RESORT with approval: Purge the queue: rabbitmqctl purge_queue <name>",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-006: Queue Depth Critical - Diagnosis & Resolution
# Usage: ./rb-rmq-006-queue-depth.sh [queue_name]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

QUEUE_NAME="${1:-}"
REPORT_FILE="/tmp/rmq-queue-depth-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

if [[ -z "${QUEUE_NAME}" ]]; then
    # Show top 20 queues by depth
    log_info "No queue specified. Showing top 20 deepest queues..." | tee -a "${REPORT_FILE}"
    rmq_api "${RMQ_NODE1}" "queues?sort=messages&sort_reverse=true&page_size=20&columns=name,vhost,messages,messages_ready,messages_unacknowledged,consumers,memory,state" 2>/dev/null | \
        python3 -c "
import json,sys
data=json.load(sys.stdin)
items=data if isinstance(data,list) else data.get('items',[])
print(f'{\"Queue\":<45} {\"Msgs\":>8} {\"Ready\":>8} {\"Unack\":>8} {\"Consumers\":>10} {\"Mem(MB)\":>8}')
print('-'*92)
for q in items[:20]:
    print(f'{q.get(\"name\",\"?\")[:43]:<45} {q.get(\"messages\",0):>8} {q.get(\"messages_ready\",0):>8} {q.get(\"messages_unacknowledged\",0):>8} {q.get(\"consumers\",0):>10} {q.get(\"memory\",0)//1048576:>8}')
" 2>/dev/null | tee -a "${REPORT_FILE}"
else
    # Detailed analysis of specific queue
    log_info "Analyzing queue: ${QUEUE_NAME}" | tee -a "${REPORT_FILE}"

    # Queue details
    QUEUE_DATA=$(rmq_api "${RMQ_NODE1}" "queues/%2F/${QUEUE_NAME}" 2>/dev/null)
    echo "${QUEUE_DATA}" | python3 -c "
import json,sys
q=json.load(sys.stdin)
print(f'Queue: {q.get(\"name\")}')
print(f'VHost: {q.get(\"vhost\")}')
print(f'Type: {q.get(\"type\",\"classic\")}')
print(f'State: {q.get(\"state\")}')
print(f'Messages: {q.get(\"messages\",0)}')
print(f'Ready: {q.get(\"messages_ready\",0)}')
print(f'Unacked: {q.get(\"messages_unacknowledged\",0)}')
print(f'Consumers: {q.get(\"consumers\",0)}')
print(f'Memory: {q.get(\"memory\",0)//1048576} MB')
print(f'Publish Rate: {q.get(\"message_stats\",{}).get(\"publish_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'Deliver Rate: {q.get(\"message_stats\",{}).get(\"deliver_get_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'Ack Rate: {q.get(\"message_stats\",{}).get(\"ack_details\",{}).get(\"rate\",0):.1f} msg/s')
" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Consumer details
    log_info "Consumer details..." | tee -a "${REPORT_FILE}"
    rmq_api "${RMQ_NODE1}" "queues/%2F/${QUEUE_NAME}/bindings" 2>/dev/null | \
        python3 -c "
import json,sys
bindings=json.load(sys.stdin)
for b in bindings:
    print(f'  Exchange: {b.get(\"source\",\"(default)\")}, Routing Key: {b.get(\"routing_key\",\"\")}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

    # Diagnosis
    CONSUMERS=$(echo "${QUEUE_DATA}" | python3 -c "import json,sys; print(json.load(sys.stdin).get('consumers',0))" 2>/dev/null || echo "0")
    MSGS=$(echo "${QUEUE_DATA}" | python3 -c "import json,sys; print(json.load(sys.stdin).get('messages',0))" 2>/dev/null || echo "0")

    echo "" | tee -a "${REPORT_FILE}"
    if [[ "${CONSUMERS}" == "0" ]]; then
        log_error "DIAGNOSIS: Queue has NO consumers. Messages will keep piling up." | tee -a "${REPORT_FILE}"
        echo "ACTION: Deploy/restart the consumer service for this queue." | tee -a "${REPORT_FILE}"
        notify_slack "Queue ${QUEUE_NAME}: ${MSGS} messages, 0 consumers. Consumer service needs restart." "critical"
    else
        log_warn "DIAGNOSIS: Queue has ${CONSUMERS} consumers but depth is still high." | tee -a "${REPORT_FILE}"
        echo "ACTION: Scale consumers or investigate slow processing." | tee -a "${REPORT_FILE}"
        notify_slack "Queue ${QUEUE_NAME}: ${MSGS} messages, ${CONSUMERS} consumers (not keeping up). Consider scaling." "warning"
    fi
fi

log_info "Report: ${REPORT_FILE}"
""",
    "If consumer service is down: Restart it (Application team). "
    "If consumers are slow: Escalate to Application team for optimization. "
    "If purge is needed: Get written approval from Application team lead."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 3 — MESSAGE FLOW RUNBOOK
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 3: Message Flow & Delivery Issues", level=1)
doc.add_paragraph("Dashboard: Message Flow & Delivery")

add_issue_block(
    "RB-RMQ-007", "Publish Rate Drop / Zero Throughput",
    "Dashboard #3 — Message Flow & Delivery",
    "P1 - Critical",
    [
        "Publish rate chart drops > 50% from baseline or goes to zero",
        "No new messages appearing in any queue",
        "Deliver rate may also drop (consequence of no publishes)",
    ],
    [
        "Check if ALL publishers stopped or just specific ones (look at per-exchange breakdown)",
        "Check for active memory/disk alarms: rabbitmqctl status | grep alarms (alarms BLOCK publishers)",
        "Check for flow control: rabbitmqctl list_connections name state | grep blocked/blocking",
        "If alarms active: Follow RB-RMQ-002 (memory) or RB-RMQ-003 (disk)",
        "If no alarms: Check if publisher applications are healthy (check app logs, process status)",
        "Check network: Can publisher hosts reach RabbitMQ? telnet <rmq-node> 5672",
        "Check for TLS certificate expiry if using TLS connections",
        "Check connection count: Did it drop? (publishers disconnected)",
        "Once identified, fix root cause and verify publish rate recovers in Datadog",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-007: Publish Rate Drop - Diagnosis
# Usage: ./rb-rmq-007-publish-drop.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-publish-drop-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Step 1: Check for alarms (most common cause)
log_info "Step 1: Checking for resource alarms (blocks publishers)..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
for n in nodes:
    name=n.get('name','?')
    mem_alarm=n.get('mem_alarm',False)
    disk_alarm=n.get('disk_free_alarm',False)
    if mem_alarm: print(f'  MEMORY ALARM ACTIVE on {name} — publishers are BLOCKED')
    if disk_alarm: print(f'  DISK ALARM ACTIVE on {name} — publishers are BLOCKED')
    if not mem_alarm and not disk_alarm: print(f'  {name}: no alarms')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 2: Check for blocked/flow-controlled connections
log_info "Step 2: Checking for blocked connections..." | tee -a "${REPORT_FILE}"
BLOCKED=$(rmq_api "${RMQ_NODE1}" "connections?page_size=100&columns=name,peer_host,state" 2>/dev/null | python3 -c "
import json,sys
data=json.load(sys.stdin)
conns=data if isinstance(data,list) else data.get('items',[])
blocked=[c for c in conns if c.get('state','') in ('blocked','blocking')]
print(f'{len(blocked)} blocked connections found')
for c in blocked[:10]:
    print(f'  {c.get(\"peer_host\",\"?\")}: state={c.get(\"state\")}')
" 2>/dev/null || echo "unknown")
echo "${BLOCKED}" | tee -a "${REPORT_FILE}"

# Step 3: Current message rates
log_info "Step 3: Current message rates..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "overview" 2>/dev/null | python3 -c "
import json,sys
d=json.load(sys.stdin)
ms=d.get('message_stats',{})
print(f'  Publish rate:  {ms.get(\"publish_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'  Deliver rate:  {ms.get(\"deliver_get_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'  Ack rate:      {ms.get(\"ack_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'  Redeliver rate:{ms.get(\"redeliver_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'  Connections:   {d.get(\"object_totals\",{}).get(\"connections\",0)}')
print(f'  Channels:      {d.get(\"object_totals\",{}).get(\"channels\",0)}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 4: Check connection count (did publishers disconnect?)
log_info "Step 4: Connection count per node..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    CONNS=$(rmq_api "${node}" "overview" 2>/dev/null | python3 -c "import json,sys; print(json.load(sys.stdin).get('object_totals',{}).get('connections',0))" 2>/dev/null || echo "?")
    log_info "  ${node}: ${CONNS} connections" | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
notify_slack "Publish rate drop investigation complete. Report: ${REPORT_FILE}" "warning"
""",
    "If alarms are the cause: Follow alarm runbooks (RB-RMQ-002/003). "
    "If publisher app is down: Escalate to Application team. "
    "If network issue: Escalate to Network team."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 4 — FAILED MESSAGES RUNBOOK
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 4: Failed / Dead-Lettered Messages", level=1)
doc.add_paragraph("Dashboard: Failed / Dead-Lettered Messages")

add_issue_block(
    "RB-RMQ-008", "Dead Letter Queue (DLQ) Depth Growing",
    "Dashboard #4 — Failed / Dead-Lettered Messages",
    "P2 - Warning (P1 if > 1000)",
    [
        "DLQ queues showing increasing message count",
        "DLQ growth rate is positive",
        "Redelivery rate may be elevated",
    ],
    [
        "Identify which DLQ(s) are growing from Datadog",
        "Determine the source queue: DLQ name usually indicates source (e.g., orders.dlq -> orders queue)",
        "Check the source queue for consumers and errors",
        "Check consumer logs for exceptions/errors causing rejections",
        "Inspect a sample DLQ message: Use Management UI or rabbitmqadmin to get message",
        "Check message headers for x-death (shows why it was dead-lettered: rejected, expired, maxlen)",
        "If rejected: Fix consumer error handling",
        "If expired (TTL): Consumers are too slow or down",
        "If maxlen: Queue overflow — increase max-length or scale consumers",
        "Process DLQ messages: Deploy DLQ processor or replay messages",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-008: DLQ Investigation
# Usage: ./rb-rmq-008-dlq-investigation.sh [dlq_queue_name]
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

DLQ_NAME="${1:-}"
REPORT_FILE="/tmp/rmq-dlq-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Find all DLQ queues
log_info "Finding all DLQ/dead-letter queues..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "queues?page_size=500&columns=name,messages,consumers,memory" 2>/dev/null | python3 -c "
import json,sys
data=json.load(sys.stdin)
items=data if isinstance(data,list) else data.get('items',[])
dlqs=[q for q in items if any(kw in q.get('name','').lower() for kw in ['dlq','dead','error','poison','failed','retry'])]
if dlqs:
    print(f'Found {len(dlqs)} DLQ-pattern queues:')
    print(f'{\"Queue\":<50} {\"Messages\":>10} {\"Consumers\":>10} {\"Mem(MB)\":>10}')
    print('-'*82)
    for q in sorted(dlqs, key=lambda x: x.get('messages',0), reverse=True):
        print(f'{q.get(\"name\",\"?\")[:48]:<50} {q.get(\"messages\",0):>10} {q.get(\"consumers\",0):>10} {q.get(\"memory\",0)//1048576:>10}')
else:
    print('No DLQ-pattern queues found.')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# If specific DLQ provided, get a sample message
if [[ -n "${DLQ_NAME}" ]]; then
    log_info "Getting sample message from: ${DLQ_NAME}..." | tee -a "${REPORT_FILE}"
    rmq_api "${RMQ_NODE1}" "queues/%2F/${DLQ_NAME}/get" \
        -X POST -H "content-type:application/json" \
        -d '{"count":1,"ackmode":"ack_requeue_true","encoding":"auto"}' 2>/dev/null | \
    python3 -c "
import json,sys
msgs=json.load(sys.stdin)
for m in msgs:
    props=m.get('properties',{})
    headers=props.get('headers',{})
    print('--- Sample Message ---')
    print(f'Routing Key: {m.get(\"routing_key\",\"?\")}')
    print(f'Exchange: {m.get(\"exchange\",\"?\")}')
    print(f'Redelivered: {m.get(\"redelivered\",\"?\")}')
    x_death=headers.get('x-death',[])
    if x_death:
        print(f'x-death (dead-letter history):')
        for d in x_death:
            print(f'  Queue: {d.get(\"queue\")}, Reason: {d.get(\"reason\")}, Count: {d.get(\"count\")}')
    print(f'Payload (first 500 chars): {m.get(\"payload\",\"\")[:500]}')
" 2>/dev/null | tee -a "${REPORT_FILE}"
fi

log_info "Report: ${REPORT_FILE}"
""",
    "If consumer code is rejecting messages: Escalate to Application team with sample message. "
    "If DLQ depth > 1000 and growing: Escalate to L3 and Application team lead."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 5 — CONSUMER/PUBLISHER PERFORMANCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 5: Consumer & Publisher Performance", level=1)

add_issue_block(
    "RB-RMQ-009", "Consumers Not Keeping Up / Consumer Lag",
    "Dashboard #5 — Consumer & Publisher Performance",
    "P2 - Warning",
    [
        "Consumer utilization < 50% while queue depth > 1000",
        "Ready messages growing despite consumers being connected",
        "Deliver rate << Publish rate",
    ],
    [
        "Check consumer count vs queue depth for affected queues",
        "Check prefetch count: Is it too low (1) or too high?",
        "Check consumer utilization metric: < 100% means consumer is idle sometimes",
        "Check unacked messages: High unacked with low prefetch = bottleneck",
        "Check if consumer is doing blocking I/O (database, API calls) in message handler",
        "Recommendation: Set prefetch to 10-50 (not 1, not unlimited)",
        "Scale consumers: Add more consumer instances",
        "If single slow message blocks others: Check for poison messages",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-009: Consumer Lag Analysis
# Usage: ./rb-rmq-009-consumer-lag.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-consumer-lag-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Queues where ready > 100 and consumers > 0 (lag = consumers can't keep up)
log_info "Queues with consumer lag (ready > 100, consumers > 0)..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "queues?sort=messages_ready&sort_reverse=true&page_size=50&columns=name,messages,messages_ready,messages_unacknowledged,consumers,consumer_utilisation,message_stats" 2>/dev/null | python3 -c "
import json,sys
data=json.load(sys.stdin)
items=data if isinstance(data,list) else data.get('items',[])
lagging=[q for q in items if q.get('consumers',0)>0 and q.get('messages_ready',0)>100]
if lagging:
    print(f'{\"Queue\":<40} {\"Ready\":>8} {\"Unack\":>8} {\"Consumers\":>10} {\"Util%\":>8} {\"Deliver/s\":>10}')
    print('-'*88)
    for q in lagging[:20]:
        util=q.get('consumer_utilisation')
        util_str=f'{util*100:.0f}%' if util is not None else 'N/A'
        deliver=q.get('message_stats',{}).get('deliver_get_details',{}).get('rate',0)
        print(f'{q.get(\"name\",\"?\")[:38]:<40} {q.get(\"messages_ready\",0):>8} {q.get(\"messages_unacknowledged\",0):>8} {q.get(\"consumers\",0):>10} {util_str:>8} {deliver:>10.1f}')
else:
    print('No queues with significant consumer lag found.')
" 2>/dev/null | tee -a "${REPORT_FILE}"

log_info "Recommendations:" | tee -a "${REPORT_FILE}"
echo "  - If Util% < 100%: Increase prefetch_count (optimal: 10-50)" | tee -a "${REPORT_FILE}"
echo "  - If Util% = 100% and lag growing: Add more consumer instances" | tee -a "${REPORT_FILE}"
echo "  - If Unack is high: Consumer processing is slow — optimize app code" | tee -a "${REPORT_FILE}"
log_info "Report: ${REPORT_FILE}"
""",
    "If consumers are at capacity: Escalate to Application team to scale consumer deployment. "
    "If prefetch needs tuning: Application team to adjust client configuration."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 6/7 — CONNECTIONS & RESOURCES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 6: Connection & Resource Issues", level=1)

add_issue_block(
    "RB-RMQ-010", "Connection Storm / Connection Leak",
    "Dashboard #7 — Connection & Channel Monitoring",
    "P1 - Critical (storm) / P2 - Warning (leak)",
    [
        "Connection count rapidly increasing (storm: > 100/sec)",
        "Connection count monotonically increasing over hours (leak)",
        "FD usage rising in correlation",
        "Possible 'blocked' connections appearing",
    ],
    [
        "Check connection rate: Is it a sudden spike (storm) or gradual increase (leak)?",
        "For STORM: Identify source IP(s): rabbitmqctl list_connections peer_host | sort | uniq -c | sort -rn",
        "For STORM: Check if a service is in a reconnect loop (crash -> connect -> crash)",
        "For LEAK: Identify which application is not closing connections properly",
        "Check connection ages: Old connections with no activity = leak",
        "If identified: Notify application team to fix client code",
        "Emergency: Close problematic connections: rabbitmqctl close_connection <pid> 'storm'",
        "Consider: Set connection rate limit in rabbitmq.conf: connection_max = 10000",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-010: Connection Storm / Leak Analysis
# Usage: ./rb-rmq-010-connection-analysis.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-connections-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Connection summary per node
log_info "Connection count per node..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    CONNS=$(rmq_api "${node}" "overview" 2>/dev/null | python3 -c "import json,sys; print(json.load(sys.stdin).get('object_totals',{}).get('connections',0))" 2>/dev/null || echo "?")
    CHANS=$(rmq_api "${node}" "overview" 2>/dev/null | python3 -c "import json,sys; print(json.load(sys.stdin).get('object_totals',{}).get('channels',0))" 2>/dev/null || echo "?")
    log_info "  ${node}: ${CONNS} connections, ${CHANS} channels" | tee -a "${REPORT_FILE}"
done

# Connections by source IP (top 20)
log_info "Connections by source IP (top 20)..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "connections?page_size=5000&columns=peer_host,name,channels,state,connected_at" 2>/dev/null | python3 -c "
import json,sys
from collections import Counter
data=json.load(sys.stdin)
conns=data if isinstance(data,list) else data.get('items',[])
ip_counts=Counter(c.get('peer_host','?') for c in conns)
print(f'{\"Source IP\":<25} {\"Connections\":>12}')
print('-'*38)
for ip,count in ip_counts.most_common(20):
    print(f'{ip:<25} {count:>12}')
print(f'{\"\":<25} {\"\":-^12}')
print(f'{\"TOTAL\":<25} {len(conns):>12}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Oldest connections (potential leaks)
log_info "Oldest 20 connections (potential leaks)..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "connections?sort=connected_at&page_size=20&columns=peer_host,name,channels,state,connected_at" 2>/dev/null | python3 -c "
import json,sys
from datetime import datetime
data=json.load(sys.stdin)
conns=data if isinstance(data,list) else data.get('items',[])
print(f'{\"Source IP\":<20} {\"Channels\":>8} {\"State\":<12} {\"Connected At\":<25}')
print('-'*70)
for c in conns[:20]:
    ts=c.get('connected_at',0)
    dt=datetime.fromtimestamp(ts/1000).strftime('%Y-%m-%d %H:%M') if ts else '?'
    print(f'{c.get(\"peer_host\",\"?\"):<20} {c.get(\"channels\",0):>8} {c.get(\"state\",\"?\"):<12} {dt:<25}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Channels per connection (top offenders)
log_info "Top connections by channel count..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "connections?sort=channels&sort_reverse=true&page_size=10&columns=peer_host,name,channels" 2>/dev/null | python3 -c "
import json,sys
data=json.load(sys.stdin)
conns=data if isinstance(data,list) else data.get('items',[])
for c in conns[:10]:
    print(f'  {c.get(\"peer_host\",\"?\")}: {c.get(\"channels\",0)} channels')
" 2>/dev/null | tee -a "${REPORT_FILE}"

log_info "Report: ${REPORT_FILE}"
""",
    "If storm from known service: Escalate to that Application team immediately. "
    "If leak source identified: Create ticket for Application team to fix client code."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 8 — USER & PERMISSION AUDIT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 7: User & Permission Issues", level=1)

add_issue_block(
    "RB-RMQ-011", "Authentication Failures / Unauthorized Access",
    "Dashboard #8 — User & Permission Audit",
    "P1 - Critical (if spike) / P3 - Info (occasional)",
    [
        "Authentication failure count spiking in logs",
        "Failed auth attempts from unknown IP addresses",
        "New admin user created without change request",
    ],
    [
        "Check auth failure logs: grep 'authentication' /var/log/rabbitmq/*.log | tail -50",
        "Identify source IPs of failed attempts",
        "Check if it's a legitimate app with wrong credentials vs potential attack",
        "If credential rotation in progress: Coordinate with application teams",
        "If unauthorized: Block source IP at firewall/security group level",
        "Audit current users: rabbitmqctl list_users",
        "Audit permissions: rabbitmqctl list_permissions -p /",
        "Check if guest account is enabled for non-localhost: SECURITY RISK",
        "If new admin found: Verify with change management, revoke if unauthorized",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-011: Security Audit & Auth Failure Investigation
# Usage: ./rb-rmq-011-security-audit.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-security-audit-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

# Step 1: List all users and roles
log_info "Step 1: All RabbitMQ users..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    if check_node_reachable "${node}"; then
        remote_exec "${node}" "sudo rabbitmqctl list_users" 2>/dev/null | tee -a "${REPORT_FILE}"
        break  # Only need from one node
    fi
done

# Step 2: List permissions per vhost
log_info "Step 2: Permissions per vhost..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "permissions" 2>/dev/null | python3 -c "
import json,sys
perms=json.load(sys.stdin)
print(f'{\"User\":<25} {\"VHost\":<15} {\"Configure\":<15} {\"Write\":<15} {\"Read\":<15}')
print('-'*87)
for p in perms:
    print(f'{p.get(\"user\",\"?\"):<25} {p.get(\"vhost\",\"?\"):<15} {p.get(\"configure\",\"?\"):<15} {p.get(\"write\",\"?\"):<15} {p.get(\"read\",\"?\"):<15}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 3: Check for admin users
log_info "Step 3: Admin users..." | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "users" 2>/dev/null | python3 -c "
import json,sys
users=json.load(sys.stdin)
admins=[u for u in users if 'administrator' in u.get('tags','')]
print(f'Admin users ({len(admins)}):')
for a in admins:
    print(f'  {a[\"name\"]} (tags: {a.get(\"tags\",\"\")})')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# Step 4: Check guest account
log_info "Step 4: Guest account status..." | tee -a "${REPORT_FILE}"
GUEST_EXISTS=$(rmq_api "${RMQ_NODE1}" "users/guest" 2>/dev/null)
if echo "${GUEST_EXISTS}" | grep -q "guest"; then
    log_warn "SECURITY: guest account EXISTS. Verify it's restricted to localhost only." | tee -a "${REPORT_FILE}"
else
    log_ok "guest account does not exist or is disabled." | tee -a "${REPORT_FILE}"
fi

# Step 5: Recent auth failures from logs
log_info "Step 5: Recent authentication failures (last 100 lines)..." | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    log_info "--- ${node} ---" | tee -a "${REPORT_FILE}"
    remote_exec "${node}" "sudo grep -i 'authentication\\|access_refused\\|login\\|AUTH' ${RMQ_LOG_DIR}/rabbit@*.log 2>/dev/null | tail -20" 2>/dev/null | tee -a "${REPORT_FILE}"
done

log_info "Report: ${REPORT_FILE}"
""",
    "If potential security incident: Escalate to Security team immediately. "
    "If unauthorized admin user found: Revoke access and escalate to Security + Management."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  DASHBOARD 9 — RESTART & AVAILABILITY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Runbook 8: Restart & Availability Issues", level=1)

add_issue_block(
    "RB-RMQ-012", "Full Cluster Restart / Recovery",
    "Dashboard #9 — Restart & Availability Tracking",
    "P1 - Critical",
    [
        "All 3 nodes show recent restart (uptime < 1 hour for all)",
        "Or: All nodes are down simultaneously",
        "Or: Cluster won't form after restarts",
    ],
    [
        "CRITICAL: Identify which node stopped LAST — it must start FIRST",
        "If unknown: Check logs on each node for last shutdown timestamp",
        "Start the last-stopped node first: sudo systemctl start rabbitmq-server",
        "If it won't start: Use force_boot: sudo rabbitmqctl force_boot && sudo systemctl start rabbitmq-server",
        "Wait for first node to be fully running: sudo rabbitmqctl status",
        "Start second node: It should automatically join the cluster",
        "Start third node: It should also join",
        "Verify cluster: sudo rabbitmqctl cluster_status — should show all 3",
        "Verify queues: sudo rabbitmqctl list_queues name messages consumers leader online",
        "Check for data loss: Compare message counts with pre-restart values if available",
        "Run health check: Publish and consume a test message",
    ],
    """#!/bin/bash
# ============================================================
# RB-RMQ-012: Full Cluster Recovery
# Usage: ./rb-rmq-012-cluster-recovery.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-cluster-recovery-$(date +%Y%m%d-%H%M%S).log"
report_header | tee "${REPORT_FILE}"

log_warn "=== FULL CLUSTER RECOVERY PROCEDURE ===" | tee -a "${REPORT_FILE}"
notify_slack "CLUSTER RECOVERY initiated for ${CLUSTER_NAME}" "critical"

# Step 1: Check current state of all nodes
log_info "Step 1: Checking state of all nodes..." | tee -a "${REPORT_FILE}"
LAST_NODE=""
LAST_TIME=0

for node in ${RMQ_NODES}; do
    log_info "--- ${node} ---" | tee -a "${REPORT_FILE}"

    if ! check_node_reachable "${node}"; then
        log_error "  HOST UNREACHABLE: ${node}" | tee -a "${REPORT_FILE}"
        continue
    fi

    # Check if RabbitMQ is running
    RMQ_RUNNING=$(remote_exec "${node}" "pgrep -f beam.smp" 2>/dev/null && echo "yes" || echo "no")
    log_info "  RabbitMQ running: ${RMQ_RUNNING}" | tee -a "${REPORT_FILE}"

    # Check last shutdown time from logs
    SHUTDOWN_TIME=$(remote_exec "${node}" "sudo grep -i 'shutdown\\|stopped\\|stop_app' ${RMQ_LOG_DIR}/rabbit@*.log 2>/dev/null | tail -1" || echo "unknown")
    log_info "  Last shutdown: ${SHUTDOWN_TIME}" | tee -a "${REPORT_FILE}"
done

# Step 2: Start first node (with force_boot if needed)
log_info "" | tee -a "${REPORT_FILE}"
log_info "Step 2: Starting cluster..." | tee -a "${REPORT_FILE}"
log_warn "The node that stopped LAST should start FIRST." | tee -a "${REPORT_FILE}"
log_warn "If unsure, start Node1 with force_boot:" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # On the first node to start:" | tee -a "${REPORT_FILE}"
echo "  ssh ${SSH_USER}@${RMQ_NODE1}" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl force_boot" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start rabbitmq-server" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl status  # Wait for this to succeed" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # Then on each remaining node:" | tee -a "${REPORT_FILE}"
echo "  ssh ${SSH_USER}@${RMQ_NODE2}" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start rabbitmq-server" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  ssh ${SSH_USER}@${RMQ_NODE3}" | tee -a "${REPORT_FILE}"
echo "  sudo systemctl start rabbitmq-server" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # Verify cluster:" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl cluster_status" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl list_queues name messages consumers leader online" | tee -a "${REPORT_FILE}"

# Step 3: Post-recovery health check
log_info "Step 3: After cluster is up, run this health check:" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # Verify all 3 nodes in cluster:" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl cluster_status | grep -A5 'running_nodes'" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # Check for alarms:" | tee -a "${REPORT_FILE}"
echo "  sudo rabbitmqctl status | grep -A5 'alarms'" | tee -a "${REPORT_FILE}"
echo "" | tee -a "${REPORT_FILE}"
echo "  # Test publish/consume:" | tee -a "${REPORT_FILE}"
echo "  rabbitmqadmin publish routing_key=test payload='healthcheck' exchange=amq.default" | tee -a "${REPORT_FILE}"
echo "  rabbitmqadmin get queue=test count=1" | tee -a "${REPORT_FILE}"

log_info "Report: ${REPORT_FILE}"
""",
    "If force_boot fails: Escalate to L3/Principal Engineer. "
    "If data loss detected: Escalate to L3 and Application team. "
    "Create post-incident review."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  COMPREHENSIVE HEALTH CHECK SCRIPT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Comprehensive Health Check Script", level=1)
doc.add_paragraph(
    "This script performs a full health check of the RabbitMQ cluster. "
    "Run this periodically or before/after any maintenance activity."
)

add_code("""#!/bin/bash
# ============================================================
# RabbitMQ Comprehensive Health Check
# Usage: ./rmq-health-check.sh
# ============================================================
source "$(dirname "$0")/../env/${ENVIRONMENT:-production}.env"
source "$(dirname "$0")/../lib/common.sh"
validate_env

REPORT_FILE="/tmp/rmq-health-check-$(date +%Y%m%d-%H%M%S).log"
EXIT_CODE=0

report_header | tee "${REPORT_FILE}"

# ---- Node Health ----
log_info "=== NODE HEALTH ===" | tee -a "${REPORT_FILE}"
for node in ${RMQ_NODES}; do
    if ! check_node_reachable "${node}"; then
        log_error "FAIL: ${node} is unreachable" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
        continue
    fi

    RMQ_PID=$(remote_exec "${node}" "pgrep -f beam.smp" 2>/dev/null || echo "")
    if [[ -z "${RMQ_PID}" ]]; then
        log_error "FAIL: ${node} — RabbitMQ not running" | tee -a "${REPORT_FILE}"
        EXIT_CODE=1
    else
        log_ok "PASS: ${node} — RabbitMQ running (PID: ${RMQ_PID})" | tee -a "${REPORT_FILE}"
    fi
done

# ---- Cluster Status ----
log_info "=== CLUSTER STATUS ===" | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "nodes" 2>/dev/null | python3 -c "
import json,sys
nodes=json.load(sys.stdin)
for n in nodes:
    name=n.get('name','?')
    running=n.get('running',False)
    mem_alarm=n.get('mem_alarm',False)
    disk_alarm=n.get('disk_free_alarm',False)
    partitions=n.get('partitions',[])
    fd_used=n.get('fd_used',0)
    fd_total=n.get('fd_total',1)
    mem_used=n.get('mem_used',0)
    mem_limit=n.get('mem_limit',1)
    uptime=n.get('uptime',0)//1000

    status='OK'
    if not running: status='CRITICAL: Not running'
    if mem_alarm: status='CRITICAL: Memory alarm'
    if disk_alarm: status='CRITICAL: Disk alarm'
    if partitions: status=f'CRITICAL: Partition detected: {partitions}'

    print(f'{name}:')
    print(f'  Status: {status}')
    print(f'  Memory: {mem_used//1048576}MB / {mem_limit//1048576}MB ({round(mem_used/mem_limit*100,1)}%)')
    print(f'  FDs: {fd_used}/{fd_total} ({round(fd_used/fd_total*100,1)}%)')
    print(f'  Uptime: {uptime//3600}h {(uptime%3600)//60}m')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# ---- Queue Health ----
log_info "=== QUEUE HEALTH ===" | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "queues?page_size=500&columns=name,messages,consumers" 2>/dev/null | python3 -c "
import json,sys
data=json.load(sys.stdin)
items=data if isinstance(data,list) else data.get('items',[])
total_queues=len(items)
deep_queues=[q for q in items if q.get('messages',0)>10000]
no_consumer_queues=[q for q in items if q.get('consumers',0)==0 and q.get('messages',0)>0]
print(f'Total queues: {total_queues}')
print(f'Queues > 10K messages: {len(deep_queues)}')
for q in deep_queues:
    print(f'  {q[\"name\"]}: {q[\"messages\"]} msgs')
print(f'Queues with messages but 0 consumers: {len(no_consumer_queues)}')
for q in no_consumer_queues[:5]:
    print(f'  {q[\"name\"]}: {q[\"messages\"]} msgs')
" 2>/dev/null | tee -a "${REPORT_FILE}"

# ---- Message Rates ----
log_info "=== MESSAGE RATES ===" | tee -a "${REPORT_FILE}"
rmq_api "${RMQ_NODE1}" "overview" 2>/dev/null | python3 -c "
import json,sys
d=json.load(sys.stdin)
ms=d.get('message_stats',{})
print(f'Publish:  {ms.get(\"publish_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'Deliver:  {ms.get(\"deliver_get_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'Ack:      {ms.get(\"ack_details\",{}).get(\"rate\",0):.1f} msg/s')
print(f'Redeliver:{ms.get(\"redeliver_details\",{}).get(\"rate\",0):.1f} msg/s')
ot=d.get('object_totals',{})
print(f'Connections: {ot.get(\"connections\",0)}')
print(f'Channels:    {ot.get(\"channels\",0)}')
print(f'Consumers:   {ot.get(\"consumers\",0)}')
" 2>/dev/null | tee -a "${REPORT_FILE}"

log_info "Health check complete. Report: ${REPORT_FILE}"
exit ${EXIT_CODE}
""")

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RabbitMQ Operations Runbook — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output = "/Users/tejasodanapalli/rabbitmq/RabbitMQ/docs/monitoring/RabbitMQ_Operations_Runbook.docx"
doc.save(output)
print(f"Document saved: {output}")
