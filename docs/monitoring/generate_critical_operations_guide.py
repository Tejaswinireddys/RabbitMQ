#!/usr/bin/env python3
"""
Critical Operations Guide Generator
=====================================
Generates a comprehensive Word document with architect-level operational notes,
OS patching playbooks, and critical warnings for RabbitMQ and Redis clusters.

Output: Critical_Operations_Guide_RabbitMQ_Redis.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "Critical_Operations_Guide_RabbitMQ_Redis.docx")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "2C3E50")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F3F4")

    return table


def add_warning_box(doc, text, severity="CRITICAL"):
    """Add a visually distinct warning paragraph."""
    colors = {
        "CRITICAL": RGBColor(192, 0, 0),
        "WARNING": RGBColor(196, 120, 0),
        "IMPORTANT": RGBColor(0, 100, 180),
        "NOTE": RGBColor(0, 128, 0),
    }
    p = doc.add_paragraph()
    prefix = p.add_run(f"  {severity}: ")
    prefix.bold = True
    prefix.font.color.rgb = colors.get(severity, RGBColor(192, 0, 0))
    prefix.font.size = Pt(10)
    body = p.add_run(text)
    body.font.size = Pt(10)
    body.font.color.rgb = colors.get(severity, RGBColor(192, 0, 0))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc, code):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_numbered_steps(doc, steps):
    """Add numbered steps as a list."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}", style='List Number')
        for r in p.runs:
            r.font.size = Pt(10)


def add_bullet_list(doc, items):
    """Add bullet points."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for r in p.runs:
            r.font.size = Pt(10)


def build_title_page(doc):
    """Build title page."""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CRITICAL OPERATIONS GUIDE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(192, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("RabbitMQ & Redis HA Clusters")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Architect-Level Operational Notes, OS Patching Playbooks,\n"
        "Quorum Rules, and Critical Warnings\n\n"
        "FOR L1/L2 ENGINEERS AND INFRASTRUCTURE TEAMS"
    )
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warn.add_run(
        "READ THIS DOCUMENT BEFORE performing any maintenance,\n"
        "patching, or restart operations on these clusters."
    )
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def build_toc(doc):
    """Build table of contents page."""
    doc.add_heading('Table of Contents', level=1)

    sections = [
        "Section A: RabbitMQ Critical Operational Notes",
        "  A.1  pause_minority Quorum Rules",
        "  A.2  Rolling Restart Procedure",
        "  A.3  Erlang Cookie — Cluster Identity",
        "  A.4  Mnesia Database — Never Delete",
        "  A.5  Quorum Queue Majority Rules",
        "  A.6  Memory & Disk Alarms — Cluster-Wide Impact",
        "  A.7  Network Partition Handling",
        "  A.8  Feature Flags — One-Way Gate",
        "  A.9  Port & Firewall Requirements",
        "  A.10 What NOT To Do — RabbitMQ",
        "",
        "Section B: Redis + Sentinel Critical Operational Notes",
        "  B.1  Sentinel Quorum (2 of 3)",
        "  B.2  Never Restart Master First",
        "  B.3  min-replicas-to-write Constraint",
        "  B.4  Sentinel Config Auto-Rewrite",
        "  B.5  RDB Fork Memory Requirement",
        "  B.6  AOF + RDB Persistence Interaction",
        "  B.7  Renamed/Disabled Commands",
        "  B.8  Replication Backlog & Full Resync",
        "  B.9  What NOT To Do — Redis",
        "",
        "Section C: OS Patching Playbook",
        "  C.1  Pre-Patch Checklist",
        "  C.2  RabbitMQ Patching Sequence",
        "  C.3  Redis Patching Sequence",
        "  C.4  Post-Patch Validation",
        "  C.5  Rollback Procedures",
        "",
        "Section D: Emergency Quick Reference",
        "  D.1  DO and DO NOT — RabbitMQ",
        "  D.2  DO and DO NOT — Redis",
        "  D.3  Decision Tree: Can I Take This Node Down?",
        "",
        "Section E: Directory Structure & Variable Reference",
        "  E.1  RabbitMQ Directory Layout (RPM / Package)",
        "  E.2  Redis Directory Layout (RPM / Extracted)",
        "  E.3  Environment Configuration Variables",
    ]

    for s in sections:
        if s == "":
            doc.add_paragraph()
        elif s.startswith("Section"):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(s)
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_page_break()


def build_section_a(doc):
    """RabbitMQ Critical Operational Notes."""
    doc.add_heading('Section A: RabbitMQ Critical Operational Notes', level=1)

    # A.1 pause_minority
    doc.add_heading('A.1  pause_minority Quorum Rules', level=2)
    doc.add_paragraph(
        'Your cluster uses cluster_partition_handling = pause_minority. '
        'This is the safest partition handling mode, but it has strict operational rules:'
    )

    add_warning_box(doc,
        "In a 3-node cluster with pause_minority, at least 2 nodes MUST be running at all times. "
        "If 2 nodes go down simultaneously, the remaining single node will PAUSE ITSELF — "
        "it stops serving ALL operations: no publishes, no consumes, no management API, no healthchecks. "
        "The cluster becomes completely unavailable.",
        "CRITICAL"
    )

    doc.add_paragraph("How pause_minority works:")
    add_bullet_list(doc, [
        "When a network partition occurs, each side counts how many nodes it can see",
        "The side with FEWER nodes (minority) pauses itself automatically",
        "In a 3-node cluster: 2 nodes = majority (keeps running), 1 node = minority (pauses)",
        "If the cluster splits 1-1-1 (all isolated): ALL nodes pause — total outage",
        "Paused nodes stop accepting ANY client connections",
        "When connectivity restores, paused nodes automatically resume",
    ])

    add_styled_table(doc,
        ["Scenario", "Nodes Up", "Cluster Status", "Impact"],
        [
            ["Normal", "3 of 3", "HEALTHY", "Full capacity"],
            ["1 node down (planned)", "2 of 3", "DEGRADED but FUNCTIONAL", "Quorum maintained, slight capacity reduction"],
            ["2 nodes down", "1 of 3", "COMPLETELY DOWN — node self-pauses", "Total outage, no reads/writes/API"],
            ["Network split (2+1)", "2+1", "Majority side runs, minority pauses", "Clients on minority side lose connection"],
            ["Network split (1+1+1)", "1+1+1", "ALL nodes pause", "Total outage"],
        ]
    )

    add_warning_box(doc,
        "NEVER restart or patch 2 nodes at the same time. Always wait for the first node to fully rejoin "
        "the cluster and synchronize before touching the next node.",
        "CRITICAL"
    )

    # A.2 Rolling Restart
    doc.add_heading('A.2  Rolling Restart Procedure', level=2)
    doc.add_paragraph(
        'Rolling restarts must follow a strict sequence to maintain quorum at every step.'
    )

    add_numbered_steps(doc, [
        "BEFORE starting: Verify all 3 nodes are healthy\n"
        "   sudo ${RABBITMQCTL} cluster_status\n"
        "   Confirm: 3 running nodes, 0 partitions, no alarms",

        "Drain Node 3 (last joined / least critical):\n"
        "   sudo ${RABBITMQ_UPGRADE} drain\n"
        "   Wait for existing connections to close (check management UI)",

        "Stop Node 3:\n"
        "   sudo systemctl stop ${RMQ_SERVICE}\n"
        "   Verify remaining 2 nodes are still healthy",

        "Perform maintenance on Node 3 (patch, upgrade, etc.)",

        "Start Node 3:\n"
        "   sudo systemctl start ${RMQ_SERVICE}",

        "Revive Node 3:\n"
        "   sudo ${RABBITMQ_UPGRADE} revive",

        "Wait for Node 3 to fully sync:\n"
        "   sudo ${RABBITMQCTL} await_online_quorum_plus_one\n"
        "   This ensures quorum queues have replicated to the restarted node",

        "VERIFY before moving to next node:\n"
        "   sudo ${RABBITMQCTL} cluster_status\n"
        "   Confirm: 3 running nodes, 0 partitions, no alarms\n"
        "   Check queue sync: sudo ${RABBITMQ_DIAGNOSTICS} check_if_node_is_quorum_critical",

        "Repeat steps 2-8 for Node 2",

        "Repeat steps 2-8 for Node 1",
    ])

    add_warning_box(doc,
        "ALWAYS run 'rabbitmqctl await_online_quorum_plus_one' after bringing a node back up. "
        "This blocks until quorum queues are fully replicated. Moving to the next node before "
        "this completes can result in DATA LOSS.",
        "CRITICAL"
    )

    # A.3 Erlang Cookie
    doc.add_heading('A.3  Erlang Cookie — Cluster Identity', level=2)
    doc.add_paragraph(
        'The Erlang cookie is how RabbitMQ nodes authenticate each other for clustering.'
    )

    add_warning_box(doc,
        "ALL nodes in the cluster MUST have the IDENTICAL Erlang cookie. "
        "If even one character differs, that node cannot join the cluster. "
        "Location: ${RMQ_COOKIE_FILE} (default: /var/lib/rabbitmq/.erlang.cookie)",
        "CRITICAL"
    )

    doc.add_paragraph("Rules:")
    add_bullet_list(doc, [
        "File must be owned by the rabbitmq user: chown ${RMQ_RUN_USER}:${RMQ_RUN_USER} ${RMQ_COOKIE_FILE}",
        "File must have permissions 400: chmod 400 ${RMQ_COOKIE_FILE}",
        "NEVER change the cookie on a running cluster",
        "If you must reset the cookie: stop ALL nodes, change cookie on ALL nodes, start one node first, then others",
        "The cookie is a plain text string — treat it as a secret (do not log or expose)",
        "When deploying new nodes, copy the cookie FROM an existing cluster member",
    ])

    add_code_block(doc,
        "# Verify cookie consistency across all nodes:\n"
        "for node in ${RMQ_NODES[@]}; do\n"
        "    echo -n \"${node}: \"\n"
        "    ssh ${node} 'sudo cat ${RMQ_COOKIE_FILE}'\n"
        "done"
    )

    # A.4 Mnesia Database
    doc.add_heading('A.4  Mnesia Database — Never Delete', level=2)

    add_warning_box(doc,
        "NEVER delete the Mnesia database directory (${RMQ_MNESIA_DIR}) unless you are intentionally "
        "performing a full cluster reset. Mnesia contains ALL cluster metadata: queue definitions, "
        "exchange bindings, user permissions, vhost configs, and quorum queue Raft state.",
        "CRITICAL"
    )

    doc.add_paragraph("What Mnesia contains:")
    add_bullet_list(doc, [
        "Queue declarations and bindings",
        "Exchange definitions",
        "User accounts and permissions",
        "Virtual host configurations",
        "Policy definitions",
        "Quorum queue Raft logs and snapshots (the actual message data)",
        "Stream segment indices",
    ])

    doc.add_paragraph("If Mnesia is corrupted or deleted:")
    add_bullet_list(doc, [
        "The node will start as a blank node with no queue/exchange metadata",
        "It will NOT automatically rejoin the cluster",
        "Quorum queue data on that node is PERMANENTLY LOST",
        "You must manually reset and re-add the node to the cluster",
        "Messages in quorum queues on that node must be re-replicated from other nodes",
    ])

    # A.5 Quorum Queue Rules
    doc.add_heading('A.5  Quorum Queue Majority Rules', level=2)
    doc.add_paragraph(
        'Quorum queues use the Raft consensus protocol. They require a majority of members '
        'to be online for write operations.'
    )

    add_styled_table(doc,
        ["Cluster Size", "Quorum Needed", "Max Nodes Down", "Behavior When Below Quorum"],
        [
            ["3 nodes", "2", "1", "Queue becomes unavailable for publishes and acknowledges"],
            ["5 nodes", "3", "2", "Same — no writes below quorum"],
            ["7 nodes", "4", "3", "Same pattern"],
        ]
    )

    add_warning_box(doc,
        "With 3 nodes: if 2 nodes are down, ALL quorum queues become UNAVAILABLE. "
        "No messages can be published or acknowledged. Consumers may still read uncommitted messages "
        "but new messages cannot be enqueued.",
        "CRITICAL"
    )

    doc.add_paragraph("Key quorum queue behaviors:")
    add_bullet_list(doc, [
        "Quorum queues replicate across ALL cluster nodes by default (replication factor = cluster size)",
        "Leader election happens via Raft — one node leads, others follow",
        "If the leader node goes down, a new leader is elected (automatic, takes seconds)",
        "After a node restart, it must catch up on missed Raft log entries before serving",
        "Use 'rabbitmq-diagnostics check_if_node_is_quorum_critical' before stopping a node — "
        "it tells you if that node is the LAST member of any quorum queue",
    ])

    add_code_block(doc,
        "# Check if a node is critical for any quorum queue:\n"
        "sudo ${RABBITMQ_DIAGNOSTICS} check_if_node_is_quorum_critical\n"
        "# If this returns a warning, DO NOT stop this node until others catch up"
    )

    # A.6 Memory & Disk Alarms
    doc.add_heading('A.6  Memory & Disk Alarms — Cluster-Wide Impact', level=2)

    add_warning_box(doc,
        "Memory and disk alarms are CLUSTER-WIDE. When ONE node triggers an alarm, "
        "ALL publishers across the ENTIRE cluster are blocked — not just publishers connected "
        "to the alarmed node.",
        "CRITICAL"
    )

    doc.add_paragraph("Memory alarm (vm_memory_high_watermark):")
    add_bullet_list(doc, [
        "Triggers when a node's memory usage exceeds the configured watermark (default 60% of RAM)",
        "ALL publishers on ALL nodes are blocked immediately",
        "Consumers continue to work (this is by design — to drain messages and free memory)",
        "Management API remains available for diagnostics",
        "The alarm clears automatically when memory drops below the threshold",
        "Common causes: message backlog in queues, too many connections, large messages in memory",
    ])

    doc.add_paragraph("Disk alarm (disk_free_limit):")
    add_bullet_list(doc, [
        "Triggers when free disk drops below the configured limit (your config: 5GB)",
        "Same behavior: ALL publishers blocked cluster-wide",
        "Consumers continue working",
        "Common causes: Mnesia growth, log files, RDB/checkpoint files, message store",
    ])

    # A.7 Network Partition
    doc.add_heading('A.7  Network Partition Handling', level=2)
    doc.add_paragraph(
        'With pause_minority, network partitions are handled automatically but you must '
        'understand what happens:'
    )

    add_styled_table(doc,
        ["Event", "What Happens", "Your Action"],
        [
            ["Node loses network to others",
             "If it's minority: self-pauses\nIf it's majority side: continues serving",
             "Fix network. Minority node auto-resumes when connectivity restores."],
            ["Switch/router failure isolates 1 node",
             "1 node pauses (minority), 2 nodes continue (majority)",
             "Fix network path. No data loss if fixed promptly."],
            ["Full network partition (all isolated)",
             "ALL 3 nodes pause (each sees itself as minority of 1)",
             "URGENT: Fix network immediately. Total outage until resolved."],
            ["Node comes back after pause",
             "Auto-resumes, re-syncs with cluster",
             "Monitor sync progress. Verify no stuck queues."],
        ]
    )

    add_warning_box(doc,
        "After a partition heals, quorum queues re-synchronize automatically. Classic mirrored queues "
        "(if any exist) may require manual intervention. Always check 'rabbitmqctl cluster_status' "
        "after a partition event.",
        "IMPORTANT"
    )

    # A.8 Feature Flags
    doc.add_heading('A.8  Feature Flags — One-Way Gate', level=2)

    add_warning_box(doc,
        "Feature flags in RabbitMQ are ONE-WAY. Once enabled, they CANNOT be disabled. "
        "Enabling a feature flag permanently blocks downgrading to a version that doesn't support it.",
        "CRITICAL"
    )

    doc.add_paragraph("Rules:")
    add_bullet_list(doc, [
        "Before enabling any feature flag: confirm you will NEVER need to downgrade",
        "RabbitMQ 4.x enables many feature flags by default on fresh install",
        "When upgrading from 3.x to 4.x: feature flags are NOT auto-enabled — you must enable them",
        "Some features (e.g., quorum queues, streams) require specific feature flags",
        "List current flags: sudo ${RABBITMQCTL} list_feature_flags",
        "Enable a flag: sudo ${RABBITMQCTL} enable_feature_flag <flag_name>",
    ])

    # A.9 Port Requirements
    doc.add_heading('A.9  Port & Firewall Requirements', level=2)
    doc.add_paragraph(
        'ALL these ports must be open between ALL cluster nodes. Missing any port causes '
        'partial or complete cluster failure.'
    )

    add_styled_table(doc,
        ["Port", "Protocol", "Purpose", "If Blocked"],
        [
            ["${RMQ_PORT} (5672)", "TCP", "AMQP client connections", "Clients cannot connect"],
            ["${RMQ_MGMT_PORT} (15672)", "TCP", "Management UI & API", "No dashboard, no API monitoring"],
            ["${RMQ_PROM_PORT} (15692)", "TCP", "Prometheus metrics", "Datadog/monitoring blind"],
            ["${RMQ_DIST_PORT} (25672)", "TCP", "Erlang distribution (inter-node)", "CLUSTER BREAKS — nodes cannot communicate"],
            ["4369", "TCP", "EPMD (Erlang Port Mapper)", "Nodes cannot discover each other"],
        ]
    )

    add_warning_box(doc,
        "Port 25672 (Erlang distribution) and 4369 (EPMD) are the most critical. "
        "If a firewall rule blocks these ports between nodes, the cluster will partition immediately.",
        "CRITICAL"
    )

    # A.10 What NOT To Do
    doc.add_heading('A.10  What NOT To Do — RabbitMQ', level=2)

    danger_items = [
        ("NEVER stop 2+ nodes simultaneously",
         "With pause_minority, the remaining node pauses itself. Total cluster outage."),
        ("NEVER delete the Mnesia directory",
         "Permanently destroys queue definitions, user accounts, and quorum queue data."),
        ("NEVER change the Erlang cookie on a running cluster",
         "Nodes cannot re-authenticate. Cluster splits permanently until all nodes restarted."),
        ("NEVER force-remove a node that has quorum queue data",
         "rabbitmqctl forget_cluster_node --offline will discard that node's Raft state permanently."),
        ("NEVER restart all nodes at the same time",
         "Race condition on startup — last-node-to-stop must start first for quorum queue recovery."),
        ("NEVER skip 'await_online_quorum_plus_one' during rolling restarts",
         "Moving to next node before sync completes can lose quorum queue data."),
        ("NEVER disable the management plugin during an incident",
         "You lose API access for diagnostics and monitoring."),
        ("NEVER run rabbitmqctl reset on a node you want to keep in the cluster",
         "This wipes Mnesia and removes the node from cluster membership."),
        ("NEVER upgrade Erlang without checking RabbitMQ compatibility",
         "Wrong Erlang version can prevent RabbitMQ from starting entirely."),
    ]

    add_styled_table(doc,
        ["Action", "Consequence"],
        [(a, c) for a, c in danger_items]
    )

    doc.add_page_break()


def build_section_b(doc):
    """Redis + Sentinel Critical Operational Notes."""
    doc.add_heading('Section B: Redis + Sentinel Critical Operational Notes', level=1)

    # B.1 Sentinel Quorum
    doc.add_heading('B.1  Sentinel Quorum (2 of 3)', level=2)
    doc.add_paragraph(
        'Your cluster runs 3 Sentinel instances (one per node) with quorum = 2. '
        'Sentinel is the automatic failover mechanism — if it loses quorum, failover stops working.'
    )

    add_warning_box(doc,
        "With 3 Sentinels and quorum=2: if 2 Sentinels are down, automatic failover is IMPOSSIBLE. "
        "If the master then fails, there is NO automatic promotion. You must perform MANUAL failover.",
        "CRITICAL"
    )

    add_styled_table(doc,
        ["Sentinels Up", "Quorum Met?", "Auto-Failover?", "Risk Level"],
        [
            ["3 of 3", "Yes (3 >= 2)", "Yes", "Normal operations"],
            ["2 of 3", "Yes (2 >= 2)", "Yes, but no margin", "HIGH — one more failure disables failover"],
            ["1 of 3", "No (1 < 2)", "NO", "CRITICAL — manual failover only"],
            ["0 of 3", "No", "NO", "EMERGENCY — no monitoring, no failover"],
        ]
    )

    doc.add_paragraph("Sentinel failover process:")
    add_bullet_list(doc, [
        "Sentinel detects master is down (SDOWN — subjective down, after down-after-milliseconds=5000ms)",
        "Sentinel asks other Sentinels to confirm (ODOWN — objective down, needs quorum agreement)",
        "One Sentinel is elected as leader to perform failover",
        "Leader promotes the best replica to master",
        "Other replicas are reconfigured to replicate from new master",
        "Old master is reconfigured as replica when it comes back",
    ])

    # B.2 Never Restart Master First
    doc.add_heading('B.2  Never Restart Master First', level=2)

    add_warning_box(doc,
        "During planned maintenance: ALWAYS restart replicas first. NEVER restart the master node first. "
        "If you must take the master down, trigger a manual failover BEFORE stopping it.",
        "CRITICAL"
    )

    doc.add_paragraph("Why this matters:")
    add_bullet_list(doc, [
        "Restarting the master triggers an unplanned failover — Sentinel promotes a replica",
        "During failover, there is a brief window where writes can be lost (not yet replicated)",
        "If min-replicas-to-write=1 and you already took a replica down, the master refuses writes BEFORE failover",
        "Manual failover is graceful — it waits for replication sync before switching",
        "Unplanned failover is NOT graceful — it promotes whatever replica has the most data",
    ])

    doc.add_paragraph("Correct order for planned master maintenance:")
    add_numbered_steps(doc, [
        "Identify current master:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL get-master-addr-by-name ${SENTINEL_MASTER_NAME}",
        "Trigger manual failover (from any Sentinel):\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL failover ${SENTINEL_MASTER_NAME}",
        "Wait 15 seconds for failover to complete",
        "Verify new master:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL get-master-addr-by-name ${SENTINEL_MASTER_NAME}",
        "Now the old master is a replica — safe to restart/patch",
    ])

    # B.3 min-replicas-to-write
    doc.add_heading('B.3  min-replicas-to-write Constraint', level=2)

    add_warning_box(doc,
        "Your Redis config has min-replicas-to-write = 1 and min-replicas-max-lag = 10. "
        "This means: if ZERO replicas are connected (or all lag > 10 seconds), the master "
        "STOPS ACCEPTING WRITES. Applications will receive errors.",
        "CRITICAL"
    )

    doc.add_paragraph("Implications for maintenance:")
    add_bullet_list(doc, [
        "In a 3-node cluster (1 master + 2 replicas): you can safely take DOWN 1 replica",
        "If you take down BOTH replicas: master stops accepting writes immediately",
        "During OS patching: only patch ONE replica at a time, verify it's back and synced before next",
        "If a replica's replication lag exceeds 10 seconds, it doesn't count as 'connected'",
        "Monitor with: INFO replication → check connected_slaves and slave lag",
    ])

    # B.4 Sentinel Config Auto-Rewrite
    doc.add_heading('B.4  Sentinel Config Auto-Rewrite', level=2)

    add_warning_box(doc,
        "Sentinel REWRITES its own configuration file (sentinel.conf) at runtime. "
        "It adds discovered replicas, other sentinels, failover epoch, and its own ID. "
        "NEVER replace sentinel.conf with a template copy during maintenance — you will lose "
        "runtime state and break Sentinel's awareness of the cluster.",
        "CRITICAL"
    )

    doc.add_paragraph("What Sentinel writes to its config:")
    add_bullet_list(doc, [
        "sentinel myid <unique-40-char-hex> — unique identity, MUST NOT be duplicated",
        "sentinel known-replica <master> <ip> <port> — discovered replicas",
        "sentinel known-sentinel <master> <ip> <port> <runid> — discovered sentinels",
        "sentinel config-epoch <master> <epoch> — failover epoch counter",
        "sentinel leader-epoch <master> <epoch>",
        "sentinel current-epoch <epoch>",
    ])

    add_warning_box(doc,
        "If you need to redeploy sentinel.conf, KEEP the 'sentinel myid' line and all 'known-*' lines "
        "from the running config. Only modify the settings portion above the runtime section.",
        "WARNING"
    )

    # B.5 RDB Fork Memory
    doc.add_heading('B.5  RDB Fork Memory Requirement', level=2)

    add_warning_box(doc,
        "When Redis performs BGSAVE (RDB snapshot), it FORKS the process. The fork requires "
        "copy-on-write memory. On a server with 12GB maxmemory, the fork can temporarily need "
        "up to 12GB additional memory if there are many writes during the save. "
        "A 16GB server with 12GB Redis is NOT enough — it will OOM-kill.",
        "CRITICAL"
    )

    doc.add_paragraph("Memory planning:")
    add_bullet_list(doc, [
        "Rule of thumb: server RAM >= 2x maxmemory (for fork headroom) + OS overhead",
        "Example: 12GB maxmemory → need at least 28GB RAM (12GB data + 12GB fork + 4GB OS)",
        "vm.overcommit_memory should be set to 1 (allow overcommit) to prevent fork failures",
        "Alternative: set to 2 with enough swap as safety net",
        "If BGSAVE fails with 'Can't save in background: fork: Cannot allocate memory' — this is the cause",
        "AOF rewrite also forks — same memory concern applies",
    ])

    add_code_block(doc,
        "# Check overcommit setting (should be 1):\n"
        "sysctl vm.overcommit_memory\n\n"
        "# Set it (persistent):\n"
        "echo 'vm.overcommit_memory = 1' >> /etc/sysctl.d/redis.conf\n"
        "sysctl -p /etc/sysctl.d/redis.conf"
    )

    # B.6 AOF + RDB Interaction
    doc.add_heading('B.6  AOF + RDB Persistence Interaction', level=2)
    doc.add_paragraph(
        'Your cluster has BOTH RDB (save directives) and AOF (appendonly yes) enabled.'
    )

    add_warning_box(doc,
        "When both RDB and AOF are enabled, Redis loads AOF on startup (not RDB), because AOF is "
        "more complete. If the AOF file is corrupted, Redis will REFUSE TO START. "
        "You must repair the AOF before starting Redis.",
        "CRITICAL"
    )

    doc.add_paragraph("Recovery if AOF is corrupted:")
    add_numbered_steps(doc, [
        "Attempt auto-repair:\n"
        "   ${REDIS_CHECK_AOF} --fix ${REDIS_DATA_DIR}/appendonlydir/appendonly.aof.1.incr.aof",
        "If auto-repair fails, disable AOF temporarily to start from RDB:\n"
        "   Edit redis.conf: set 'appendonly no'\n"
        "   Start Redis — it will load from RDB (you lose data between last RDB and crash)",
        "Re-enable AOF after startup:\n"
        "   ${REDIS_CLI} ${REDIS_CMD_CONFIG} SET appendonly yes\n"
        "   This rewrites a fresh AOF from the current dataset",
        "Restore config file: set 'appendonly yes' back in redis.conf",
    ])

    # B.7 Renamed Commands
    doc.add_heading('B.7  Renamed/Disabled Commands', level=2)
    doc.add_paragraph(
        'Your production Redis config renames dangerous commands for security. '
        'L1/L2 engineers must know the renamed versions:'
    )

    add_styled_table(doc,
        ["Original Command", "Status", "Renamed To", "Why"],
        [
            ["FLUSHDB", "DISABLED", "(empty string)", "Would delete all keys in a database"],
            ["FLUSHALL", "DISABLED", "(empty string)", "Would delete ALL keys across ALL databases"],
            ["KEYS", "DISABLED", "(empty string)", "O(N) scan — blocks Redis on large datasets"],
            ["DEBUG", "DISABLED", "(empty string)", "Can crash Redis intentionally"],
            ["CONFIG", "RENAMED", "${REDIS_CMD_CONFIG}", "Prevents unauthorized config changes"],
            ["SHUTDOWN", "RENAMED", "${REDIS_CMD_SHUTDOWN}", "Prevents accidental shutdown"],
            ["BGSAVE", "RENAMED", "${REDIS_CMD_BGSAVE}", "Controls who can trigger saves"],
            ["BGREWRITEAOF", "RENAMED", "${REDIS_CMD_BGREWRITEAOF}", "Controls AOF rewrites"],
            ["REPLICAOF", "RENAMED", "${REDIS_CMD_REPLICAOF}", "Prevents changing replication topology"],
            ["SLAVEOF", "DISABLED", "(empty string)", "Deprecated, use REPLICAOF"],
        ]
    )

    add_warning_box(doc,
        "If you need to run CONFIG, BGSAVE, or SHUTDOWN, you MUST use the renamed version. "
        "Running the original command name will return '(error) ERR unknown command'. "
        "These renamed command tokens are secrets — do not share outside the ops team.",
        "IMPORTANT"
    )

    # B.8 Replication Backlog
    doc.add_heading('B.8  Replication Backlog & Full Resync', level=2)
    doc.add_paragraph(
        'When a replica reconnects after being down, it attempts a PARTIAL resync using the '
        'replication backlog. If the replica was down too long, a FULL resync is required.'
    )

    add_warning_box(doc,
        "Full resync is EXPENSIVE: the master forks (needing 2x memory), generates an RDB, "
        "and transfers it over the network. During transfer, the replica is unavailable. "
        "For a 12GB dataset on a 1Gbps link, this takes ~2 minutes minimum.",
        "WARNING"
    )

    doc.add_paragraph("Your config: repl-backlog-size = 256mb, repl-backlog-ttl = 3600")
    add_bullet_list(doc, [
        "The backlog holds the last 256MB of write commands",
        "If a replica misses more than 256MB of writes, it triggers full resync",
        "At 10MB/s write rate, the backlog covers ~25 seconds of writes",
        "At 1MB/s write rate, the backlog covers ~256 seconds (~4 minutes)",
        "If replica downtime exceeds backlog coverage → FULL RESYNC → high load on master",
        "For long maintenance windows: consider increasing repl-backlog-size temporarily",
    ])

    # B.9 What NOT To Do
    doc.add_heading('B.9  What NOT To Do — Redis', level=2)

    danger_items = [
        ("NEVER restart all 3 nodes simultaneously",
         "Total outage. Sentinel cannot failover. Data loss risk."),
        ("NEVER restart the master without failover first",
         "Triggers unplanned failover with potential data loss window."),
        ("NEVER replace sentinel.conf with a template",
         "Destroys runtime state (myid, known replicas, epochs). Sentinel breaks."),
        ("NEVER run FLUSHALL/FLUSHDB (they are disabled)",
         "Would wipe all data. Disabled in your config for safety."),
        ("NEVER change requirepass without updating masterauth on replicas",
         "Replicas cannot authenticate to master. Replication breaks immediately."),
        ("NEVER set maxmemory > 75% of physical RAM",
         "BGSAVE fork will OOM-kill Redis or other processes."),
        ("NEVER delete RDB/AOF files on a running master",
         "Next restart will start with empty dataset. All data lost."),
        ("NEVER run KEYS * in production",
         "Disabled in your config. Even if enabled: blocks Redis for seconds on large datasets."),
        ("NEVER stop both replicas while master has min-replicas-to-write=1",
         "Master immediately stops accepting writes. Application errors."),
    ]

    add_styled_table(doc,
        ["Action", "Consequence"],
        [(a, c) for a, c in danger_items]
    )

    doc.add_page_break()


def build_section_c(doc):
    """OS Patching Playbook."""
    doc.add_heading('Section C: OS Patching Playbook', level=1)
    doc.add_paragraph(
        'Step-by-step procedures for patching operating systems on RabbitMQ and Redis nodes. '
        'These procedures maintain cluster availability throughout the patching process.'
    )

    # C.1 Pre-Patch Checklist
    doc.add_heading('C.1  Pre-Patch Checklist (Both Systems)', level=2)

    add_warning_box(doc,
        "Complete ALL items in this checklist BEFORE starting any patching activity. "
        "If any check fails, resolve it before proceeding.",
        "IMPORTANT"
    )

    add_styled_table(doc,
        ["#", "Check", "Command / Action", "Expected Result"],
        [
            ["1", "RabbitMQ: All 3 nodes running",
             "sudo ${RABBITMQCTL} cluster_status",
             "3 running_nodes, 0 partitions"],
            ["2", "RabbitMQ: No alarms",
             "sudo ${RABBITMQCTL} cluster_status | grep alarms",
             "No memory or disk alarms"],
            ["3", "RabbitMQ: No critical quorum queues",
             "sudo ${RABBITMQ_DIAGNOSTICS} check_if_node_is_quorum_critical",
             "No warnings"],
            ["4", "Redis: All 3 nodes responding",
             "${REDIS_CLI} -h <each-node> PING",
             "PONG from all 3"],
            ["5", "Redis: Sentinel quorum healthy",
             "${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL ckquorum ${SENTINEL_MASTER_NAME}",
             "OK 3 usable Sentinels"],
            ["6", "Redis: Replication lag = 0",
             "${REDIS_CLI} INFO replication | grep lag",
             "lag=0 on all replicas"],
            ["7", "Redis: Last BGSAVE successful",
             "${REDIS_CLI} INFO persistence | grep rdb_last_bgsave",
             "rdb_last_bgsave_status:ok"],
            ["8", "Notify stakeholders",
             "Send maintenance window notification",
             "Acknowledgment received"],
            ["9", "Verify backups exist",
             "Check recent RDB snapshot and RabbitMQ definitions export",
             "Backups < 24h old"],
            ["10", "Confirm rollback plan",
             "Review rollback section C.5",
             "Team understands rollback steps"],
        ]
    )

    # C.2 RabbitMQ Patching Sequence
    doc.add_heading('C.2  RabbitMQ Patching Sequence', level=2)

    add_warning_box(doc,
        "Patch ONE node at a time. Order: Node3 (least critical) → Node2 → Node1. "
        "Wait for FULL cluster sync between each node. Minimum 2 nodes MUST be running at all times.",
        "CRITICAL"
    )

    doc.add_heading('Step-by-Step: Patch RabbitMQ Node 3', level=3)
    add_numbered_steps(doc, [
        "Verify cluster is healthy (3 nodes, no alarms, no partitions):\n"
        "   sudo ${RABBITMQCTL} cluster_status",

        "Export definitions as backup:\n"
        "   curl -s -u ${RMQ_ADMIN_USER}:${RMQ_ADMIN_PASS} http://${RMQ_NODE1}:${RMQ_MGMT_PORT}/api/definitions > /tmp/rmq_definitions_backup.json",

        "Drain Node 3 — stop accepting new connections:\n"
        "   ssh ${RMQ_NODE3} 'sudo ${RABBITMQ_UPGRADE} drain'\n"
        "   Wait 30-60 seconds for existing connections to close",

        "Stop RabbitMQ on Node 3:\n"
        "   ssh ${RMQ_NODE3} 'sudo systemctl stop ${RMQ_SERVICE}'",

        "Verify 2 remaining nodes are healthy:\n"
        "   sudo ${RABBITMQCTL} cluster_status\n"
        "   Expected: 2 running_nodes, Node3 listed in disc nodes but not running",

        "Perform OS patching on Node 3:\n"
        "   ssh ${RMQ_NODE3} 'sudo yum update -y'   # or dnf, apt\n"
        "   ssh ${RMQ_NODE3} 'sudo reboot'           # if kernel update",

        "After reboot, start RabbitMQ on Node 3:\n"
        "   ssh ${RMQ_NODE3} 'sudo systemctl start ${RMQ_SERVICE}'",

        "Revive Node 3 — accept connections again:\n"
        "   ssh ${RMQ_NODE3} 'sudo ${RABBITMQ_UPGRADE} revive'",

        "Wait for quorum queue sync — THIS IS CRITICAL:\n"
        "   ssh ${RMQ_NODE3} 'sudo ${RABBITMQCTL} await_online_quorum_plus_one'\n"
        "   This blocks until all quorum queues have replicated to Node 3\n"
        "   DO NOT proceed to Node 2 until this returns successfully",

        "Final verification:\n"
        "   sudo ${RABBITMQCTL} cluster_status\n"
        "   Expected: 3 running_nodes, 0 partitions, no alarms",
    ])

    add_warning_box(doc,
        "Repeat the EXACT same steps for Node 2, then Node 1. "
        "NEVER skip step 9 (await_online_quorum_plus_one).",
        "CRITICAL"
    )

    # C.3 Redis Patching Sequence
    doc.add_heading('C.3  Redis Patching Sequence', level=2)

    add_warning_box(doc,
        "Patch REPLICAS first, then FAILOVER the master, then patch the old master. "
        "Order: Replica2 → Replica1 → Manual Failover → Old Master. "
        "NEVER patch the master while it is still master.",
        "CRITICAL"
    )

    doc.add_heading('Step-by-Step: Patch Redis Replica 2 (e.g., Node 3)', level=3)
    add_numbered_steps(doc, [
        "Identify current roles:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL get-master-addr-by-name ${SENTINEL_MASTER_NAME}\n"
        "   ${REDIS_CLI} -h ${REDIS_NODE1} ROLE\n"
        "   ${REDIS_CLI} -h ${REDIS_NODE2} ROLE\n"
        "   ${REDIS_CLI} -h ${REDIS_NODE3} ROLE",

        "Verify replication is current (lag = 0):\n"
        "   ${REDIS_CLI} -h <master-node> INFO replication",

        "Trigger BGSAVE on master (create fresh backup):\n"
        "   ${REDIS_CLI} -h <master-node> ${REDIS_CMD_BGSAVE}\n"
        "   Wait for: ${REDIS_CLI} -h <master-node> INFO persistence | grep rdb_bgsave_in_progress\n"
        "   (wait until rdb_bgsave_in_progress:0)",

        "Stop Redis on Replica 2 (the node being patched):\n"
        "   ssh ${REDIS_NODE3} 'sudo systemctl stop ${REDIS_SERVICE}'\n"
        "   ssh ${REDIS_NODE3} 'sudo systemctl stop ${SENTINEL_SERVICE}'",

        "Verify master still has 1 connected replica:\n"
        "   ${REDIS_CLI} -h <master-node> INFO replication | grep connected_slaves\n"
        "   Expected: connected_slaves:1 (min-replicas-to-write still satisfied)",

        "Perform OS patching on Node 3:\n"
        "   ssh ${REDIS_NODE3} 'sudo yum update -y'\n"
        "   ssh ${REDIS_NODE3} 'sudo reboot'   # if kernel update",

        "After reboot, start Redis + Sentinel:\n"
        "   ssh ${REDIS_NODE3} 'sudo systemctl start ${REDIS_SERVICE}'\n"
        "   ssh ${REDIS_NODE3} 'sudo systemctl start ${SENTINEL_SERVICE}'",

        "Verify replica rejoined and is syncing:\n"
        "   ${REDIS_CLI} -h ${REDIS_NODE3} INFO replication\n"
        "   Expected: role:slave, master_link_status:up\n"
        "   Wait for master_repl_offset to match master",

        "Verify Sentinel sees all 3 nodes:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL ckquorum ${SENTINEL_MASTER_NAME}",
    ])

    doc.add_paragraph()
    doc.add_heading('Step-by-Step: Patch the Current Master', level=3)
    add_numbered_steps(doc, [
        "After BOTH replicas are patched and synced, identify the master:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL get-master-addr-by-name ${SENTINEL_MASTER_NAME}",

        "Trigger MANUAL FAILOVER (graceful — waits for sync):\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL failover ${SENTINEL_MASTER_NAME}\n"
        "   Wait 15 seconds",

        "Verify failover completed — new master should be a different node:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL get-master-addr-by-name ${SENTINEL_MASTER_NAME}",

        "Old master is now a replica — safe to stop:\n"
        "   ssh <old-master> 'sudo systemctl stop ${REDIS_SERVICE}'\n"
        "   ssh <old-master> 'sudo systemctl stop ${SENTINEL_SERVICE}'",

        "Patch the old master node:\n"
        "   ssh <old-master> 'sudo yum update -y && sudo reboot'",

        "After reboot, start Redis + Sentinel:\n"
        "   ssh <old-master> 'sudo systemctl start ${REDIS_SERVICE}'\n"
        "   ssh <old-master> 'sudo systemctl start ${SENTINEL_SERVICE}'",

        "Final verification — all 3 nodes up, Sentinel quorum healthy:\n"
        "   ${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL ckquorum ${SENTINEL_MASTER_NAME}\n"
        "   ${REDIS_CLI} -h <new-master> INFO replication",
    ])

    # C.4 Post-Patch Validation
    doc.add_heading('C.4  Post-Patch Validation', level=2)

    add_styled_table(doc,
        ["#", "System", "Validation", "Command"],
        [
            ["1", "RabbitMQ", "All 3 nodes running",
             "sudo ${RABBITMQCTL} cluster_status"],
            ["2", "RabbitMQ", "No alarms (memory/disk)",
             "sudo ${RABBITMQCTL} cluster_status | grep alarms"],
            ["3", "RabbitMQ", "No partitions",
             "sudo ${RABBITMQCTL} cluster_status | grep partitions"],
            ["4", "RabbitMQ", "Quorum queues healthy",
             "sudo ${RABBITMQ_DIAGNOSTICS} check_if_node_is_quorum_critical"],
            ["5", "RabbitMQ", "Consumers reconnected",
             "Check management UI → Connections tab"],
            ["6", "RabbitMQ", "Message rates normal",
             "Check management UI → Overview"],
            ["7", "Redis", "All 3 nodes PONG",
             "${REDIS_CLI} -h <node> PING (each node)"],
            ["8", "Redis", "Replication healthy",
             "${REDIS_CLI} -h <master> INFO replication"],
            ["9", "Redis", "Sentinel quorum OK",
             "${REDIS_CLI} -p ${SENTINEL_PORT} SENTINEL ckquorum ${SENTINEL_MASTER_NAME}"],
            ["10", "Redis", "No replication lag",
             "INFO replication → check offset match"],
            ["11", "Both", "Datadog monitors green",
             "Check Datadog dashboard — all monitors OK"],
            ["12", "Both", "Application health",
             "Verify application logs — no connection errors"],
        ]
    )

    # C.5 Rollback Procedures
    doc.add_heading('C.5  Rollback Procedures', level=2)

    doc.add_heading('RabbitMQ Rollback', level=3)
    add_bullet_list(doc, [
        "If a node fails to rejoin after patching: check Erlang cookie, check Mnesia directory exists",
        "If Erlang version changed during patch and RabbitMQ won't start: downgrade Erlang package",
        "If quorum queues are in bad state: let 2 healthy nodes serve, investigate the 3rd",
        "If definitions are lost: restore from backup:\n"
        "   curl -s -u admin:pass -X POST -H 'content-type:application/json' "
        "-d @/tmp/rmq_definitions_backup.json http://node:15672/api/definitions",
        "LAST RESORT: If cluster is unrecoverable, reset one node and rebuild:\n"
        "   sudo ${RABBITMQCTL} force_boot   (on the node with latest data)\n"
        "   Then join other nodes to it",
    ])

    doc.add_heading('Redis Rollback', level=3)
    add_bullet_list(doc, [
        "If a replica won't start after patching: check redis.conf syntax, check data directory permissions",
        "If AOF corruption prevents startup: repair with redis-check-aof --fix",
        "If Sentinel lost track of master: restart all 3 Sentinels (they re-discover)",
        "If data loss occurred: restore from last RDB backup:\n"
        "   1. Stop Redis on target node\n"
        "   2. Copy dump.rdb to ${REDIS_DATA_DIR}/\n"
        "   3. Ensure appendonly is disabled temporarily\n"
        "   4. Start Redis\n"
        "   5. Re-enable AOF: ${REDIS_CLI} ${REDIS_CMD_CONFIG} SET appendonly yes",
        "If all nodes are down: start the node with the LATEST RDB/AOF first (check file timestamps)",
    ])

    doc.add_page_break()


def build_section_d(doc):
    """Emergency Quick Reference."""
    doc.add_heading('Section D: Emergency Quick Reference', level=1)

    # D.1 RabbitMQ DO and DO NOT
    doc.add_heading('D.1  DO and DO NOT — RabbitMQ', level=2)

    add_styled_table(doc,
        ["DO", "DO NOT"],
        [
            ["Drain before stopping: sudo ${RABBITMQ_UPGRADE} drain",
             "Stop without draining — causes abrupt client disconnections"],
            ["Wait for sync: sudo ${RABBITMQCTL} await_online_quorum_plus_one",
             "Skip sync wait — risks quorum queue data loss"],
            ["Check quorum critical: sudo ${RABBITMQ_DIAGNOSTICS} check_if_node_is_quorum_critical",
             "Stop a quorum-critical node — makes queues unavailable"],
            ["Patch one node at a time",
             "Patch 2+ nodes simultaneously — cluster pauses (pause_minority)"],
            ["Export definitions before maintenance",
             "Assume definitions are safe — Mnesia corruption happens"],
            ["Verify Erlang cookie after node rebuild",
             "Change cookie on a live cluster — instant partition"],
            ["Monitor partitions: rabbitmqctl cluster_status",
             "Ignore partition warnings — split-brain causes data inconsistency"],
            ["Use management API for diagnostics during incidents",
             "Disable management plugin to 'save resources' during incident"],
        ]
    )

    # D.2 Redis DO and DO NOT
    doc.add_heading('D.2  DO and DO NOT — Redis', level=2)

    add_styled_table(doc,
        ["DO", "DO NOT"],
        [
            ["Failover before master maintenance: SENTINEL failover",
             "Stop master directly — unplanned failover loses data"],
            ["Patch replicas first, then failover, then patch old master",
             "Patch master first — write outage if min-replicas-to-write=1"],
            ["Check Sentinel quorum before maintenance: SENTINEL ckquorum",
             "Assume Sentinel is fine — quorum loss = no auto-failover"],
            ["Backup sentinel.conf before changes",
             "Replace sentinel.conf with template — loses runtime state"],
            ["Use renamed commands (${REDIS_CMD_CONFIG}, etc.)",
             "Try to run CONFIG, SHUTDOWN directly — they are disabled"],
            ["Verify repl lag = 0 before stopping a replica",
             "Stop replica with high lag — may trigger full resync on return"],
            ["Keep vm.overcommit_memory = 1",
             "Set to 0 — BGSAVE fork will fail on large datasets"],
            ["Monitor connected_slaves count",
             "Ignore — if count drops to 0, master stops writes"],
        ]
    )

    # D.3 Decision Tree
    doc.add_heading('D.3  Decision Tree: Can I Take This Node Down?', level=2)

    doc.add_heading('RabbitMQ', level=3)
    add_styled_table(doc,
        ["Question", "Answer", "Action"],
        [
            ["How many nodes are currently running?", "3", "Safe to take 1 down (follow procedure)"],
            ["How many nodes are currently running?", "2", "DO NOT take another down — cluster will halt"],
            ["How many nodes are currently running?", "1", "EMERGENCY — cluster is already halted"],
            ["Is target node quorum-critical?", "Yes", "Wait for other nodes to sync first"],
            ["Is target node quorum-critical?", "No", "Safe to proceed"],
            ["Are there active memory/disk alarms?", "Yes", "Resolve alarms FIRST, then consider maintenance"],
            ["Are there active partitions?", "Yes", "Resolve partition FIRST — do not add more disruption"],
        ]
    )

    doc.add_heading('Redis', level=3)
    add_styled_table(doc,
        ["Question", "Answer", "Action"],
        [
            ["Is the target node the master?", "Yes", "FAILOVER first, then stop as replica"],
            ["Is the target node the master?", "No (replica)", "Proceed (check replica count)"],
            ["How many replicas are connected?", "2", "Safe to stop 1 replica"],
            ["How many replicas are connected?", "1", "Stopping this replica will halt writes (min-replicas-to-write=1)"],
            ["How many replicas are connected?", "0", "EMERGENCY — master already refusing writes"],
            ["How many Sentinels are running?", "3", "Safe to stop 1 (quorum maintained)"],
            ["How many Sentinels are running?", "2", "WARNING — stopping 1 more breaks quorum"],
            ["How many Sentinels are running?", "1", "DO NOT stop — failover already impossible"],
            ["Is replication lag > 0?", "Yes", "Wait for sync before stopping replica"],
        ]
    )

    doc.add_page_break()


def build_section_e(doc):
    """Directory Structure & Variable Reference."""
    doc.add_heading('Section E: Directory Structure & Variable Reference', level=1)

    # E.1 RabbitMQ Directory Layout
    doc.add_heading('E.1  RabbitMQ Directory Layout', level=2)
    doc.add_paragraph(
        'RabbitMQ can be installed via package manager (yum/dnf/apt) or from extracted RPMs. '
        'The table below shows both layouts and the environment variable to configure each path.'
    )

    add_styled_table(doc,
        ["Component", "Package Install Path", "RPM Extracted Path (example)", "Variable"],
        [
            ["Home / Root", "/usr/lib/rabbitmq", "/opt/rabbitmq/4.1.4", "RMQ_HOME"],
            ["Binaries (rabbitmqctl, etc.)", "/usr/lib/rabbitmq/bin  or  /usr/sbin", "${RMQ_HOME}/sbin", "RMQ_SBIN_DIR"],
            ["rabbitmqctl", "/usr/sbin/rabbitmqctl", "${RMQ_SBIN_DIR}/rabbitmqctl", "RABBITMQCTL"],
            ["rabbitmq-diagnostics", "/usr/sbin/rabbitmq-diagnostics", "${RMQ_SBIN_DIR}/rabbitmq-diagnostics", "RABBITMQ_DIAGNOSTICS"],
            ["rabbitmq-plugins", "/usr/sbin/rabbitmq-plugins", "${RMQ_SBIN_DIR}/rabbitmq-plugins", "RABBITMQ_PLUGINS"],
            ["rabbitmq-upgrade", "/usr/sbin/rabbitmq-upgrade", "${RMQ_SBIN_DIR}/rabbitmq-upgrade", "RABBITMQ_UPGRADE"],
            ["Config directory", "/etc/rabbitmq", "/etc/rabbitmq (or custom)", "RMQ_CONF_DIR"],
            ["rabbitmq.conf", "/etc/rabbitmq/rabbitmq.conf", "${RMQ_CONF_DIR}/rabbitmq.conf", "RMQ_CONFIG"],
            ["enabled_plugins", "/etc/rabbitmq/enabled_plugins", "${RMQ_CONF_DIR}/enabled_plugins", "RMQ_ENABLED_PLUGINS"],
            ["SSL/TLS certs", "/etc/rabbitmq/ssl/", "${RMQ_CONF_DIR}/ssl/", "RMQ_SSL_DIR"],
            ["Plugins", "/usr/lib/rabbitmq/plugins", "${RMQ_HOME}/plugins", "RMQ_PLUGINS_DIR"],
            ["Data (Mnesia)", "/var/lib/rabbitmq", "/var/lib/rabbitmq (or custom)", "RMQ_DATA_DIR"],
            ["Mnesia DB", "/var/lib/rabbitmq/mnesia", "${RMQ_DATA_DIR}/mnesia", "RMQ_MNESIA_DIR"],
            ["Erlang cookie", "/var/lib/rabbitmq/.erlang.cookie", "${RMQ_DATA_DIR}/.erlang.cookie", "RMQ_COOKIE_FILE"],
            ["Logs", "/var/log/rabbitmq", "/var/log/rabbitmq (or custom)", "RMQ_LOG_DIR"],
            ["Erlang/OTP", "/usr/lib64/erlang", "/opt/erlang/26.x", "ERLANG_HOME"],
            ["Service name", "rabbitmq-server", "rabbitmq-server (or custom)", "RMQ_SERVICE"],
            ["Run-as user", "rabbitmq", "rabbitmq (or custom)", "RMQ_RUN_USER"],
        ]
    )

    # E.2 Redis Directory Layout
    doc.add_heading('E.2  Redis Directory Layout', level=2)
    doc.add_paragraph(
        'Redis is installed from extracted tarball/RPM under /opt/cached/. '
        'A symlink /opt/cached/current points to the active version.'
    )

    add_styled_table(doc,
        ["Component", "Your Current Path", "Variable", "Notes"],
        [
            ["Home / Root", "/opt/cached/current (symlink)", "REDIS_HOME", "Symlink to /opt/cached/redis-8.2.2"],
            ["Binaries", "/opt/cached/current/bin", "REDIS_BIN_DIR", "Contains redis-server, redis-cli, redis-sentinel"],
            ["redis-cli", "/opt/cached/current/bin/redis-cli", "REDIS_CLI", "Used by all runbook scripts"],
            ["redis-server", "/opt/cached/current/bin/redis-server", "REDIS_SERVER_BIN", "The server binary"],
            ["redis-sentinel", "/opt/cached/current/bin/redis-sentinel", "REDIS_SENTINEL_BIN", "Sentinel binary"],
            ["redis-check-aof", "/opt/cached/current/bin/redis-check-aof", "REDIS_CHECK_AOF", "AOF repair tool"],
            ["redis-check-rdb", "/opt/cached/current/bin/redis-check-rdb", "REDIS_CHECK_RDB", "RDB validation tool"],
            ["Config directory", "/opt/cached/current/conf", "REDIS_CONF_DIR", "Contains redis.conf, sentinel.conf"],
            ["redis.conf", "/opt/cached/current/conf/redis.conf", "REDIS_CONF_FILE", "Main config"],
            ["sentinel.conf", "/opt/cached/current/conf/sentinel.conf", "SENTINEL_CONF_FILE", "AUTO-REWRITTEN by Sentinel"],
            ["Data directory", "/opt/cached/current/data", "REDIS_DATA_DIR", "dump.rdb, appendonly dir"],
            ["Log directory", "/opt/cached/current/logs", "REDIS_LOG_DIR", "redis.log, sentinel.log"],
            ["PID directory", "/opt/cached/current/run", "REDIS_PID_DIR", "redis.pid, sentinel.pid"],
            ["Scripts", "/opt/cached/current/scripts", "REDIS_SCRIPTS_DIR", "Operational scripts"],
            ["Service name", "redis", "REDIS_SERVICE", "systemd unit"],
            ["Sentinel service", "redis-sentinel", "SENTINEL_SERVICE", "systemd unit"],
            ["Run-as user", "redis", "REDIS_RUN_USER", "Non-root user"],
        ]
    )

    # E.3 Environment Configuration
    doc.add_heading('E.3  How to Configure for Your Environment', level=2)

    doc.add_paragraph(
        'All variables are defined in scripts/runbook/env/environment.conf with sensible defaults. '
        'To customize for your environment, you have two options:'
    )

    doc.add_paragraph("Option 1: Set environment variables before running scripts:")
    add_code_block(doc,
        "export RMQ_HOME=/app/rabbitmq/4.1.4\n"
        "export RMQ_SBIN_DIR=/app/rabbitmq/4.1.4/sbin\n"
        "export REDIS_HOME=/opt/cached/redis-8.2.2\n"
        "export RMQ_NODE1=prod-rmq-01.example.com\n"
        "export REDIS_NODE1=prod-redis-01.example.com\n"
        "./scripts/runbook/rabbitmq/rb-rmq-007-health-check.sh"
    )

    doc.add_paragraph("Option 2: Create an environment-specific override file:")
    add_code_block(doc,
        "# Create: scripts/runbook/env/environment.production.conf\n"
        "# This is auto-loaded when ENVIRONMENT=production\n\n"
        "RMQ_HOME=/app/rabbitmq/4.1.4\n"
        "RMQ_SBIN_DIR=${RMQ_HOME}/sbin\n"
        "RMQ_NODE1=prod-rmq-01.example.com\n"
        "RMQ_NODE2=prod-rmq-02.example.com\n"
        "RMQ_NODE3=prod-rmq-03.example.com\n"
        "RMQ_ADMIN_PASS=your-secure-password\n\n"
        "REDIS_HOME=/opt/cached/redis-8.2.2\n"
        "REDIS_NODE1=prod-redis-01.example.com\n"
        "REDIS_NODE2=prod-redis-02.example.com\n"
        "REDIS_NODE3=prod-redis-03.example.com\n"
        "REDIS_AUTH_PASS=your-redis-password"
    )

    add_warning_box(doc,
        "Environment override files may contain passwords. Add them to .gitignore and "
        "restrict file permissions: chmod 600 environment.*.conf",
        "IMPORTANT"
    )


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Build document
    build_title_page(doc)
    build_toc(doc)
    build_section_a(doc)
    build_section_b(doc)
    build_section_c(doc)
    build_section_d(doc)
    build_section_e(doc)

    # Footer note
    doc.add_page_break()
    doc.add_heading('Document Control', level=1)
    add_styled_table(doc,
        ["Field", "Value"],
        [
            ["Document Title", "Critical Operations Guide — RabbitMQ & Redis HA Clusters"],
            ["Version", "1.0"],
            ["Classification", "INTERNAL — Operations Team Only"],
            ["RabbitMQ Version", "4.1.4 (Quorum Queues, pause_minority)"],
            ["Redis Version", "8.2.2 (Sentinel HA, 1 Master + 2 Replicas)"],
            ["Cluster Size", "3 nodes each (RabbitMQ + Redis)"],
            ["Partition Handling", "pause_minority (RabbitMQ), Sentinel quorum=2 (Redis)"],
            ["Persistence", "RDB + AOF (Redis), Mnesia + Quorum Raft (RabbitMQ)"],
        ]
    )

    doc.save(OUTPUT_FILE)
    print(f"Document saved: {OUTPUT_FILE}")


if __name__ == '__main__':
    main()
