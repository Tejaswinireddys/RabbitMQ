#!/usr/bin/env python3
"""
Generate comprehensive RabbitMQ Datadog Monitoring Dashboard Document
For a 3-Node RabbitMQ Cluster
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RabbitMQ 3-Node Cluster\nDatadog Monitoring & Dashboard Strategy")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph("")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Enterprise Monitoring Architecture & Dashboard Design")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")
doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    f"Version: 1.0",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "Classification: Internal / Operations",
    "Author: Monitoring Architecture Team",
]:
    meta.add_run(line + "\n").font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Monitoring Architecture Overview",
    "3. Datadog Integration Setup",
    "4. Dashboard #1 — Cluster Health Overview",
    "5. Dashboard #2 — Queue Monitoring & Analytics",
    "6. Dashboard #3 — Message Flow & Delivery",
    "7. Dashboard #4 — Failed / Dead-Lettered Messages",
    "8. Dashboard #5 — Consumer & Publisher Performance",
    "9. Dashboard #6 — Node Resource Utilization",
    "10. Dashboard #7 — Connection & Channel Monitoring",
    "11. Dashboard #8 — User & Permission Audit",
    "12. Dashboard #9 — Restart & Availability Tracking",
    "13. Dashboard #10 — Network Partition & Split-Brain",
    "14. Dashboard #11 — Quorum Queue Health",
    "15. Dashboard #12 — Executive / SLA Summary",
    "16. Alerting Strategy & Escalation Matrix",
    "17. Tagging & Naming Conventions",
    "18. Runbook References",
    "19. Appendix — Full Metric Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document defines the complete Datadog monitoring strategy for our production "
    "3-node RabbitMQ cluster. It covers 12 purpose-built dashboards, an alerting escalation "
    "matrix, tagging conventions, and runbook references designed to provide full operational "
    "visibility — from executive SLA summaries down to individual queue-level diagnostics."
)
doc.add_paragraph(
    "The monitoring architecture follows the RED method (Rate, Errors, Duration) and USE method "
    "(Utilization, Saturation, Errors) combined with RabbitMQ-specific best practices to ensure "
    "zero blind spots across messaging, clustering, and infrastructure layers."
)

doc.add_heading("Cluster Topology", level=2)
add_table(
    ["Property", "Value"],
    [
        ["Cluster Nodes", "3 (rabbitmq-node-1, rabbitmq-node-2, rabbitmq-node-3)"],
        ["RabbitMQ Version", "4.x (with Quorum Queues)"],
        ["Erlang/OTP", "26.x"],
        ["Queue Type", "Quorum Queues (primary), Classic Mirrored (legacy)"],
        ["Plugins", "Management, Prometheus, Shovel, Federation, Consistent Hash"],
        ["Monitoring Tool", "Datadog Agent + RabbitMQ Integration"],
        ["Metrics Source", "Prometheus endpoint + Management API"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. MONITORING ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Monitoring Architecture Overview", level=1)
doc.add_heading("2.1 Data Collection Layers", level=2)
doc.add_paragraph("A robust monitoring architecture collects data at multiple layers:")

add_table(
    ["Layer", "Source", "Collection Method", "Frequency"],
    [
        ["Application Metrics", "RabbitMQ Prometheus Plugin", "Datadog OpenMetrics Check", "15s"],
        ["Management API", "RabbitMQ Management Plugin", "Datadog RabbitMQ Integration", "15s"],
        ["System Metrics", "Host OS (CPU, Memory, Disk, Network)", "Datadog Agent", "15s"],
        ["Logs", "RabbitMQ log files + Erlang crash dumps", "Datadog Log Agent", "Real-time"],
        ["APM Traces", "Publisher/Consumer applications", "Datadog APM", "Real-time"],
        ["Process Checks", "Erlang beam.smp process", "Datadog Process Check", "30s"],
        ["Custom Checks", "rabbitmqctl commands", "Datadog Custom Check", "60s"],
        ["Synthetic Checks", "Heartbeat queue publish/consume", "Datadog Synthetic Monitor", "60s"],
    ]
)

doc.add_heading("2.2 Monitoring Pillars", level=2)
pillars = [
    ("Availability", "Is the cluster up? Are all nodes healthy and synchronized?"),
    ("Performance", "What is the message throughput? Are consumers keeping up?"),
    ("Saturation", "Are resources (memory, disk, FDs, connections) approaching limits?"),
    ("Errors", "Are messages failing? Are there dead-lettered or unroutable messages?"),
    ("Change Tracking", "When did restarts, deployments, or config changes happen?"),
    ("Security", "Who has access? Are permissions correct? Any unauthorized attempts?"),
]
for pillar, desc in pillars:
    p = doc.add_paragraph()
    run = p.add_run(f"{pillar}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. DATADOG INTEGRATION SETUP
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Datadog Integration Setup", level=1)

doc.add_heading("3.1 Datadog Agent Configuration", level=2)
doc.add_paragraph("Install the Datadog Agent on all 3 RabbitMQ nodes. Key configuration:")
add_code("""# /etc/datadog-agent/conf.d/rabbitmq.d/conf.yaml
init_config:

instances:
  - rabbitmq_api_url: http://localhost:15672/api/
    rabbitmq_user: datadog_monitor
    rabbitmq_pass: <ENCRYPTED_PASSWORD>
    tag_families: true
    collect_node_metrics: true
    collect_overview_metrics: true
    exchanges: [".*"]
    exchanges_regexes: [".*"]
    queues: [".*"]
    queues_regexes: [".*"]
    vhosts: [".*"]
    tags:
      - env:production
      - cluster:rabbitmq-prod
      - service:rabbitmq""")

doc.add_heading("3.2 Prometheus Metrics Collection", level=2)
add_code("""# /etc/datadog-agent/conf.d/openmetrics.d/conf.yaml
instances:
  - openmetrics_endpoint: http://localhost:15692/metrics
    namespace: rabbitmq
    metrics:
      - rabbitmq_queue_messages
      - rabbitmq_queue_messages_ready
      - rabbitmq_queue_messages_unacked
      - rabbitmq_queue_consumers
      - rabbitmq_channel_messages_published_total
      - rabbitmq_channel_messages_delivered_total
      - rabbitmq_channel_messages_acked_total
      - rabbitmq_channel_messages_redelivered_total
      - rabbitmq_connections
      - rabbitmq_channels
      - rabbitmq_process_resident_memory_bytes
      - rabbitmq_disk_space_available_bytes
      - rabbitmq_erlang_processes_used
    tags:
      - env:production
      - cluster:rabbitmq-prod""")

doc.add_heading("3.3 Dedicated Monitoring User", level=2)
doc.add_paragraph("Create a dedicated read-only Datadog monitoring user on RabbitMQ:")
add_code("""rabbitmqctl add_user datadog_monitor <secure_password>
rabbitmqctl set_user_tags datadog_monitor monitoring
rabbitmqctl set_permissions -p / datadog_monitor "" "" ".*" """)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. DASHBOARD #1 — CLUSTER HEALTH OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Dashboard #1 — Cluster Health Overview", level=1)
doc.add_paragraph(
    "The primary landing dashboard providing at-a-glance health status of the entire "
    "3-node RabbitMQ cluster. This is the first dashboard NOC/SRE teams should check."
)

doc.add_heading("4.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Cluster Status", "Query Value (Green/Red)", "rabbitmq.node.running per node", "Show if all 3 nodes are UP"],
        ["Node Health Grid", "Heatmap (3 cells)", "rabbitmq.node.running, mem_used, fd_used", "Per-node health at a glance"],
        ["Total Queues", "Query Value", "count(rabbitmq.queue.messages) by {queue}", "Total number of queues"],
        ["Total Connections", "Query Value", "sum:rabbitmq.connections{*}", "Active connections across all nodes"],
        ["Total Channels", "Query Value", "sum:rabbitmq.channels{*}", "Active channels across all nodes"],
        ["Messages Ready", "Timeseries", "sum:rabbitmq.queue.messages_ready{*}", "Messages waiting for consumers"],
        ["Messages Unacked", "Timeseries", "sum:rabbitmq.queue.messages_unacked{*}", "Delivered but not acknowledged"],
        ["Publish Rate", "Timeseries", "rate(rabbitmq.queue.messages.publish.count{*})", "Cluster-wide publish rate/sec"],
        ["Deliver Rate", "Timeseries", "rate(rabbitmq.queue.messages.deliver.count{*})", "Cluster-wide deliver rate/sec"],
        ["Cluster Alarms", "Event Stream", "rabbitmq alarms (memory, disk)", "Active resource alarms"],
        ["Erlang Process Usage", "Gauge per node", "rabbitmq.node.proc_used / proc_total", "Erlang VM process saturation"],
        ["File Descriptor Usage", "Gauge per node", "rabbitmq.node.fd_used / fd_total", "File descriptor saturation"],
        ["Memory Usage", "Gauge per node", "rabbitmq.node.mem_used / mem_limit", "Memory watermark proximity"],
        ["Disk Free Space", "Gauge per node", "rabbitmq.node.disk_free", "Disk space remaining per node"],
        ["Uptime per Node", "Query Value (3x)", "rabbitmq.node.uptime", "How long each node has been running"],
    ]
)

doc.add_heading("4.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Node Down", "rabbitmq.node.running", "== 0 for any node > 30s", "P1 - Critical", "Page on-call, trigger auto-recovery"],
        ["Memory Alarm", "rabbitmq.node.mem_alarm", "== 1", "P1 - Critical", "Investigate memory consumers"],
        ["Disk Alarm", "rabbitmq.node.disk_alarm", "== 1", "P1 - Critical", "Check log rotation, purge old data"],
        ["FD Usage High", "rabbitmq.node.fd_used_pct", "> 80%", "P2 - Warning", "Check connection leaks"],
        ["Erlang Procs High", "rabbitmq.node.proc_used_pct", "> 85%", "P2 - Warning", "Investigate runaway processes"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. DASHBOARD #2 — QUEUE MONITORING & ANALYTICS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Dashboard #2 — Queue Monitoring & Analytics", level=1)
doc.add_paragraph(
    "Deep-dive into queue-level metrics. Covers queue counts, depth trends, growth rates, "
    "and identifies problematic queues accumulating messages."
)

doc.add_heading("5.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Total Queue Count", "Query Value", "count:rabbitmq.queue.messages{*} by {queue}", "Number of declared queues"],
        ["Queue Count by VHost", "Top List", "count by {vhost}", "Queue distribution per vhost"],
        ["Queue Count by Type", "Pie Chart", "count by {queue_type}", "Quorum vs classic ratio"],
        ["Top 20 Deepest Queues", "Top List", "rabbitmq.queue.messages by {queue} top 20", "Queues with most messages"],
        ["Queue Depth Over Time", "Timeseries", "rabbitmq.queue.messages by {queue}", "Historical depth trend"],
        ["Queue Growth Rate", "Timeseries", "diff(rabbitmq.queue.messages) by {queue}", "Rate of depth change"],
        ["Empty vs Non-Empty", "Query Value", "count where messages == 0 / > 0", "Quick health indicator"],
        ["No Consumers", "Top List", "consumers == 0 AND messages > 0", "Orphaned queues"],
        ["Consumer Count by Queue", "Top List", "rabbitmq.queue.consumers by {queue}", "Consumer distribution"],
        ["Queue Memory Usage", "Top List", "rabbitmq.queue.memory by {queue}", "Memory per queue"],
        ["Queue Churn", "Timeseries", "Queue created/deleted events", "Queue lifecycle activity"],
        ["Idle Queues (>24h)", "Top List", "No publish/consume for 24h", "Cleanup candidates"],
    ]
)

doc.add_heading("5.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Queue Depth Critical", "rabbitmq.queue.messages", "> 100,000 for > 5m", "P1 - Critical", "Check consumer health, scale consumers"],
        ["Queue Depth Warning", "rabbitmq.queue.messages", "> 10,000 for > 15m", "P2 - Warning", "Investigate slow consumers"],
        ["No Consumers", "rabbitmq.queue.consumers", "== 0 AND messages > 100 for > 10m", "P2 - Warning", "Deploy/restart consumer service"],
        ["Queue Growth Anomaly", "diff(queue.messages)", "Anomaly detection (2 sigma)", "P3 - Info", "Investigate traffic spike"],
        ["Too Many Queues", "Total queue count", "> 5,000", "P2 - Warning", "Check for queue leak in app"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. DASHBOARD #3 — MESSAGE FLOW & DELIVERY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Dashboard #3 — Message Flow & Delivery", level=1)
doc.add_paragraph(
    "Tracks the complete message lifecycle: publish -> route -> deliver -> acknowledge. "
    "Provides visibility into throughput, latency, and delivery success rates."
)

doc.add_heading("6.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Publish Rate (msg/s)", "Timeseries", "rate(messages.publish.count)", "Cluster-wide publish throughput"],
        ["Deliver Rate (msg/s)", "Timeseries", "rate(messages.deliver_get.count)", "Cluster-wide delivery throughput"],
        ["Acknowledge Rate", "Timeseries", "rate(messages.ack.count)", "Consumer ack rate"],
        ["Publish vs Deliver", "Timeseries overlay", "publish_rate vs deliver_rate", "Gap = accumulation"],
        ["Publish/Deliver Ratio", "Query Value", "publish_rate / deliver_rate", "~1.0 target; >1 = growing backlog"],
        ["Published (24h)", "Query Value", "sum over 24h publish count", "Daily volume metric"],
        ["Delivered (24h)", "Query Value", "sum over 24h deliver count", "Daily delivery metric"],
        ["Redelivery Rate", "Timeseries", "rate(messages.redeliver.count)", "Messages redelivered (nack/reject)"],
        ["Unroutable Messages", "Timeseries", "messages.return_unroutable.count", "No matching queue"],
        ["Per-Queue Publish Rate", "Top List", "rate(publish) by {queue}", "Hotspot identification"],
        ["Per-Queue Deliver Rate", "Top List", "rate(deliver) by {queue}", "Consumer throughput per queue"],
        ["Message Size", "Histogram", "Custom APM metric", "Payload size patterns"],
    ]
)

doc.add_heading("6.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Publish Rate Drop", "publish.rate", "< 50% of baseline for > 5m", "P1 - Critical", "Check publisher health"],
        ["Delivery Stalled", "deliver.rate", "== 0 for > 5m with ready > 0", "P1 - Critical", "Consumer failure investigation"],
        ["High Redelivery", "redeliver / deliver", "> 10%", "P2 - Warning", "Consumer error handling issue"],
        ["Unroutable Messages", "return_unroutable", "> 0 for > 5m", "P2 - Warning", "Check exchange bindings"],
        ["Publish/Deliver Gap", "publish - deliver rate", "> 1000 msg/s for > 10m", "P2 - Warning", "Backlog building"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. DASHBOARD #4 — FAILED / DEAD-LETTERED MESSAGES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Dashboard #4 — Failed / Dead-Lettered Messages", level=1)
doc.add_paragraph(
    "Monitors message failures: dead-lettered, rejected, expired (TTL), and poison messages. "
    "Critical for detecting data loss risks."
)

doc.add_heading("7.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["DLQ Depth", "Timeseries", "queue.messages{queue:*dlq*,*dead*}", "Messages in DLQs"],
        ["DLQ Growth Rate", "Timeseries", "diff(dlq_messages)", "Failure accumulation rate"],
        ["Total DLQ Messages", "Query Value", "sum of all DLQ depths", "Total failed count"],
        ["Rejected Rate", "Timeseries", "rate(messages.reject.count)", "Consumer rejections/sec"],
        ["Nacked Rate", "Timeseries", "rate(messages.nack.count)", "Negative acks/sec"],
        ["Expired (TTL)", "Timeseries", "messages dropped by TTL", "Expired before consumption"],
        ["Max-Length Dropped", "Timeseries", "messages dropped by max-length", "Overflow drops"],
        ["DLQ by Source", "Top List", "DLQ by x-death header", "Which queues produce failures"],
        ["Failure Rate %", "Query Value", "(rejected + expired) / published * 100", "Overall failure %"],
        ["Oldest DLQ Message", "Query Value", "Custom: age of oldest DLQ msg", "Staleness indicator"],
        ["DLQ Consumer Status", "Status Widget", "Are DLQ processors running?", "DLQ processing health"],
        ["Poison Message Alert", "Event Stream", "Same msg redelivered > 5x", "Infinite retry detection"],
    ]
)

doc.add_heading("7.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["DLQ Depth Critical", "DLQ messages", "> 1,000 for > 10m", "P1 - Critical", "Investigate consumer failures"],
        ["DLQ Growing Fast", "DLQ growth rate", "> 100 msg/min for > 5m", "P1 - Critical", "Immediate investigation"],
        ["High Rejection Rate", "reject / deliver", "> 5%", "P2 - Warning", "Check message format/schema"],
        ["TTL Expiry Spike", "expired messages", "Anomaly detection", "P2 - Warning", "Consumers too slow or down"],
        ["DLQ Not Processed", "DLQ growing + consumer == 0", "> 30m", "P2 - Warning", "Start DLQ processor"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. DASHBOARD #5 — CONSUMER & PUBLISHER PERFORMANCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Dashboard #5 — Consumer & Publisher Performance", level=1)
doc.add_paragraph(
    "Performance characteristics of message producers and consumers: throughput, "
    "prefetch utilization, and consumer lag."
)

doc.add_heading("8.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Active Consumers", "Query Value", "sum:rabbitmq.queue.consumers{*}", "Total consumer count"],
        ["Consumer Utilization %", "Gauge per queue", "rabbitmq.queue.consumer_utilisation", "% time busy (target >90%)"],
        ["Prefetch Analysis", "Top List", "channel_prefetch_count by channel", "Prefetch optimization"],
        ["Consumer Lag", "Timeseries", "messages_ready by {queue}", "Backlog per queue"],
        ["Publish Confirm Latency", "Timeseries P50/P95/P99", "Custom APM metric", "Publisher confirm RTT"],
        ["Publisher Channels", "Timeseries", "channels with publish activity", "Active publisher count"],
        ["Per-Consumer Throughput", "Top List", "deliver_rate / consumers per queue", "msg/consumer/sec"],
        ["Slow Consumers", "Top List", "ready growing + consumers > 0", "Falling behind"],
        ["Consumer Cancellations", "Event Stream", "consumer.cancel events", "Unexpected disconnects"],
        ["Channel Errors", "Timeseries", "channel.error events", "Protocol errors"],
        ["Flow Control Active", "Status Widget", "rabbitmq.node.connection_flow", "Connections throttled"],
    ]
)

doc.add_heading("8.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Consumer Utilization Low", "consumer_utilisation", "< 50% with depth > 1000", "P2 - Warning", "Increase prefetch or add consumers"],
        ["Consumer Count Drop", "consumer count", "Drops > 50% in 5m", "P1 - Critical", "Consumer service likely crashed"],
        ["Flow Control Active", "connection_flow", "> 0 for > 2m", "P2 - Warning", "Publishers being throttled"],
        ["Publish Confirm Slow", "confirm latency P99", "> 500ms for > 5m", "P2 - Warning", "Cluster under pressure"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. DASHBOARD #6 — NODE RESOURCE UTILIZATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Dashboard #6 — Node Resource Utilization", level=1)
doc.add_paragraph(
    "Infrastructure-level monitoring for each of the 3 RabbitMQ nodes: CPU, memory, "
    "disk I/O, network, and Erlang VM internals."
)

doc.add_heading("9.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["CPU per Node", "Timeseries (3 lines)", "system.cpu.user by {host}", "CPU per node"],
        ["Memory per Node", "Timeseries (3 lines)", "system.mem.used / total by {host}", "Host memory"],
        ["RabbitMQ Memory", "Timeseries (3 lines)", "rabbitmq.node.mem_used by {node}", "RMQ-specific memory"],
        ["Memory Breakdown", "Stacked Area", "connections, queues, mnesia, binaries", "Memory consumers"],
        ["Disk IOPS", "Timeseries", "system.io.r_s, system.io.w_s by {host}", "Disk I/O patterns"],
        ["Disk Usage %", "Gauge per node", "system.disk.in_use by {host}", "Disk capacity"],
        ["Disk Free", "Timeseries", "rabbitmq.node.disk_free by {node}", "Disk free metric"],
        ["Network In/Out", "Timeseries", "system.net.bytes_rcvd/sent by {host}", "Network throughput"],
        ["Erlang VM Memory", "Stacked Area", "process, ets, atom, binary, code", "Erlang breakdown"],
        ["Erlang Processes", "Timeseries", "rabbitmq.node.proc_used by {node}", "Process count"],
        ["GC Rate", "Timeseries", "erlang gc rate", "GC frequency"],
        ["Beam CPU %", "Timeseries", "beam.smp process CPU by {host}", "Erlang VM CPU"],
        ["File Descriptors", "Gauge per node", "rabbitmq.node.fd_used by {node}", "FD consumption"],
        ["Socket Descriptors", "Gauge per node", "rabbitmq.node.sockets_used by {node}", "Socket consumption"],
    ]
)

doc.add_heading("9.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["CPU High", "system.cpu.user", "> 85% for > 10m", "P2 - Warning", "Check hot queues, GC pressure"],
        ["Memory High", "mem_used_pct", "> 80% of watermark", "P1 - Critical", "Risk of memory alarm"],
        ["Disk Space Low", "disk_free", "< 2x disk_free_limit", "P1 - Critical", "Risk of disk alarm"],
        ["High GC Rate", "erlang.gc.rate", "> baseline + 3 sigma", "P3 - Info", "Possible memory pressure"],
        ["Network Saturation", "system.net.bytes", "> 80% NIC capacity", "P2 - Warning", "Check inter-node traffic"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. DASHBOARD #7 — CONNECTION & CHANNEL MONITORING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Dashboard #7 — Connection & Channel Monitoring", level=1)
doc.add_paragraph(
    "Connection and channel lifecycle, distribution across nodes, detection of "
    "connection leaks or storms."
)

doc.add_heading("10.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Total Connections", "Timeseries", "rabbitmq.connections by {node}", "Per-node connections"],
        ["Total Channels", "Timeseries", "rabbitmq.channels by {node}", "Per-node channels"],
        ["Connection Rate", "Timeseries", "rate(connections created)", "New connections/sec"],
        ["Connection Close Rate", "Timeseries", "rate(connections closed)", "Closures/sec"],
        ["Connection Churn", "Timeseries", "created - closed rate", "Net connection change"],
        ["Connections by Client", "Top List", "grouped by client_properties", "Top consumers of connections"],
        ["Connections by Protocol", "Pie Chart", "AMQP 0-9-1 / 1.0 / MQTT / STOMP", "Protocol distribution"],
        ["Channels per Connection", "Histogram", "channels / connections", "Channel-heavy connections"],
        ["Blocked Connections", "Query Value", "rabbitmq.connections.blocked", "Flow-controlled connections"],
        ["Connection Errors", "Event Stream", "connection.error from logs", "Auth failures, protocol errors"],
        ["TLS Connections", "Query Value", "connections using TLS", "Security compliance"],
        ["Long-Running Connections", "Top List", "connections by age", "Stale connections"],
    ]
)

doc.add_heading("10.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Connection Storm", "connection creation rate", "> 100/s for > 1m", "P1 - Critical", "Possible reconnect loop"],
        ["Connection Limit Near", "connections / max", "> 80%", "P2 - Warning", "Increase limit or find leaks"],
        ["High Channel Count", "channels / connection", "> 50 avg", "P2 - Warning", "App misusing channels"],
        ["Blocked Connections", "blocked connections", "> 0 for > 2m", "P2 - Warning", "Resource alarm active"],
        ["Auth Failures Spike", "connection.error (auth)", "> 10/min", "P2 - Warning", "Check credentials"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. DASHBOARD #8 — USER & PERMISSION AUDIT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Dashboard #8 — User & Permission Audit", level=1)
doc.add_paragraph(
    "Tracks RabbitMQ user accounts, permissions, and authentication activity. "
    "Essential for security compliance and access governance."
)

doc.add_heading("11.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Total Users", "Query Value", "Custom: rabbitmqctl list_users count", "User count"],
        ["Users by Role", "Pie Chart", "administrator, monitoring, management, none", "Role distribution"],
        ["Admin User Count", "Query Value", "Users with administrator tag", "Admin tracking"],
        ["Permission Matrix", "Table Widget", "Custom: user + vhost + permissions", "Full permission view"],
        ["Auth Success Rate", "Timeseries", "success / (success + failure)", "Auth health"],
        ["Failed Auth Attempts", "Timeseries", "auth.failures by {username, source_ip}", "Security monitoring"],
        ["Auth Failures by IP", "Top List", "failures by source IP", "Brute force detection"],
        ["Permission Changes", "Event Stream", "Log: set_permissions, add/delete user", "Audit trail"],
        ["User Activity Map", "Heatmap", "Connections per user over time", "Usage patterns"],
        ["Inactive Users (>30d)", "Table", "No connections in 30 days", "Deactivation candidates"],
        ["Service vs Human", "Pie Chart", "Account classification", "Account type distribution"],
        ["VHost Access Matrix", "Table", "User-to-vhost mapping", "Access coverage"],
    ]
)

doc.add_heading("11.2 Custom Check Script", level=2)
doc.add_paragraph("Since Datadog doesn't natively collect user/permission data, deploy a custom check:")
add_code("""# /etc/datadog-agent/checks.d/rabbitmq_users.py
import subprocess, json
from datadog_checks.base import AgentCheck

class RabbitMQUserCheck(AgentCheck):
    def check(self, instance):
        result = subprocess.run(
            ['rabbitmqctl', 'list_users', '--formatter', 'json'],
            capture_output=True, text=True
        )
        users = json.loads(result.stdout)
        self.gauge('rabbitmq.users.total', len(users))
        admin_count = sum(1 for u in users
                         if 'administrator' in u.get('tags', []))
        self.gauge('rabbitmq.users.admin_count', admin_count)""")

doc.add_heading("11.3 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["New Admin User", "users.admin_count", "Increases from baseline", "P1 - Critical", "Verify authorization"],
        ["Auth Failure Spike", "auth.failures", "> 20/min", "P1 - Critical", "Possible credential stuffing"],
        ["Permission Change", "Event: set_permissions", "Any occurrence", "P3 - Info", "Audit notification"],
        ["Guest Active Remotely", "guest connections", "> 0 from non-localhost", "P1 - Critical", "Security vuln — disable guest"],
        ["Inactive User", "last connection", "> 90 days", "P3 - Info", "Review and remove"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. DASHBOARD #9 — RESTART & AVAILABILITY TRACKING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Dashboard #9 — Restart & Availability Tracking", level=1)
doc.add_paragraph(
    "Tracks node restarts, uptime, availability SLA, and correlates restarts with incidents."
)

doc.add_heading("12.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Uptime per Node", "Query Value (3x)", "rabbitmq.node.uptime by {node}", "Current uptime"],
        ["Uptime Timeline", "Timeseries", "uptime (reset = restart)", "Visual restart detection"],
        ["Restart Events", "Event Timeline", "Uptime resets + log 'RabbitMQ started'", "Restart timestamps"],
        ["Restart Frequency", "Bar Chart", "Restarts per node per week/month", "Trend analysis"],
        ["Restart Reason", "Table", "Shutdown reason from logs", "Why did it restart?"],
        ["MTBF per Node", "Query Value", "Mean time between restarts", "Reliability metric"],
        ["Availability % (30d)", "Query Value", "(total - downtime) / total * 100", "SLA compliance"],
        ["Availability Timeline", "Heatmap", "Per-node per hour (green/red)", "Visual availability"],
        ["Cluster Availability", "Query Value", "Time >= 2/3 nodes UP", "Cluster-level SLA"],
        ["Rolling Restart Status", "Status Widget", "Which node restarting", "Maintenance visibility"],
        ["Restart Impact", "Timeseries overlay", "Msg rate + connections around restart", "Impact correlation"],
        ["Restart Annotation", "Event Overlay", "Markers on all dashboards", "Cross-dashboard correlation"],
    ]
)

doc.add_heading("12.2 Restart Detection via Logs", level=2)
add_code("""# Datadog Log Pipeline Configuration
# Source: /var/log/rabbitmq/rabbit@<node>.log

# Restart Detection Pattern
grok_parser:
  pattern: "Starting RabbitMQ %{DATA:version} on Erlang %{DATA:erlang_version}"
  metric: rabbitmq.node.restart_detected
  tags: [node, version]

# Shutdown Detection Pattern
grok_parser:
  pattern: "RabbitMQ is asked to stop"
  metric: rabbitmq.node.shutdown_detected
  tags: [node, reason]""")

doc.add_heading("12.3 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Unplanned Restart", "uptime reset outside window", "Any occurrence", "P1 - Critical", "Immediate RCA"],
        ["Frequent Restarts", "restart count", "> 3/node/week", "P2 - Warning", "Root cause investigation"],
        ["SLA Breach Risk", "availability %", "< 99.9% (30d rolling)", "P1 - Critical", "SLA recovery plan"],
        ["Node Down > 5min", "node.running == 0", "> 5 minutes", "P1 - Critical", "Manual intervention"],
        ["Full Cluster Restart", "All 3 nodes uptime < 1h", "Simultaneous", "P1 - Critical", "Verify data integrity"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. DASHBOARD #10 — NETWORK PARTITION & SPLIT-BRAIN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Dashboard #10 — Network Partition & Split-Brain", level=1)
doc.add_paragraph(
    "Monitors network partitions — the most dangerous failure mode in RabbitMQ. "
    "Split-brain can cause data loss and message duplication."
)

doc.add_heading("13.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Partition Status", "Status (Green/Red)", "rabbitmq.node.partitions by {node}", "Active partitions?"],
        ["Partition History", "Event Timeline", "Partition events from logs", "Historical occurrences"],
        ["Inter-Node Latency", "Timeseries (3 pairs)", "Ping latency between pairs", "Network health"],
        ["Inter-Node Throughput", "Timeseries", "Erlang distribution port traffic", "Inter-node volume"],
        ["Cluster Membership", "Table", "rabbitmqctl cluster_status", "Peer visibility"],
        ["Mnesia Table Status", "Status Widget", "All tables synchronized?", "DB consistency"],
        ["Queue Leader Distribution", "Pie Chart", "Leaders per node", "Balance check"],
        ["Queue Replica Status", "Table", "Queues with insufficient replicas", "Replication health"],
        ["Partition Recovery Mode", "Status Widget", "autoheal / pause_minority / ignore", "Current strategy"],
    ]
)

doc.add_heading("13.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Partition Detected", "node.partitions", "> 0", "P0 - EMERGENCY", "Immediate intervention, follow runbook"],
        ["Inter-Node Latency High", "node-to-node ping", "> 50ms", "P2 - Warning", "Network investigation"],
        ["Under-Replicated Queue", "replicas < target", "Any queue", "P2 - Warning", "Node down or partition starting"],
        ["Mnesia Sync Failure", "table not synced", "> 5m", "P1 - Critical", "DB inconsistency risk"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. DASHBOARD #11 — QUORUM QUEUE HEALTH
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Dashboard #11 — Quorum Queue Health", level=1)
doc.add_paragraph(
    "Dedicated monitoring for Quorum Queues (Raft-based), the recommended queue type in "
    "RabbitMQ 4.x. Monitors Raft consensus, log replication, and leader elections."
)

doc.add_heading("14.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Quorum Queue Count", "Query Value", "count queue_type=quorum", "Total quorum queues"],
        ["Leader Distribution", "Pie Chart", "Leaders per node", "Balance (should be ~equal)"],
        ["Leader Elections", "Timeseries", "Raft leader election events", "Election frequency"],
        ["Raft Log Size", "Timeseries", "Raft log entries per queue", "Log compaction health"],
        ["Raft Commit Rate", "Timeseries", "rate of raft commits", "Consensus throughput"],
        ["Follower Lag", "Timeseries", "leader - follower commit index", "Replication lag"],
        ["Raft Term Changes", "Event Timeline", "Raft term increments", "Leadership changes"],
        ["WAL Segment Count", "Timeseries per node", "Raft WAL segments", "WAL growth"],
        ["WAL Disk Usage", "Gauge per node", "Raft WAL directory size", "Disk impact"],
        ["Quorum Queue Memory", "Top List", "Memory per quorum queue", "Memory impact"],
        ["In-Memory Messages", "Timeseries", "In-memory message count", "Memory pressure"],
        ["Snapshot Frequency", "Timeseries", "Raft snapshot operations", "Compaction activity"],
    ]
)

doc.add_heading("14.2 Alerts", level=2)
add_table(
    ["Alert Name", "Metric", "Condition", "Severity", "Action"],
    [
        ["Frequent Elections", "election count", "> 5/min", "P2 - Warning", "Cluster instability, check network"],
        ["Follower Lag High", "follower lag", "> 10,000 entries", "P2 - Warning", "Slow follower, check I/O"],
        ["WAL Disk Growing", "WAL dir size", "> 5GB/node", "P2 - Warning", "Check snapshot/compaction"],
        ["Quorum Lost", "replicas < majority", "Any queue", "P1 - Critical", "Queue unavailable, recover node"],
        ["Leader Imbalance", "leader count variance", "> 30% skew", "P3 - Info", "Rebalance leaders"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 15. DASHBOARD #12 — EXECUTIVE / SLA SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Dashboard #12 — Executive / SLA Summary", level=1)
doc.add_paragraph(
    "High-level dashboard for management and stakeholders. SLA compliance, key business "
    "metrics, and overall health without technical details."
)

doc.add_heading("15.1 Dashboard Layout", level=2)
add_table(
    ["Widget", "Type", "Metric / Query", "Purpose"],
    [
        ["Cluster Health", "Traffic Light", "Composite health score", "At-a-glance status"],
        ["Availability SLA %", "Query Value (large)", "Uptime % (target: 99.95%)", "SLA compliance"],
        ["Messages Today", "Query Value", "Total published + delivered", "Business volume"],
        ["Messages 7d Trend", "Timeseries", "Daily volumes", "Volume trend"],
        ["Failure Rate %", "Query Value", "Failed / Total * 100", "Error rate KPI"],
        ["Avg Delivery Latency", "Query Value", "Publish-to-consume P50", "Performance KPI"],
        ["P99 Delivery Latency", "Query Value", "99th percentile", "Tail latency KPI"],
        ["Incidents This Month", "Query Value", "P1/P2 alerts count", "Operational health"],
        ["MTTR", "Query Value", "Mean time to recover", "Operational efficiency"],
        ["Capacity Utilization", "Gauge", "Load / Max capacity %", "Capacity planning"],
        ["Cost per Million Msgs", "Query Value", "Infra cost / volume", "Cost efficiency"],
        ["Upcoming Maintenance", "Note Widget", "Next maintenance window", "Planning visibility"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 16. ALERTING STRATEGY & ESCALATION MATRIX
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Alerting Strategy & Escalation Matrix", level=1)

doc.add_heading("16.1 Alert Severity Levels", level=2)
add_table(
    ["Severity", "Response Time", "Notification", "Escalation", "Examples"],
    [
        ["P0 - Emergency", "< 5 min", "PagerDuty + Phone + Slack", "VP Eng in 15 min", "Partition, full cluster down"],
        ["P1 - Critical", "< 15 min", "PagerDuty + Slack", "Lead in 30 min", "Node down, alarms, SLA breach"],
        ["P2 - Warning", "< 1 hour", "Slack + Email", "Next standup", "Queue depth, FD usage, slow consumers"],
        ["P3 - Info", "Next business day", "Slack", "Weekly review", "Anomalies, trends, cleanup"],
    ]
)

doc.add_heading("16.2 Alert Routing", level=2)
add_table(
    ["Category", "Primary Team", "Secondary", "Notification Handle"],
    [
        ["Cluster / Nodes", "Platform SRE", "DevOps", "@pagerduty-rabbitmq-cluster"],
        ["Queue / Messages", "Application Team", "Platform SRE", "@slack-rabbitmq-alerts"],
        ["Consumer / Publisher", "Application Team", "Platform SRE", "@slack-app-team"],
        ["Infrastructure", "Infra Team", "Platform SRE", "@pagerduty-infra"],
        ["Security / Auth", "Security Team", "Platform SRE", "@pagerduty-security"],
        ["Network Partition", "SRE + Network", "VP Engineering", "@pagerduty-rabbitmq-emergency"],
    ]
)

doc.add_heading("16.3 Alert Best Practices", level=2)
for practice in [
    "Use composite monitors to reduce noise (e.g., queue depth high AND consumer count low)",
    "Set different thresholds for business hours vs. off-hours",
    "Use anomaly detection for variable baselines (publish rate, connections)",
    "Include runbook links in every alert notification body",
    "Auto-resolve alerts that recover within evaluation window",
    "Use Datadog Downtime to suppress alerts during maintenance",
    "Review and tune thresholds monthly based on false-positive rate",
    "Tag all monitors with team, service, and environment",
]:
    doc.add_paragraph(practice, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 17. TAGGING & NAMING CONVENTIONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("17. Tagging & Naming Conventions", level=1)

doc.add_heading("17.1 Mandatory Tags", level=2)
add_table(
    ["Tag Key", "Example Values", "Purpose"],
    [
        ["env", "production, staging, qa", "Environment identification"],
        ["cluster", "rabbitmq-prod, rabbitmq-staging", "Cluster identification"],
        ["node", "rabbitmq-node-1, -2, -3", "Individual node"],
        ["service", "rabbitmq", "Service catalog mapping"],
        ["team", "platform-sre, app-team-orders", "Ownership"],
        ["vhost", "/, orders, payments", "Virtual host"],
        ["queue", "orders.created, payments.process", "Queue name"],
        ["queue_type", "quorum, classic", "Queue type filtering"],
    ]
)

doc.add_heading("17.2 Dashboard Naming", level=2)
doc.add_paragraph("Pattern: [Environment] RabbitMQ - [Purpose] - [Scope]")
for ex in [
    "PROD RabbitMQ - Cluster Health Overview - All Nodes",
    "PROD RabbitMQ - Queue Monitoring - All VHosts",
    "PROD RabbitMQ - Message Flow & Delivery - Global",
    "PROD RabbitMQ - Executive SLA Summary - Monthly",
]:
    doc.add_paragraph(ex, style='List Bullet')

doc.add_heading("17.3 Monitor Naming", level=2)
doc.add_paragraph("Pattern: [Severity] RabbitMQ [Cluster] - [What] - [Where]")
for ex in [
    "[P1] RabbitMQ Prod - Node Down - rabbitmq-node-2",
    "[P2] RabbitMQ Prod - Queue Depth > 100K - orders.created",
    "[P0] RabbitMQ Prod - Network Partition - All Nodes",
]:
    doc.add_paragraph(ex, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 18. RUNBOOK REFERENCES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("18. Runbook References", level=1)
doc.add_paragraph("Every alert must link to a runbook:")

add_table(
    ["Runbook", "Trigger", "Key Steps"],
    [
        ["RB-001: Node Down", "node.running == 0", "Check host > Check Erlang > Restart > Cluster status > Rejoin"],
        ["RB-002: Memory Alarm", "mem_alarm == 1", "Memory breakdown > Large queues > Purge if safe > Increase watermark > Fix root cause"],
        ["RB-003: Disk Alarm", "disk_alarm == 1", "Disk usage > Clear logs > Remove temp > Expand disk"],
        ["RB-004: Partition", "partitions > 0", "Identify nodes > Check network > Follow partition strategy > Verify consistency"],
        ["RB-005: Queue Depth", "messages > 100K", "Consumer health > Scale consumers > Check poison msgs > Purge (with approval)"],
        ["RB-006: Consumer Failure", "consumers == 0", "Service health > Check logs > Restart service > Scale"],
        ["RB-007: Connection Storm", "connections spike", "Identify source > Check reconnect loop > Rate limit > Fix client"],
        ["RB-008: Full Cluster Restart", "All nodes restarted", "Verify quorum > Check sync > Verify integrity > Health checks > Notify"],
        ["RB-009: Security Incident", "Auth failures spike", "Identify IPs > Block > Rotate creds > Audit logs"],
        ["RB-010: SLA Breach", "Availability < 99.95%", "RCA > Incident report > Preventive measures > Update capacity plan"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 19. APPENDIX — FULL METRIC REFERENCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("19. Appendix — Full Metric Reference", level=1)

doc.add_heading("19.1 Datadog RabbitMQ Integration Metrics", level=2)
add_table(
    ["Metric Name", "Type", "Description"],
    [
        ["rabbitmq.node.running", "Gauge", "1 if running, 0 otherwise"],
        ["rabbitmq.node.mem_used", "Gauge", "Memory used (bytes)"],
        ["rabbitmq.node.mem_limit", "Gauge", "Memory watermark (bytes)"],
        ["rabbitmq.node.mem_alarm", "Gauge", "1 if memory alarm active"],
        ["rabbitmq.node.disk_free", "Gauge", "Free disk space (bytes)"],
        ["rabbitmq.node.disk_alarm", "Gauge", "1 if disk alarm active"],
        ["rabbitmq.node.fd_used", "Gauge", "File descriptors in use"],
        ["rabbitmq.node.fd_total", "Gauge", "File descriptors available"],
        ["rabbitmq.node.sockets_used", "Gauge", "Sockets in use"],
        ["rabbitmq.node.proc_used", "Gauge", "Erlang processes in use"],
        ["rabbitmq.node.proc_total", "Gauge", "Erlang processes available"],
        ["rabbitmq.node.uptime", "Gauge", "Uptime in milliseconds"],
        ["rabbitmq.node.partitions", "Gauge", "Network partitions seen"],
        ["rabbitmq.connections", "Gauge", "Total connections"],
        ["rabbitmq.channels", "Gauge", "Total channels"],
        ["rabbitmq.queue.messages", "Gauge", "Total messages in queue"],
        ["rabbitmq.queue.messages_ready", "Gauge", "Messages ready for delivery"],
        ["rabbitmq.queue.messages_unacked", "Gauge", "Delivered but unacked"],
        ["rabbitmq.queue.consumers", "Gauge", "Consumers on queue"],
        ["rabbitmq.queue.consumer_utilisation", "Gauge", "Consumer utilization (0-1)"],
        ["rabbitmq.queue.memory", "Gauge", "Memory used by queue"],
        ["rabbitmq.overview.messages.publish.count", "Counter", "Messages published"],
        ["rabbitmq.overview.messages.deliver_get.count", "Counter", "Messages delivered"],
        ["rabbitmq.overview.messages.ack.count", "Counter", "Messages acknowledged"],
        ["rabbitmq.overview.messages.redeliver.count", "Counter", "Messages redelivered"],
        ["rabbitmq.overview.messages.return_unroutable.count", "Counter", "Unroutable messages"],
    ]
)

doc.add_heading("19.2 Recommended Custom Metrics", level=2)
add_table(
    ["Custom Metric", "Source", "Purpose"],
    [
        ["rabbitmq.custom.restart_detected", "Log parsing", "Detect node restarts"],
        ["rabbitmq.custom.users.total", "rabbitmqctl", "Total user count"],
        ["rabbitmq.custom.users.admin_count", "rabbitmqctl", "Admin user count"],
        ["rabbitmq.custom.permissions.total", "rabbitmqctl", "Permission entries"],
        ["rabbitmq.custom.queue_count.total", "Management API", "Total queue count"],
        ["rabbitmq.custom.dlq.total_messages", "Management API", "Total DLQ messages"],
        ["rabbitmq.custom.synthetic.roundtrip_ms", "Synthetic check", "Publish-consume RTT"],
        ["rabbitmq.custom.quorum.leader_distribution", "Management API", "Leader distribution variance"],
    ]
)

# ── Footer ──
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RabbitMQ Datadog Monitoring Strategy - Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output = "/Users/tejasodanapalli/rabbitmq/RabbitMQ/docs/monitoring/RabbitMQ_Datadog_Monitoring_Dashboard_Strategy.docx"
doc.save(output)
print(f"Document saved: {output}")
