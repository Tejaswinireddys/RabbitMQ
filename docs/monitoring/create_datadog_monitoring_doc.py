#!/usr/bin/env python3
"""
Generate comprehensive RabbitMQ 3-Node Cluster Datadog Monitoring Dashboard Document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RabbitMQ 3-Node Cluster\nDatadog Monitoring & Dashboard Strategy')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Enterprise Monitoring Architecture & Dashboard Design Document')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    f'Version: 1.0\n'
    f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n'
    f'Classification: Internal / Infrastructure Team\n'
    f'Prepared by: Monitoring Architecture Team'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (placeholder)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Monitoring Architecture Overview',
    '3. Datadog Integration Setup',
    '4. Dashboard #1 — Cluster Health & Node Overview',
    '5. Dashboard #2 — Queue Monitoring & Analytics',
    '6. Dashboard #3 — Message Flow & Throughput',
    '7. Dashboard #4 — Message Delivery & Acknowledgement',
    '8. Dashboard #5 — Failed / Dead-Lettered Messages',
    '9. Dashboard #6 — Connection & Channel Monitoring',
    '10. Dashboard #7 — User & Permission Audit',
    '11. Dashboard #8 — Node Restart & Availability Tracking',
    '12. Dashboard #9 — Resource Utilization (Memory, Disk, CPU)',
    '13. Dashboard #10 — Exchange & Binding Analytics',
    '14. Dashboard #11 — Quorum Queue & Raft Consensus',
    '15. Dashboard #12 — TLS / Security Monitoring',
    '16. Dashboard #13 — Shovel & Federation Health',
    '17. Dashboard #14 — SLA & Executive Summary',
    '18. Alerting Strategy & Escalation Matrix',
    '19. Datadog Monitor Definitions',
    '20. Runbook References',
    '21. Appendix — Metric Reference Table',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document defines the comprehensive Datadog monitoring and dashboard strategy for our '
    'production RabbitMQ 3-node cluster. It is designed from the perspective of a Monitoring Architect '
    'and covers every critical dimension — from cluster health, queue depth, message throughput, '
    'delivery success/failure, user auditing, node restarts, resource utilization, and security posture.'
)
doc.add_paragraph(
    'The strategy encompasses 14 purpose-built dashboards, a tiered alerting framework with '
    'escalation paths, and integration with existing incident management workflows. Each dashboard '
    'section below includes the widgets, metrics, queries, thresholds, and recommended layout.'
)

doc.add_heading('1.1 Cluster Topology', level=2)
add_table_with_header(doc,
    ['Property', 'Value'],
    [
        ['Cluster Size', '3 Nodes (HA)'],
        ['RabbitMQ Version', '4.1.x'],
        ['Erlang/OTP Version', '26.x'],
        ['Queue Type', 'Quorum Queues (default) + Classic Mirrored'],
        ['Plugins Enabled', 'Management, Prometheus, Shovel, Federation, Consistent Hash Exchange, Top'],
        ['Monitoring Tool', 'Datadog Agent + RabbitMQ Integration'],
        ['Metric Source', 'Prometheus endpoint (/metrics) + Management API'],
        ['Environments', 'Production, Staging, QA'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  2. MONITORING ARCHITECTURE OVERVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Monitoring Architecture Overview', level=1)

doc.add_heading('2.1 Data Collection Architecture', level=2)
doc.add_paragraph(
    'Datadog collects RabbitMQ metrics through two complementary channels:'
)
add_bullet(doc, 'Datadog Agent RabbitMQ Integration — Polls the Management HTTP API (port 15672) every 15 seconds for queue, exchange, connection, channel, and node-level metrics.')
add_bullet(doc, 'Prometheus Endpoint Scraping — The rabbitmq_prometheus plugin exposes detailed Erlang VM and internal metrics on /metrics (port 15692). Datadog OpenMetrics check scrapes these.')
add_bullet(doc, 'Log Collection — Datadog Agent tails RabbitMQ logs for error patterns, restart events, and audit entries.')
add_bullet(doc, 'APM Traces — Application-side traces showing publish/consume latency.')
add_bullet(doc, 'Custom StatsD Metrics — Application-level counters for business-specific message tracking.')

doc.add_heading('2.2 Tagging Strategy', level=2)
doc.add_paragraph(
    'Consistent tagging is critical for filtering dashboards across the 3-node cluster:'
)
add_table_with_header(doc,
    ['Tag Key', 'Example Values', 'Purpose'],
    [
        ['env', 'production, staging, qa', 'Environment isolation'],
        ['rabbitmq_cluster', 'prod-rmq-cluster-01', 'Cluster identification'],
        ['rabbitmq_node', 'rabbit@node1, rabbit@node2, rabbit@node3', 'Per-node drill-down'],
        ['rabbitmq_vhost', '/, payments, notifications', 'Virtual host filtering'],
        ['rabbitmq_queue', 'order.processing, email.outbound', 'Queue-level filtering'],
        ['rabbitmq_queue_type', 'quorum, classic, stream', 'Queue type segmentation'],
        ['team', 'platform, payments, notifications', 'Team ownership'],
        ['service', 'order-service, email-service', 'Application-level correlation'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  3. DATADOG INTEGRATION SETUP
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Datadog Integration Setup', level=1)

doc.add_heading('3.1 Datadog Agent Configuration (Each Node)', level=2)
doc.add_paragraph('File: /etc/datadog-agent/conf.d/rabbitmq.d/conf.yaml')
config_text = (
    'init_config:\n\n'
    'instances:\n'
    '  - rabbitmq_api_url: http://localhost:15672/api/\n'
    '    rabbitmq_user: datadog_monitor\n'
    '    rabbitmq_pass: <VAULT_SECRET>\n'
    '    tag_families: true\n'
    '    collect_node_metrics: true\n'
    '    collect_overview_metrics: true\n'
    '    exchanges: [".*"]\n'
    '    exchanges_regexes: [".*"]\n'
    '    queues: [".*"]\n'
    '    queues_regexes: [".*"]\n'
    '    vhosts: [".*"]\n'
    '    tags:\n'
    '      - env:production\n'
    '      - rabbitmq_cluster:prod-rmq-cluster-01\n'
    '      - service:rabbitmq\n'
)
p = doc.add_paragraph()
run = p.add_run(config_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('3.2 Prometheus OpenMetrics Check', level=2)
doc.add_paragraph('File: /etc/datadog-agent/conf.d/openmetrics.d/conf.yaml')
prom_config = (
    'instances:\n'
    '  - prometheus_url: http://localhost:15692/metrics\n'
    '    namespace: rabbitmq\n'
    '    metrics:\n'
    '      - rabbitmq_queue_messages\n'
    '      - rabbitmq_queue_messages_ready\n'
    '      - rabbitmq_queue_messages_unacknowledged\n'
    '      - rabbitmq_queue_consumers\n'
    '      - rabbitmq_channel_messages_published_total\n'
    '      - rabbitmq_channel_messages_delivered_total\n'
    '      - rabbitmq_channel_messages_acknowledged_total\n'
    '      - rabbitmq_channel_messages_redelivered_total\n'
    '      - rabbitmq_connections\n'
    '      - rabbitmq_channels\n'
    '      - rabbitmq_node_mem_used\n'
    '      - rabbitmq_node_disk_free\n'
    '      - rabbitmq_erlang_processes_used\n'
    '      - rabbitmq_raft_term_total\n'
    '      - rabbitmq_raft_log_commit_index\n'
)
p = doc.add_paragraph()
run = p.add_run(prom_config)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('3.3 Log Collection', level=2)
doc.add_paragraph('File: /etc/datadog-agent/conf.d/rabbitmq.d/conf.yaml (logs section)')
log_config = (
    'logs:\n'
    '  - type: file\n'
    '    path: /var/log/rabbitmq/rabbit@*.log\n'
    '    service: rabbitmq\n'
    '    source: rabbitmq\n'
    '    tags:\n'
    '      - env:production\n'
    '    log_processing_rules:\n'
    '      - type: multi_line\n'
    '        name: rabbit_multiline\n'
    '        pattern: "^\\d{4}-\\d{2}-\\d{2}"\n'
)
p = doc.add_paragraph()
run = p.add_run(log_config)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  4. DASHBOARD #1 — CLUSTER HEALTH & NODE OVERVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Dashboard #1 — Cluster Health & Node Overview', level=1)
doc.add_paragraph(
    'Purpose: The single pane of glass for the RabbitMQ cluster. Shows overall cluster state, '
    'node availability, partition status, and Erlang VM health across all 3 nodes.'
)

doc.add_heading('4.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Cluster Status', 'Check Status', 'rabbitmq.aliveness (per node)', 'Green=All 3 up, Red=Any down'],
        ['Node 1 Status', 'Query Value', 'rabbitmq.node.running{rabbitmq_node:rabbit@node1}', '1=Running, 0=Down'],
        ['Node 2 Status', 'Query Value', 'rabbitmq.node.running{rabbitmq_node:rabbit@node2}', '1=Running, 0=Down'],
        ['Node 3 Status', 'Query Value', 'rabbitmq.node.running{rabbitmq_node:rabbit@node3}', '1=Running, 0=Down'],
        ['Network Partition Detected', 'Check Status', 'rabbitmq.node.partitions{*}', 'Alert if > 0'],
        ['Erlang Processes Used', 'Timeseries (grouped by node)', 'rabbitmq.node.proc_used / rabbitmq.node.proc_total * 100', 'Warn > 70%, Critical > 85%'],
        ['File Descriptors Used', 'Timeseries (grouped by node)', 'rabbitmq.node.fd_used / rabbitmq.node.fd_total * 100', 'Warn > 70%, Critical > 85%'],
        ['Socket Descriptors Used', 'Timeseries (grouped by node)', 'rabbitmq.node.sockets_used / rabbitmq.node.sockets_total * 100', 'Warn > 70%, Critical > 85%'],
        ['Uptime per Node', 'Query Value (3 widgets)', 'rabbitmq.node.uptime{rabbitmq_node:*}', 'Display in days/hours'],
        ['Erlang VM Memory', 'Timeseries (stacked by node)', 'rabbitmq.node.mem_used{rabbitmq_node:*}', 'Per-node memory trend'],
        ['GC Runs / sec', 'Timeseries', 'rate(rabbitmq_erlang_gc_runs_total)', 'Erlang GC pressure indicator'],
        ['Cluster Event Log', 'Log Stream', 'source:rabbitmq status:(error OR warning)', 'Filtered to cluster events'],
    ]
)

doc.add_heading('4.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Node Down', 'rabbitmq.aliveness == 0 for any node for 1 min', 'P1 — Critical', 'PagerDuty + Slack #rabbitmq-alerts'],
        ['Network Partition', 'rabbitmq.node.partitions > 0 for 30s', 'P1 — Critical', 'PagerDuty + Slack + Email'],
        ['Erlang Processes > 85%', 'proc_used/proc_total > 0.85 for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['File Descriptors > 80%', 'fd_used/fd_total > 0.80 for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  5. DASHBOARD #2 — QUEUE MONITORING & ANALYTICS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Dashboard #2 — Queue Monitoring & Analytics', level=1)
doc.add_paragraph(
    'Purpose: Deep visibility into all queues — count, depth, growth rate, consumer count, '
    'idle queues, and queues at risk of backing up.'
)

doc.add_heading('5.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Total Queue Count', 'Query Value', 'rabbitmq.overview.queues', 'Total across cluster'],
        ['Queues by VHost', 'Top List', 'rabbitmq.overview.queues grouped by rabbitmq_vhost', 'Breakdown per vhost'],
        ['Queues by Type', 'Pie Chart', 'count by rabbitmq_queue_type (quorum, classic, stream)', 'Quorum vs Classic distribution'],
        ['Queue Depth — Top 20', 'Top List', 'rabbitmq.queue.messages{*} top 20', 'Sorted descending by depth'],
        ['Queue Depth Trend (All)', 'Timeseries', 'rabbitmq.queue.messages{*} by {rabbitmq_queue}', 'Last 24h trend'],
        ['Ready Messages — Top 20', 'Top List', 'rabbitmq.queue.messages.ready{*} top 20', 'Messages waiting for consumers'],
        ['Unacked Messages — Top 20', 'Top List', 'rabbitmq.queue.messages.unacknowledged{*} top 20', 'Potential consumer issues'],
        ['Queue Growth Rate', 'Timeseries', 'rate(rabbitmq.queue.messages{*}) by {rabbitmq_queue}', 'Positive = growing, Negative = draining'],
        ['Consumers per Queue', 'Top List', 'rabbitmq.queue.consumers{*} by {rabbitmq_queue}', 'Warn if 0 consumers'],
        ['Queues with 0 Consumers', 'Query Value', 'count_nonzero(rabbitmq.queue.consumers{*} < 1)', 'Alert if persistent > 5 min'],
        ['Idle Queues (no activity 24h)', 'Top List', 'queues with 0 publish + 0 deliver for 24h', 'Cleanup candidates'],
        ['Queue Memory Usage', 'Top List', 'rabbitmq.queue.memory{*} top 20', 'Memory hogs'],
        ['Queue Messages RAM vs Disk', 'Stacked Bar', 'rabbitmq.queue.messages.ram + messages.persistent', 'RAM pressure indicator'],
        ['Queue Auto-Delete & Exclusive', 'Table', 'List of auto-delete/exclusive queues', 'Housekeeping view'],
    ]
)

doc.add_heading('5.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Queue Depth Spike', 'rabbitmq.queue.messages > 10,000 for any queue for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Queue Depth Critical', 'rabbitmq.queue.messages > 50,000 for any queue for 5 min', 'P1 — Critical', 'PagerDuty + Slack'],
        ['Queue with 0 Consumers', 'rabbitmq.queue.consumers == 0 for 5 min (non-DLQ)', 'P2 — Warning', 'Slack + Email to queue owner'],
        ['Queue Count Anomaly', 'Anomaly detection on rabbitmq.overview.queues', 'P3 — Info', 'Slack #rabbitmq-info'],
        ['Unacked Messages Growing', 'rate(rabbitmq.queue.messages.unacknowledged) > 100/s for 10 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  6. DASHBOARD #3 — MESSAGE FLOW & THROUGHPUT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Dashboard #3 — Message Flow & Throughput', level=1)
doc.add_paragraph(
    'Purpose: Real-time and historical view of message publishing, routing, and consumption rates '
    'across the entire cluster.'
)

doc.add_heading('6.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Messages Published/sec', 'Query Value', 'rate(rabbitmq.overview.messages.publish_in.rate)', 'Current publish rate'],
        ['Messages Delivered/sec', 'Query Value', 'rate(rabbitmq.overview.messages.deliver_get.rate)', 'Current consume rate'],
        ['Publish vs Deliver Trend', 'Timeseries (dual axis)', 'publish.rate vs deliver_get.rate', 'Gap = backlog building'],
        ['Publish Rate by Exchange', 'Top List', 'rate(rabbitmq.exchange.messages.publish_in.rate) by {exchange}', 'Hotspot detection'],
        ['Deliver Rate by Queue', 'Top List', 'rate(rabbitmq.queue.messages.deliver.rate) by {queue}', 'Consumer throughput'],
        ['Message Rate by VHost', 'Timeseries', 'publish + deliver rates grouped by vhost', 'VHost-level traffic'],
        ['Message Rate by Node', 'Timeseries (3 lines)', 'publish rate by {rabbitmq_node}', 'Load distribution across nodes'],
        ['Publish/Deliver Ratio', 'Query Value', 'publish.rate / deliver.rate', '> 1.0 means backlog growing'],
        ['Throughput Heatmap', 'Heatmap', 'rate(rabbitmq.overview.messages.publish_in.rate)', 'Hourly pattern analysis'],
        ['Message Size Distribution', 'Distribution', 'rabbitmq.queue.message_bytes / rabbitmq.queue.messages', 'Avg message size trend'],
        ['Routed vs Unroutable', 'Timeseries', 'publish_in.rate vs return_unroutable.rate', 'Routing failures'],
    ]
)

doc.add_heading('6.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Publish Rate Drop > 50%', 'Anomaly detection on publish rate, > 50% drop for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Deliver Rate Drop > 50%', 'Anomaly detection on deliver rate, > 50% drop for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Publish/Deliver Gap Widening', 'publish_rate - deliver_rate > 500/s for 10 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Unroutable Messages', 'return_unroutable.rate > 0 for 5 min', 'P3 — Info', 'Slack #rabbitmq-info'],
        ['Zero Throughput', 'publish.rate == 0 AND deliver.rate == 0 for 10 min (during business hours)', 'P1 — Critical', 'PagerDuty'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  7. DASHBOARD #4 — MESSAGE DELIVERY & ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Dashboard #4 — Message Delivery & Acknowledgement', level=1)
doc.add_paragraph(
    'Purpose: Track the full message lifecycle from delivery to consumer acknowledgment. '
    'Identifies slow consumers, redelivery storms, and acknowledgment failures.'
)

doc.add_heading('7.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Messages Delivered/sec', 'Query Value', 'rabbitmq.overview.messages.deliver_get.rate', 'Cluster-wide delivery rate'],
        ['Messages Acknowledged/sec', 'Query Value', 'rabbitmq.overview.messages.ack.rate', 'Successful processing rate'],
        ['Deliver vs Ack Trend', 'Timeseries (overlay)', 'deliver.rate vs ack.rate', 'Gap = processing backlog'],
        ['Acknowledgment Latency', 'Timeseries', 'Custom: time between deliver and ack (APM trace)', 'p50, p95, p99 latency'],
        ['Redelivered Messages/sec', 'Timeseries', 'rabbitmq.overview.messages.redeliver.rate', 'Consumer rejections/failures'],
        ['Redelivery Rate by Queue', 'Top List', 'rabbitmq.queue.messages.redeliver.rate by {queue}', 'Problematic queues'],
        ['Unacknowledged Count', 'Timeseries (by node)', 'rabbitmq.queue.messages.unacknowledged by {node}', 'Consumer stuck indicator'],
        ['Prefetch Utilization', 'Timeseries', 'unacked / prefetch_count per channel', 'Consumer capacity'],
        ['Consumer Utilization %', 'Top List', 'rabbitmq.queue.consumer_utilisation by {queue}', '< 100% = consumer idle sometimes'],
        ['Delivery Mode Split', 'Pie Chart', 'Persistent vs Transient messages', 'Durability distribution'],
        ['Nacked Messages/sec', 'Timeseries', 'rabbitmq.channel.messages.nacked.rate', 'Explicit consumer rejections'],
    ]
)

doc.add_heading('7.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['High Redelivery Rate', 'redeliver.rate > 50/s for any queue for 5 min', 'P2 — Warning', 'Slack + queue owner'],
        ['Unacked Growing Continuously', 'monotonic increase in unacked for 15 min', 'P1 — Critical', 'PagerDuty'],
        ['Consumer Utilization < 50%', 'consumer_utilisation < 0.5 while queue depth > 1000', 'P3 — Info', 'Slack #rabbitmq-info'],
        ['Ack Rate Drop', 'ack.rate drops > 80% from baseline for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  8. DASHBOARD #5 — FAILED / DEAD-LETTERED MESSAGES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Dashboard #5 — Failed / Dead-Lettered Messages', level=1)
doc.add_paragraph(
    'Purpose: Critical visibility into message failures — dead-letter queues, rejected messages, '
    'expired (TTL) messages, and poison messages requiring manual intervention.'
)

doc.add_heading('8.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Total DLQ Messages', 'Query Value', 'sum(rabbitmq.queue.messages{rabbitmq_queue:*.dlq OR rabbitmq_queue:*.dead-letter})', 'Total failed messages'],
        ['DLQ Depth Trend', 'Timeseries', 'rabbitmq.queue.messages{queue:*.dlq} by {queue}', 'Growing = unresolved failures'],
        ['DLQ Ingress Rate', 'Timeseries', 'rate(rabbitmq.queue.messages.publish_in{queue:*.dlq})', 'New failures per second'],
        ['Messages Rejected (basic.reject)', 'Timeseries', 'rate(rabbitmq.channel.messages.rejected.rate)', 'Consumer-side rejections'],
        ['Messages Expired (TTL)', 'Timeseries', 'Custom metric for TTL-expired messages', 'Messages timing out before consumption'],
        ['Messages Dropped (max-length)', 'Timeseries', 'Custom metric for overflow drops', 'Queue overflow events'],
        ['DLQ by Source Queue', 'Table', 'DLQ messages tagged with x-first-death-queue header', 'Root cause identification'],
        ['DLQ by Rejection Reason', 'Pie Chart', 'x-first-death-reason: rejected / expired / maxlen', 'Failure distribution'],
        ['Poison Message Candidates', 'Top List', 'Messages with x-death count > 3', 'Repeatedly failing messages'],
        ['DLQ Age — Oldest Message', 'Query Value', 'Max age of message in DLQ', 'Staleness indicator'],
        ['Failed Message Log', 'Log Stream', 'source:rabbitmq "dead letter" OR "rejected" OR "expired"', 'Real-time failure events'],
    ]
)

doc.add_heading('8.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['DLQ Depth Increasing', 'rate(DLQ messages) > 0 for 10 min', 'P2 — Warning', 'Slack + queue owner team'],
        ['DLQ Depth Critical', 'DLQ messages > 1,000', 'P1 — Critical', 'PagerDuty + Slack'],
        ['High Rejection Rate', 'rejected.rate > 100/s for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Poison Messages Detected', 'Any message with x-death count > 5', 'P1 — Critical', 'PagerDuty + Manual intervention'],
        ['TTL Expiry Spike', 'TTL expired messages > 50/s for 5 min', 'P2 — Warning', 'Slack + consumer team'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  9. DASHBOARD #6 — CONNECTION & CHANNEL MONITORING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Dashboard #6 — Connection & Channel Monitoring', level=1)
doc.add_paragraph(
    'Purpose: Monitor client connections and channels to detect connection leaks, channel storms, '
    'and client-side issues affecting cluster stability.'
)

doc.add_heading('9.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Total Connections', 'Query Value', 'rabbitmq.overview.connections', 'Cluster-wide count'],
        ['Total Channels', 'Query Value', 'rabbitmq.overview.channels', 'Cluster-wide count'],
        ['Connections by Node', 'Timeseries (3 lines)', 'rabbitmq.connections by {rabbitmq_node}', 'Load distribution'],
        ['Channels by Node', 'Timeseries (3 lines)', 'rabbitmq.channels by {rabbitmq_node}', 'Load distribution'],
        ['Connections by VHost', 'Top List', 'rabbitmq.connections by {rabbitmq_vhost}', 'VHost connection breakdown'],
        ['Connections by Client App', 'Top List', 'rabbitmq.connections by {client_properties.connection_name}', 'Identify heavy clients'],
        ['Connection Open Rate', 'Timeseries', 'rate(rabbitmq.connections.opened)', 'Churn detection'],
        ['Connection Close Rate', 'Timeseries', 'rate(rabbitmq.connections.closed)', 'Churn detection'],
        ['Connection Churn (Open - Close)', 'Timeseries', 'rate(opened) - rate(closed)', 'Leak detection: positive = leak'],
        ['Channels per Connection', 'Distribution', 'channels / connections', 'Avg and outlier detection'],
        ['Blocked Connections', 'Query Value', 'rabbitmq.connections.state:blocked', 'Memory alarm triggered'],
        ['Connection Errors', 'Log Stream', 'source:rabbitmq "connection_closed_abruptly" OR "handshake_timeout"', 'Client errors'],
    ]
)

doc.add_heading('9.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Connection Count Spike', 'Anomaly detection: connections > 2x normal', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Connection Leak Detected', 'Monotonic increase in connections for 30 min', 'P2 — Warning', 'Slack + app team'],
        ['Blocked Connections', 'blocked_connections > 0 for 1 min', 'P1 — Critical', 'PagerDuty'],
        ['Channel Count > 10,000', 'rabbitmq.overview.channels > 10000', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['High Connection Churn', 'connection open rate > 100/s for 5 min', 'P3 — Info', 'Slack #rabbitmq-info'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  10. DASHBOARD #7 — USER & PERMISSION AUDIT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Dashboard #7 — User & Permission Audit', level=1)
doc.add_paragraph(
    'Purpose: Track user accounts, their permissions, active sessions, and audit changes. '
    'Essential for security compliance and access governance.'
)

doc.add_heading('10.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Total Users', 'Query Value', 'Custom check: rabbitmq management API /api/users count', 'Baseline user count'],
        ['Users by Tag/Role', 'Table', 'API /api/users — name, tags (administrator, monitoring, management)', 'Role distribution'],
        ['User Permissions Matrix', 'Table', 'API /api/permissions — user, vhost, configure, write, read', 'Full permission audit view'],
        ['Active User Sessions', 'Table', 'Connections grouped by authenticated user', 'Who is connected now'],
        ['Connections per User', 'Top List', 'rabbitmq.connections grouped by user', 'Heavy user identification'],
        ['User Permission Changes', 'Log Stream/Event', 'source:rabbitmq "set_permissions" OR "set_user_tags" OR "add_user" OR "delete_user"', 'Audit trail'],
        ['Failed Authentication Attempts', 'Timeseries', 'Logs: "authentication failure" OR "ACCESS_REFUSED"', 'Security monitoring'],
        ['Admin User Activity', 'Log Stream', 'Management API audit log for administrator-tagged users', 'Privileged access audit'],
        ['Guest Account Status', 'Check Status', 'Custom check: guest account disabled', 'Security compliance'],
        ['Users Without Connections (Stale)', 'Table', 'Users in /api/users not seen in connections for 30 days', 'Cleanup candidates'],
        ['Permission Change Timeline', 'Event Timeline', 'Datadog events tagged rabbitmq:permission_change', 'Change history'],
    ]
)

doc.add_heading('10.2 Custom Datadog Check for User Monitoring', level=2)
doc.add_paragraph(
    'Since Datadog does not natively track RabbitMQ user/permission metrics, a custom Agent check is required:'
)
custom_check = (
    '# /etc/datadog-agent/checks.d/rabbitmq_users.py\n'
    'import requests\n'
    'from datadog_checks.base import AgentCheck\n\n'
    'class RabbitMQUserCheck(AgentCheck):\n'
    '    def check(self, instance):\n'
    '        url = instance.get("rabbitmq_api_url", "http://localhost:15672/api")\n'
    '        auth = (instance["user"], instance["pass"])\n\n'
    '        # Total users\n'
    '        users = requests.get(f"{url}/users", auth=auth).json()\n'
    '        self.gauge("rabbitmq.users.total", len(users))\n'
    '        admin_count = sum(1 for u in users if "administrator" in u.get("tags", ""))\n'
    '        self.gauge("rabbitmq.users.administrators", admin_count)\n\n'
    '        # Check guest account\n'
    '        guest_exists = any(u["name"] == "guest" for u in users)\n'
    '        self.gauge("rabbitmq.users.guest_enabled", 1 if guest_exists else 0)\n\n'
    '        # Permissions per vhost\n'
    '        perms = requests.get(f"{url}/permissions", auth=auth).json()\n'
    '        self.gauge("rabbitmq.permissions.total", len(perms))\n'
)
p = doc.add_paragraph()
run = p.add_run(custom_check)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_heading('10.3 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['New Admin User Created', 'Event: "set_user_tags" with "administrator"', 'P1 — Critical', 'Security team + PagerDuty'],
        ['Failed Auth > 10/min', 'Log count "authentication failure" > 10 in 1 min', 'P2 — Warning', 'Security team Slack'],
        ['Guest Account Enabled', 'rabbitmq.users.guest_enabled == 1', 'P2 — Warning', 'Security team Slack'],
        ['Permission Change Detected', 'Event: "set_permissions"', 'P3 — Info', 'Slack #rabbitmq-audit'],
        ['User Deleted', 'Event: "delete_user"', 'P2 — Warning', 'Slack #rabbitmq-audit'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  11. DASHBOARD #8 — NODE RESTART & AVAILABILITY TRACKING
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Dashboard #8 — Node Restart & Availability Tracking', level=1)
doc.add_paragraph(
    'Purpose: Track when nodes restart, how long they were down, rolling restart history, '
    'and measure overall cluster availability SLA.'
)

doc.add_heading('11.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Node 1 Uptime', 'Query Value', 'rabbitmq.node.uptime{node:rabbit@node1}', 'Current uptime in hours/days'],
        ['Node 2 Uptime', 'Query Value', 'rabbitmq.node.uptime{node:rabbit@node2}', 'Current uptime in hours/days'],
        ['Node 3 Uptime', 'Query Value', 'rabbitmq.node.uptime{node:rabbit@node3}', 'Current uptime in hours/days'],
        ['Restart Event Timeline', 'Event Timeline', 'events("rabbitmq node restart" OR uptime drop to 0)', 'When did each node restart?'],
        ['Uptime Trend (All Nodes)', 'Timeseries', 'rabbitmq.node.uptime by {rabbitmq_node}', 'Drops indicate restarts'],
        ['Restart Count (Last 30 Days)', 'Query Value (per node)', 'count of uptime resets per node', 'Restart frequency'],
        ['Mean Time Between Restarts', 'Query Value (per node)', 'Avg uptime between resets', 'Stability indicator'],
        ['Restart Correlation with Errors', 'Event Overlay on Timeseries', 'Overlay restart events on error log count', 'Root cause correlation'],
        ['Cluster Availability % (30 day)', 'Query Value', '(time all 3 nodes up / total time) * 100', 'SLA metric: target 99.99%'],
        ['Partial Availability %', 'Query Value', '(time >= 2 nodes up / total time) * 100', 'Degraded but functional'],
        ['Restart Impact — Queue Re-sync Time', 'Timeseries', 'Time for quorum queues to re-elect leader after restart', 'Recovery time indicator'],
        ['Restart Reason Analysis', 'Log Stream', 'source:rabbitmq "stopped" OR "starting" OR "Shutdown" OR "boot"', 'Why did it restart?'],
        ['Process Restart (systemd)', 'Event Timeline', 'source:systemd unit:rabbitmq-server', 'OS-level restart tracking'],
    ]
)

doc.add_heading('11.2 Custom Detection: Restart Events', level=2)
doc.add_paragraph(
    'Since RabbitMQ does not emit a "restart" metric directly, detect restarts by monitoring uptime drops:'
)
restart_logic = (
    '# Datadog Monitor: Detect RabbitMQ Node Restart\n'
    '# Metric: rabbitmq.node.uptime\n'
    '# Logic: If uptime drops by > 90% between two consecutive checks, a restart occurred\n\n'
    'Monitor Query:\n'
    '  pct_change(avg(last_5m), last_1m):avg:rabbitmq.node.uptime{*} by {rabbitmq_node} < -90\n\n'
    '# Alternative: Log-based detection\n'
    'Log Query:\n'
    '  source:rabbitmq ("Starting RabbitMQ" OR "RabbitMQ is asked to stop")\n'
)
p = doc.add_paragraph()
run = p.add_run(restart_logic)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('11.3 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Node Restarted', 'Uptime drop detected (pct_change < -90%)', 'P1 — Critical', 'PagerDuty + Slack + Email'],
        ['Multiple Restarts in 24h', '> 2 restart events for same node in 24h', 'P1 — Critical', 'PagerDuty + escalate to SRE lead'],
        ['Cluster Availability < 99.9%', '30-day rolling availability < 99.9%', 'P2 — Warning', 'Email to SRE manager'],
        ['Node Down > 5 Minutes', 'Node not running for > 5 min after restart', 'P1 — Critical', 'PagerDuty + auto-recovery trigger'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  12. DASHBOARD #9 — RESOURCE UTILIZATION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Dashboard #9 — Resource Utilization (Memory, Disk, CPU)', level=1)
doc.add_paragraph(
    'Purpose: Infrastructure-level monitoring of the 3 nodes. Ensures RabbitMQ has sufficient '
    'resources and detects memory/disk alarms before they trigger flow control.'
)

doc.add_heading('12.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Memory Used per Node', 'Timeseries (3 lines)', 'rabbitmq.node.mem_used by {rabbitmq_node}', 'Memory trend per node'],
        ['Memory Limit per Node', 'Timeseries (3 lines)', 'rabbitmq.node.mem_limit by {rabbitmq_node}', 'Configured high watermark'],
        ['Memory Usage %', 'Query Value (3 widgets)', 'mem_used / mem_limit * 100', 'Warn > 70%, Critical > 85%'],
        ['Memory Alarm Active', 'Check Status', 'rabbitmq.node.mem_alarm by {node}', 'Red = alarm triggered = flow control'],
        ['Disk Free per Node', 'Timeseries (3 lines)', 'rabbitmq.node.disk_free by {rabbitmq_node}', 'Disk space trend'],
        ['Disk Free Limit', 'Query Value (3 widgets)', 'rabbitmq.node.disk_free_limit', 'Configured low watermark'],
        ['Disk Alarm Active', 'Check Status', 'rabbitmq.node.disk_free_alarm by {node}', 'Red = alarm = publishers blocked'],
        ['CPU Usage per Node', 'Timeseries', 'system.cpu.user{host:rmq-node*} by {host}', 'OS-level CPU (Datadog Agent)'],
        ['Erlang Scheduler Utilization', 'Timeseries', 'rabbitmq_erlang_scheduler_utilization', 'Erlang-specific CPU usage'],
        ['IO Read/Write per Node', 'Timeseries', 'rabbitmq.node.io_read_bytes.rate + io_write_bytes.rate', 'Disk I/O pressure'],
        ['IO Read/Write Latency', 'Timeseries', 'rabbitmq.node.io_read_avg_time + io_write_avg_time', 'Disk latency'],
        ['Network Traffic per Node', 'Timeseries', 'system.net.bytes_sent + bytes_rcvd {host:rmq-node*}', 'Inter-node + client traffic'],
        ['GC Reclaimed Bytes/sec', 'Timeseries', 'rate(rabbitmq_erlang_gc_reclaimed_bytes_total)', 'Memory pressure indicator'],
        ['Mnesia Transaction Rate', 'Timeseries', 'rabbitmq.node.mnesia_transaction_rate', 'Metadata operations'],
    ]
)

doc.add_heading('12.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Memory Alarm Triggered', 'rabbitmq.node.mem_alarm == 1', 'P1 — Critical', 'PagerDuty + Slack (publishers blocked!)'],
        ['Disk Alarm Triggered', 'rabbitmq.node.disk_free_alarm == 1', 'P1 — Critical', 'PagerDuty + Slack (publishers blocked!)'],
        ['Memory > 80% of Limit', 'mem_used / mem_limit > 0.80 for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Disk Free < 5 GB', 'disk_free < 5GB for 5 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['CPU > 80% for 15 min', 'system.cpu.user > 80 for 15 min', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['IO Latency > 10ms', 'io_write_avg_time > 10ms for 10 min', 'P3 — Info', 'Slack #rabbitmq-info'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  13. DASHBOARD #10 — EXCHANGE & BINDING ANALYTICS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Dashboard #10 — Exchange & Binding Analytics', level=1)
doc.add_paragraph(
    'Purpose: Monitor exchange types, message routing patterns, binding counts, and detect '
    'misconfigured or unused exchanges.'
)

doc.add_heading('13.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Total Exchanges', 'Query Value', 'rabbitmq.overview.exchanges', 'Total count'],
        ['Exchanges by Type', 'Pie Chart', 'direct, fanout, topic, headers, consistent-hash', 'Type distribution'],
        ['Message Rate per Exchange', 'Top List', 'rabbitmq.exchange.messages.publish_in.rate by {exchange}', 'Hotspot detection'],
        ['Bindings per Exchange', 'Top List', 'Custom: binding count per exchange', 'Complexity indicator'],
        ['Unroutable Messages per Exchange', 'Timeseries', 'rabbitmq.exchange.messages.return_unroutable.rate by {exchange}', 'Routing failures'],
        ['Exchanges with Zero Traffic', 'Table', 'Exchanges with 0 publish_in for 24h', 'Cleanup candidates'],
        ['Exchange-to-Queue Binding Map', 'Table', 'Exchange -> Queue bindings with routing keys', 'Topology documentation'],
        ['Alternate Exchange Usage', 'Table', 'Exchanges configured with alternate-exchange', 'Dead letter routing audit'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  14. DASHBOARD #11 — QUORUM QUEUE & RAFT CONSENSUS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('14. Dashboard #11 — Quorum Queue & Raft Consensus Health', level=1)
doc.add_paragraph(
    'Purpose: Since the cluster uses quorum queues for HA, this dashboard monitors the Raft '
    'consensus protocol — leader elections, log replication, and member health.'
)

doc.add_heading('14.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Quorum Queue Count', 'Query Value', 'count of queues with type=quorum', 'Total quorum queues'],
        ['Raft Term Total', 'Timeseries (by queue)', 'rabbitmq_raft_term_total', 'Increment = leader election occurred'],
        ['Leader Elections/min', 'Timeseries', 'rate(rabbitmq_raft_term_total)', 'High rate = instability'],
        ['Raft Log Entries Committed', 'Timeseries', 'rate(rabbitmq_raft_log_commit_index)', 'Replication throughput'],
        ['Raft Log Entries Uncommitted', 'Timeseries', 'raft_log_last_written_index - raft_log_commit_index', 'Replication lag'],
        ['Raft Snapshot Count', 'Timeseries', 'rabbitmq_raft_log_snapshot_index', 'Snapshot frequency'],
        ['Quorum Queue Leader Distribution', 'Pie Chart', 'Leaders per node', 'Balanced = ~33% each node'],
        ['Under-Replicated Queues', 'Query Value', 'Queues where online members < configured members', 'Data safety risk'],
        ['Quorum Queue Memory Usage', 'Top List', 'Memory per quorum queue', 'Resource consumption'],
        ['Member Health per Queue', 'Table', 'Queue -> leader + followers + status', 'Detailed membership view'],
    ]
)

doc.add_heading('14.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['Frequent Leader Elections', 'raft_term rate > 5/min for any queue', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Under-Replicated Queue', 'online_members < total_members for > 5 min', 'P1 — Critical', 'PagerDuty'],
        ['Raft Log Lag > 10,000', 'uncommitted entries > 10,000', 'P2 — Warning', 'Slack #rabbitmq-alerts'],
        ['Unbalanced Leaders', 'Any node has > 50% of leaders', 'P3 — Info', 'Slack #rabbitmq-info'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  15. DASHBOARD #12 — TLS / SECURITY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('15. Dashboard #12 — TLS / Security Monitoring', level=1)
doc.add_paragraph(
    'Purpose: Monitor TLS certificate expiry, cipher suite usage, non-TLS connections, '
    'and security-relevant events.'
)

doc.add_heading('15.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['TLS Certificate Expiry (Days)', 'Query Value (per node)', 'Custom check: days until cert expiry', 'Warn < 30 days, Critical < 7 days'],
        ['TLS vs Non-TLS Connections', 'Pie Chart', 'Connections by protocol (amqp vs amqps)', 'Security compliance: 100% TLS target'],
        ['Non-TLS Connection List', 'Table', 'Connections on port 5672 (non-TLS)', 'Compliance violations'],
        ['TLS Protocol Version Distribution', 'Pie Chart', 'TLS 1.2 vs TLS 1.3 connections', 'TLS 1.3 migration tracking'],
        ['Cipher Suite Distribution', 'Top List', 'Connections grouped by cipher suite', 'Weak cipher detection'],
        ['Failed TLS Handshakes', 'Timeseries', 'Logs: "SSL" AND ("error" OR "handshake" OR "alert")', 'Client cert issues'],
        ['Inter-Node TLS Status', 'Check Status', 'Erlang distribution using TLS', 'Node-to-node encryption'],
        ['Management UI HTTPS Status', 'Check Status', 'Port 15671 (HTTPS management)', 'UI access security'],
    ]
)

doc.add_heading('15.2 Alert Definitions', level=2)
add_table_with_header(doc,
    ['Alert Name', 'Condition', 'Severity', 'Notification'],
    [
        ['TLS Certificate Expiring < 30 Days', 'cert_expiry_days < 30', 'P2 — Warning', 'Slack + Email to cert team'],
        ['TLS Certificate Expiring < 7 Days', 'cert_expiry_days < 7', 'P1 — Critical', 'PagerDuty + cert team'],
        ['Non-TLS Connections Detected', 'non_tls_connections > 0', 'P2 — Warning', 'Security team Slack'],
        ['TLS Handshake Failures > 10/min', 'Failed handshakes > 10 per minute', 'P3 — Info', 'Slack #rabbitmq-security'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  16. DASHBOARD #13 — SHOVEL & FEDERATION HEALTH
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('16. Dashboard #13 — Shovel & Federation Health', level=1)
doc.add_paragraph(
    'Purpose: Since the cluster has Shovel and Federation plugins enabled, monitor their '
    'operational status and message transfer health.'
)

doc.add_heading('16.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Shovel Status', 'Check Status', 'Custom: /api/shovels — state per shovel', 'Green=running, Red=terminated'],
        ['Shovel Message Rate', 'Timeseries', 'Messages transferred per shovel', 'Throughput per shovel link'],
        ['Federation Link Status', 'Check Status', 'Custom: /api/federation-links — status', 'Green=running, Red=error'],
        ['Federation Link Lag', 'Timeseries', 'Messages pending in federation upstream', 'Replication delay'],
        ['Shovel Restart Count', 'Timeseries', 'Number of shovel restarts over time', 'Stability indicator'],
        ['Federation Error Log', 'Log Stream', 'source:rabbitmq "federation" AND ("error" OR "down")', 'Federation failures'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  17. DASHBOARD #14 — SLA & EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('17. Dashboard #14 — SLA & Executive Summary', level=1)
doc.add_paragraph(
    'Purpose: A high-level executive dashboard for management and stakeholders showing SLAs, '
    'availability, key performance indicators, and trend summaries.'
)

doc.add_heading('17.1 Widgets & Layout', level=2)
add_table_with_header(doc,
    ['Widget', 'Type', 'Datadog Metric / Query', 'Threshold / Notes'],
    [
        ['Cluster Availability (30d)', 'Query Value (large)', 'Composite: (uptime of all 3 nodes / total) * 100', 'SLA target: 99.99%'],
        ['Messages Processed Today', 'Query Value', 'sum of delivered + acked today', 'Daily throughput KPI'],
        ['Average End-to-End Latency', 'Query Value', 'APM: publish-to-consume p95 latency', 'SLA target: < 100ms'],
        ['Total Failed Messages (30d)', 'Query Value', 'sum of DLQ messages over 30 days', 'Failure KPI'],
        ['Error Rate %', 'Query Value', '(failed / total) * 100', 'SLA target: < 0.01%'],
        ['Weekly Throughput Trend', 'Timeseries (4 weeks)', 'Weekly message volume', 'Growth trend'],
        ['Top 5 Busiest Queues', 'Top List', 'rabbitmq.queue.messages.deliver.rate top 5', 'Key queue health'],
        ['Active Incidents', 'Query Value', 'Count of P1/P2 alerts currently firing', 'Operational status'],
        ['Node Restart History (90d)', 'Event Timeline', 'All restart events over 90 days', 'Stability trend'],
        ['Capacity Forecast', 'Timeseries + Forecast', 'Memory/Disk usage with linear forecast', '30-day capacity projection'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  18. ALERTING STRATEGY & ESCALATION MATRIX
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('18. Alerting Strategy & Escalation Matrix', level=1)

doc.add_heading('18.1 Severity Levels', level=2)
add_table_with_header(doc,
    ['Severity', 'Response Time', 'Description', 'Examples'],
    [
        ['P1 — Critical', '< 5 minutes', 'Cluster down or data at risk, publishers/consumers blocked', 'Node down, memory alarm, network partition, DLQ critical'],
        ['P2 — Warning', '< 30 minutes', 'Degraded performance or approaching limits', 'Queue depth high, connection leak, memory > 80%'],
        ['P3 — Info', '< 4 hours (business hours)', 'Anomaly detected, cleanup needed, optimization opportunity', 'Idle queues, unroutable messages, stale users'],
        ['P4 — Low', 'Next business day', 'Non-urgent observations', 'Capacity forecast approaching limits'],
    ]
)

doc.add_heading('18.2 Escalation Matrix', level=2)
add_table_with_header(doc,
    ['Time Elapsed', 'P1 Action', 'P2 Action', 'P3 Action'],
    [
        ['0 min', 'PagerDuty on-call + Slack #rabbitmq-critical', 'Slack #rabbitmq-alerts', 'Slack #rabbitmq-info'],
        ['5 min', 'If no ack: escalate to SRE lead', '—', '—'],
        ['15 min', 'If unresolved: page SRE manager', 'If no ack: page on-call', '—'],
        ['30 min', 'If unresolved: bridge call with SRE + App team', 'If unresolved: escalate to SRE lead', 'If no ack: email team lead'],
        ['1 hour', 'If unresolved: VP Engineering notified', 'If unresolved: page SRE manager', '—'],
        ['4 hours', 'Post-incident review scheduled', 'Post-incident review if customer impact', '—'],
    ]
)

doc.add_heading('18.3 Notification Channels', level=2)
add_table_with_header(doc,
    ['Channel', 'Purpose', 'Configuration'],
    [
        ['PagerDuty', 'P1 alerts — immediate paging', 'Integration key: rabbitmq-prod-cluster'],
        ['Slack #rabbitmq-critical', 'P1 alerts — team visibility', 'Datadog Slack integration'],
        ['Slack #rabbitmq-alerts', 'P2 alerts — team awareness', 'Datadog Slack integration'],
        ['Slack #rabbitmq-info', 'P3 alerts — informational', 'Datadog Slack integration'],
        ['Slack #rabbitmq-audit', 'Security / permission changes', 'Datadog Slack integration'],
        ['Email — SRE Team DL', 'Summary digests and P1 follow-up', 'sre-team@company.com'],
        ['Email — Security Team DL', 'Auth failures, permission changes', 'security-team@company.com'],
        ['Datadog Event Stream', 'All events for correlation', 'Automatic'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  19. DATADOG MONITOR DEFINITIONS (Summary)
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('19. Datadog Monitor Definitions — Complete List', level=1)
doc.add_paragraph(
    'Below is the consolidated list of all Datadog monitors to be created for the RabbitMQ cluster.'
)

monitors = [
    ['M-001', 'RMQ Node Down', 'rabbitmq.aliveness == 0', 'P1', '1 min'],
    ['M-002', 'Network Partition Detected', 'rabbitmq.node.partitions > 0', 'P1', '30 sec'],
    ['M-003', 'Memory Alarm Triggered', 'rabbitmq.node.mem_alarm == 1', 'P1', '0 (immediate)'],
    ['M-004', 'Disk Alarm Triggered', 'rabbitmq.node.disk_free_alarm == 1', 'P1', '0 (immediate)'],
    ['M-005', 'Node Restarted', 'Uptime pct_change < -90%', 'P1', '0 (immediate)'],
    ['M-006', 'Multiple Restarts (same node, 24h)', '> 2 restart events', 'P1', '0 (immediate)'],
    ['M-007', 'Blocked Connections', 'blocked_connections > 0', 'P1', '1 min'],
    ['M-008', 'Unacked Messages Growing', 'Monotonic increase for 15 min', 'P1', '15 min'],
    ['M-009', 'DLQ Depth Critical (> 1000)', 'DLQ messages > 1000', 'P1', '0 (immediate)'],
    ['M-010', 'Under-Replicated Quorum Queue', 'online_members < total_members', 'P1', '5 min'],
    ['M-011', 'New Admin User Created', 'Log: set_user_tags + administrator', 'P1', '0 (immediate)'],
    ['M-012', 'TLS Certificate < 7 Days', 'cert_expiry_days < 7', 'P1', '0 (immediate)'],
    ['M-013', 'Zero Throughput (Business Hours)', 'publish + deliver == 0 for 10 min', 'P1', '10 min'],
    ['M-014', 'Queue Depth > 50,000', 'queue.messages > 50000', 'P1', '5 min'],
    ['M-015', 'Queue Depth > 10,000', 'queue.messages > 10000', 'P2', '5 min'],
    ['M-016', 'Queue with 0 Consumers (5 min)', 'consumers == 0 for non-DLQ', 'P2', '5 min'],
    ['M-017', 'Memory > 80%', 'mem_used / mem_limit > 0.80', 'P2', '5 min'],
    ['M-018', 'Disk Free < 5 GB', 'disk_free < 5GB', 'P2', '5 min'],
    ['M-019', 'Publish Rate Drop > 50%', 'Anomaly detection', 'P2', '5 min'],
    ['M-020', 'Deliver Rate Drop > 50%', 'Anomaly detection', 'P2', '5 min'],
    ['M-021', 'High Redelivery Rate', 'redeliver > 50/s per queue', 'P2', '5 min'],
    ['M-022', 'Connection Leak', 'Monotonic connection increase 30 min', 'P2', '30 min'],
    ['M-023', 'Connection Count Anomaly', 'Anomaly detection', 'P2', '5 min'],
    ['M-024', 'Failed Auth > 10/min', 'Log count', 'P2', '1 min'],
    ['M-025', 'Guest Account Enabled', 'guest_enabled == 1', 'P2', '0 (immediate)'],
    ['M-026', 'DLQ Depth Increasing', 'rate > 0 for 10 min', 'P2', '10 min'],
    ['M-027', 'CPU > 80% for 15 min', 'system.cpu.user > 80', 'P2', '15 min'],
    ['M-028', 'Erlang Processes > 85%', 'proc_used / proc_total > 0.85', 'P2', '5 min'],
    ['M-029', 'File Descriptors > 80%', 'fd_used / fd_total > 0.80', 'P2', '5 min'],
    ['M-030', 'TLS Certificate < 30 Days', 'cert_expiry_days < 30', 'P2', '0'],
    ['M-031', 'Non-TLS Connections', 'non_tls_connections > 0', 'P2', '0'],
    ['M-032', 'Shovel Down', 'shovel_state != running', 'P2', '1 min'],
    ['M-033', 'Federation Link Down', 'federation_status != running', 'P2', '1 min'],
    ['M-034', 'High Rejection Rate', 'rejected > 100/s', 'P2', '5 min'],
    ['M-035', 'Poison Messages Detected', 'x-death count > 5', 'P1', '0'],
    ['M-036', 'Queue Count Anomaly', 'Anomaly detection on queue count', 'P3', '15 min'],
    ['M-037', 'Unroutable Messages', 'return_unroutable > 0', 'P3', '5 min'],
    ['M-038', 'Consumer Utilization < 50%', 'consumer_utilisation < 0.5', 'P3', '10 min'],
    ['M-039', 'IO Latency > 10ms', 'io_write_avg_time > 10ms', 'P3', '10 min'],
    ['M-040', 'Frequent Leader Elections', 'raft_term rate > 5/min', 'P2', '5 min'],
]

add_table_with_header(doc,
    ['Monitor ID', 'Name', 'Condition', 'Severity', 'Eval Window'],
    monitors
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  20. RUNBOOK REFERENCES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('20. Runbook References', level=1)
doc.add_paragraph(
    'Each alert should link to the corresponding runbook. Below are the recommended runbooks '
    'and their mapping to monitors.'
)

add_table_with_header(doc,
    ['Runbook', 'Covers Monitors', 'Description'],
    [
        ['RB-001: Node Down Recovery', 'M-001, M-005, M-006', 'Steps to diagnose and recover a failed node, including force-boot procedures'],
        ['RB-002: Network Partition Resolution', 'M-002', 'Identify partition type, choose resolution strategy (autoheal, pause_minority, manual)'],
        ['RB-003: Memory Alarm Response', 'M-003, M-017', 'Identify memory consumers, trigger GC, add memory, restart if needed'],
        ['RB-004: Disk Alarm Response', 'M-004, M-018', 'Free disk space, purge old data, expand volume, restart if needed'],
        ['RB-005: Queue Depth Remediation', 'M-014, M-015, M-016', 'Scale consumers, identify stuck consumers, purge if needed'],
        ['RB-006: DLQ Processing', 'M-009, M-026, M-035', 'Inspect DLQ messages, replay or discard, fix root cause'],
        ['RB-007: Connection Management', 'M-007, M-022, M-023', 'Identify leaking clients, close stuck connections, apply rate limits'],
        ['RB-008: Security Incident Response', 'M-011, M-024, M-025, M-031', 'Investigate auth failures, revoke compromised credentials, audit permissions'],
        ['RB-009: TLS Certificate Renewal', 'M-012, M-030', 'Renew certificates, deploy to all nodes, verify TLS handshake'],
        ['RB-010: Quorum Queue Recovery', 'M-010, M-040', 'Recover under-replicated queues, rebalance leaders, force leader election'],
        ['RB-011: Shovel/Federation Recovery', 'M-032, M-033', 'Restart shovels, check connectivity, verify federation links'],
        ['RB-012: Rolling Restart Procedure', 'M-005', 'Safe rolling restart procedure with queue drain and health checks'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  21. APPENDIX — METRIC REFERENCE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('21. Appendix — Key Metric Reference Table', level=1)
doc.add_paragraph(
    'Quick reference for the most important RabbitMQ metrics available in Datadog.'
)

metrics = [
    ['rabbitmq.overview.queues', 'Gauge', 'Total number of queues', 'Cluster'],
    ['rabbitmq.overview.connections', 'Gauge', 'Total connections', 'Cluster'],
    ['rabbitmq.overview.channels', 'Gauge', 'Total channels', 'Cluster'],
    ['rabbitmq.overview.consumers', 'Gauge', 'Total consumers', 'Cluster'],
    ['rabbitmq.overview.messages.publish_in.rate', 'Rate', 'Messages published per second', 'Cluster'],
    ['rabbitmq.overview.messages.deliver_get.rate', 'Rate', 'Messages delivered per second', 'Cluster'],
    ['rabbitmq.overview.messages.ack.rate', 'Rate', 'Messages acknowledged per second', 'Cluster'],
    ['rabbitmq.overview.messages.redeliver.rate', 'Rate', 'Messages redelivered per second', 'Cluster'],
    ['rabbitmq.queue.messages', 'Gauge', 'Total messages in queue', 'Queue'],
    ['rabbitmq.queue.messages.ready', 'Gauge', 'Messages ready for delivery', 'Queue'],
    ['rabbitmq.queue.messages.unacknowledged', 'Gauge', 'Messages delivered but not acked', 'Queue'],
    ['rabbitmq.queue.consumers', 'Gauge', 'Consumer count per queue', 'Queue'],
    ['rabbitmq.queue.memory', 'Gauge', 'Memory used by queue (bytes)', 'Queue'],
    ['rabbitmq.queue.consumer_utilisation', 'Gauge', 'Consumer utilization (0-1)', 'Queue'],
    ['rabbitmq.exchange.messages.publish_in.rate', 'Rate', 'Messages published to exchange', 'Exchange'],
    ['rabbitmq.exchange.messages.return_unroutable.rate', 'Rate', 'Unroutable messages returned', 'Exchange'],
    ['rabbitmq.node.running', 'Gauge', 'Node is running (1/0)', 'Node'],
    ['rabbitmq.node.uptime', 'Gauge', 'Node uptime in milliseconds', 'Node'],
    ['rabbitmq.node.mem_used', 'Gauge', 'Memory used (bytes)', 'Node'],
    ['rabbitmq.node.mem_limit', 'Gauge', 'Memory high watermark (bytes)', 'Node'],
    ['rabbitmq.node.mem_alarm', 'Gauge', 'Memory alarm active (1/0)', 'Node'],
    ['rabbitmq.node.disk_free', 'Gauge', 'Free disk space (bytes)', 'Node'],
    ['rabbitmq.node.disk_free_alarm', 'Gauge', 'Disk alarm active (1/0)', 'Node'],
    ['rabbitmq.node.fd_used', 'Gauge', 'File descriptors used', 'Node'],
    ['rabbitmq.node.fd_total', 'Gauge', 'File descriptors available', 'Node'],
    ['rabbitmq.node.sockets_used', 'Gauge', 'Sockets used', 'Node'],
    ['rabbitmq.node.proc_used', 'Gauge', 'Erlang processes used', 'Node'],
    ['rabbitmq.node.proc_total', 'Gauge', 'Erlang processes available', 'Node'],
    ['rabbitmq.node.partitions', 'Gauge', 'Network partitions seen by node', 'Node'],
    ['rabbitmq.node.io_read_avg_time', 'Gauge', 'Avg IO read time (ms)', 'Node'],
    ['rabbitmq.node.io_write_avg_time', 'Gauge', 'Avg IO write time (ms)', 'Node'],
    ['rabbitmq.aliveness', 'Service Check', 'Node aliveness check', 'Node'],
    ['rabbitmq_raft_term_total', 'Counter', 'Raft term (leader elections)', 'Prometheus'],
    ['rabbitmq_raft_log_commit_index', 'Counter', 'Raft committed log entries', 'Prometheus'],
    ['rabbitmq_erlang_gc_runs_total', 'Counter', 'Erlang GC runs', 'Prometheus'],
    ['rabbitmq_erlang_gc_reclaimed_bytes_total', 'Counter', 'Bytes reclaimed by GC', 'Prometheus'],
    ['rabbitmq_erlang_scheduler_utilization', 'Gauge', 'Erlang scheduler CPU usage', 'Prometheus'],
]

add_table_with_header(doc,
    ['Metric Name', 'Type', 'Description', 'Scope'],
    metrics
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  FINAL — DOCUMENT APPROVAL
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Document Approval', level=1)
add_table_with_header(doc,
    ['Role', 'Name', 'Date', 'Signature'],
    [
        ['Monitoring Architect', '________________', '________', '________________'],
        ['SRE Lead', '________________', '________', '________________'],
        ['Infrastructure Manager', '________________', '________', '________________'],
        ['Security Lead', '________________', '________', '________________'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = '/Users/tejasodanapalli/rabbitmq/RabbitMQ/docs/monitoring/RabbitMQ_Datadog_Monitoring_Dashboard_Strategy.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Pages (estimated): ~45-50 pages')
print(f'Dashboards: 14')
print(f'Monitors: 40')
print(f'Runbooks: 12')
